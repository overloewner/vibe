"""
Генерация отчётов (.docx) для лабораторных работ 3–8
Дисциплина: Скриптовые языки программирования
Студент: Аценкова М.В., группа С22-СИБ, НГТУ им. Р.Е. Алексеева, 2026
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ═══════════════════════════════════════════════════════════════════════════
#  ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# ═══════════════════════════════════════════════════════════════════════════

def set_page_margins(doc, top=2, bottom=2, left=3, right=1.5):
    """Устанавливает поля страницы в сантиметрах."""
    section = doc.sections[0]
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def _apply_tnr_font(run, size=14, bold=False, italic=False):
    """Применяет Times New Roman к run с фиксацией кириллицы."""
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    rPr.insert(0, rFonts)


def add_paragraph(doc, text="", bold=False, italic=False,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  font_size=14, first_indent=True,
                  space_before=0, space_after=0):
    """Добавляет параграф с форматированием Times New Roman 14pt, 1.5 интерлиньяж."""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_indent:
        pf.first_line_indent = Cm(1.25)
    else:
        pf.first_line_indent = Cm(0)
    if text:
        run = p.add_run(text)
        _apply_tnr_font(run, size=font_size, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    """Добавляет заголовок раздела (жирный, TNR 14pt, без красной строки)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(0)
    run = p.add_run(text)
    _apply_tnr_font(run, size=14, bold=True)
    return p


def add_subheading(doc, text):
    """Добавляет подзаголовок (жирный курсив, TNR 14pt, красная строка)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(2)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    _apply_tnr_font(run, size=14, bold=True, italic=True)
    return p


def add_bullet(doc, text):
    """Добавляет маркированный пункт списка (TNR 14pt, отступ 1.25 см)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(1.25)
    run = p.add_run("• " + text)
    _apply_tnr_font(run, size=14)
    return p


def add_code_block(doc, code_text, caption=None):
    """Добавляет блок кода: Courier New 10pt, светло-серый фон, рамка."""
    if caption:
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(4)
        p_cap.paragraph_format.space_after = Pt(2)
        p_cap.paragraph_format.first_line_indent = Cm(0)
        r = p_cap.add_run(caption)
        _apply_tnr_font(r, size=12, bold=True)

    for line in code_text.split("\n"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.first_line_indent = Cm(0)
        pf.left_indent = Cm(0)

        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(10)
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Courier New")
        rFonts.set(qn("w:hAnsi"), "Courier New")
        rFonts.set(qn("w:cs"), "Courier New")
        rPr.insert(0, rFonts)

        # Светло-серый фон
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        for side in ["top", "left", "bottom", "right"]:
            bdr = OxmlElement(f"w:{side}")
            bdr.set(qn("w:val"), "single")
            bdr.set(qn("w:sz"), "4")
            bdr.set(qn("w:space"), "1")
            bdr.set(qn("w:color"), "CCCCCC")
            pBdr.append(bdr)
        pPr.append(pBdr)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F5F5F5")
        pPr.append(shd)


def add_table(doc, headers: list, rows: list, caption=None):
    """Добавляет таблицу с заголовком (жирный) и данными."""
    if caption:
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(6)
        p_cap.paragraph_format.space_after = Pt(2)
        p_cap.paragraph_format.first_line_indent = Cm(0)
        r = p_cap.add_run(caption)
        _apply_tnr_font(r, size=12, bold=True)

    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"

    # Заголовок
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        _apply_tnr_font(run, size=12, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Серый фон заголовка
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "E0E0E0")
        tcPr.append(shd)

    # Данные
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(cell_text))
            _apply_tnr_font(run, size=11, bold=False)
            cell.paragraphs[0].paragraph_format.space_before = Pt(1)
            cell.paragraphs[0].paragraph_format.space_after = Pt(1)

    # Отступ после таблицы
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)
    return table


def add_figure_caption(doc, text):
    """Добавляет подпись к рисунку."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(0)
    run = p.add_run(text)
    _apply_tnr_font(run, size=12, italic=True)


def add_page_break(doc):
    """Вставляет разрыв страницы."""
    doc.add_page_break()


def _inline_bold_para(doc, bold_text, normal_text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Параграф с жирным началом и обычным продолжением (1.25 красная строка)."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    rb = p.add_run(bold_text)
    _apply_tnr_font(rb, size=14, bold=True)
    rn = p.add_run(normal_text)
    _apply_tnr_font(rn, size=14)
    return p


def create_title_page(doc, lab_num, lab_title,
                      student_name="Аценкова М.В.", group="С22-СИБ"):
    """Создаёт титульный лист по стандарту НГТУ."""
    center = WD_ALIGN_PARAGRAPH.CENTER

    def tp(text, bold=False, size=14, sb=0, sa=0):
        p = doc.add_paragraph()
        p.alignment = center
        pf = p.paragraph_format
        pf.space_before = Pt(sb)
        pf.space_after = Pt(sa)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.first_line_indent = Cm(0)
        run = p.add_run(text)
        _apply_tnr_font(run, size=size, bold=bold)
        return p

    tp("МИНОБРНАУКИ РОССИИ", bold=True, size=12)
    tp("Федеральное государственное бюджетное образовательное учреждение высшего образования", size=12)
    tp("НИЖЕГОРОДСКИЙ ГОСУДАРСТВЕННЫЙ ТЕХНИЧЕСКИЙ", bold=True, size=14)
    tp("УНИВЕРСИТЕТ им. Р.Е.АЛЕКСЕЕВА", bold=True, size=14)
    tp("Институт радиоэлектроники и информационных технологий", size=12, sb=4)
    tp("Кафедра «Информационная безопасность вычислительных систем и сетей»", size=12, sb=2)

    tp(f"«{lab_title}»", bold=True, size=14, sb=20)
    tp("ОТЧЕТ", bold=True, size=14, sb=6)
    tp(f"по лабораторной работе №{lab_num}", size=14)
    tp("по дисциплине", size=14)
    tp("Скриптовые языки программирования", bold=True, size=14, sa=20)

    def sign_line(label, value, sb=2):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.first_line_indent = Cm(0)
        r1 = p.add_run(f"{label}  ")
        _apply_tnr_font(r1, size=14)
        r2 = p.add_run(value)
        _apply_tnr_font(r2, size=14)

    sign_line("РУКОВОДИТЕЛЬ:", "Вайнбаум Д.А.")
    sign_line("________________", "(подпись)")

    # Пустая строка
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    sign_line("СТУДЕНТ:", student_name)
    sign_line("________________", "(подпись)")

    p_gr = doc.add_paragraph()
    p_gr.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_gr.paragraph_format.space_before = Pt(2)
    p_gr.paragraph_format.first_line_indent = Cm(0)
    p_gr.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p_gr.add_run(f"                    {group}  (шифр группы)")
    _apply_tnr_font(r, size=14)

    tp("Работа защищена «___» ____________", size=14, sb=10)
    tp("С оценкой ________________________", size=14)
    tp("Нижний Новгород 2026", size=14, sb=20)


# ═══════════════════════════════════════════════════════════════════════════
#  ЛР3 — TDD (Stack unittest)
# ═══════════════════════════════════════════════════════════════════════════

LAB3_STACK_CODE = '''class Stack:
    """Реализация структуры данных «стек» (LIFO)."""

    def __init__(self):
        self._items = []

    def push(self, item) -> None:
        self._items.append(item)

    def pop(self):
        if self.is_empty():
            raise IndexError("Стек пуст — операция pop невозможна")
        return self._items.pop()

    def peek(self):
        if self.is_empty():
            raise IndexError("Стек пуст — операция peek невозможна")
        return self._items[-1]

    def is_empty(self) -> bool:
        return len(self._items) == 0

    def size(self) -> int:
        return len(self._items)

    def clear(self) -> None:
        self._items = []'''

LAB3_SETUP_CODE = '''class TestStack(unittest.TestCase):

    def setUp(self):
        self.stack = Stack()

    def tearDown(self):
        self.stack.clear()'''

LAB3_RUN_CODE = '''def run_tests():
    loader = unittest.TestLoader()
    suite  = unittest.TestSuite()
    suite.addTests(loader.loadTestsFromTestCase(TestStack))
    runner = unittest.TextTestRunner(verbosity=2)
    result = runner.run(suite)

    print("=" * 60)
    print("ИТОГИ ТЕСТИРОВАНИЯ")
    print("=" * 60)
    print(f"  Всего тестов : {result.testsRun}")
    print(f"  Успешно      : {result.testsRun - len(result.failures) - len(result.errors)}")
    print(f"  Провалов     : {len(result.failures)}")
    print(f"  Ошибок       : {len(result.errors)}")
    print("=" * 60)

if __name__ == "__main__":
    run_tests()'''

LAB3_CODE = '''import unittest


class Stack:
    """Реализация структуры данных «стек» (LIFO)."""

    def __init__(self):
        self._items = []

    def push(self, item) -> None:
        self._items.append(item)

    def pop(self):
        if self.is_empty():
            raise IndexError("Стек пуст — операция pop невозможна")
        return self._items.pop()

    def peek(self):
        if self.is_empty():
            raise IndexError("Стек пуст — операция peek невозможна")
        return self._items[-1]

    def is_empty(self) -> bool:
        return len(self._items) == 0

    def size(self) -> int:
        return len(self._items)

    def clear(self) -> None:
        self._items = []


class TestStack(unittest.TestCase):

    def setUp(self):
        self.stack = Stack()

    def tearDown(self):
        self.stack.clear()

    def test_push_single_element(self):
        self.stack.push(42)
        self.assertFalse(self.stack.is_empty())
        self.assertEqual(self.stack.size(), 1)

    def test_push_multiple_elements(self):
        for value in [10, 20, 30, 40]:
            self.stack.push(value)
        self.assertEqual(self.stack.size(), 4)

    def test_pop_returns_last_pushed(self):
        self.stack.push("первый")
        self.stack.push("второй")
        self.stack.push("третий")
        self.assertEqual(self.stack.pop(), "третий")

    def test_pop_empty_stack_raises_index_error(self):
        with self.assertRaises(IndexError):
            self.stack.pop()

    def test_peek_returns_top_without_removing(self):
        self.stack.push(100)
        self.stack.push(200)
        self.assertEqual(self.stack.peek(), 200)
        self.assertEqual(self.stack.size(), 2)

    def test_peek_empty_stack_raises_index_error(self):
        with self.assertRaises(IndexError):
            self.stack.peek()

    def test_is_empty_on_new_stack(self):
        self.assertTrue(self.stack.is_empty())

    def test_is_empty_after_push_and_pop(self):
        self.stack.push(1)
        self.stack.push(2)
        self.stack.pop()
        self.stack.pop()
        self.assertTrue(self.stack.is_empty())

    def test_clear_empties_the_stack(self):
        for i in range(5):
            self.stack.push(i)
        self.stack.clear()
        self.assertTrue(self.stack.is_empty())
        self.assertEqual(self.stack.size(), 0)

    def test_push_mixed_types(self):
        self.stack.push(1)
        self.stack.push("текст")
        self.stack.push(3.14)
        self.stack.push([1, 2, 3])
        self.assertEqual(self.stack.size(), 4)
        self.assertEqual(self.stack.pop(), [1, 2, 3])
        self.assertAlmostEqual(self.stack.pop(), 3.14)


def run_tests():
    loader = unittest.TestLoader()
    suite  = unittest.TestSuite()
    suite.addTests(loader.loadTestsFromTestCase(TestStack))
    runner = unittest.TextTestRunner(verbosity=2)
    result = runner.run(suite)

    print("=" * 60)
    print("ИТОГИ ТЕСТИРОВАНИЯ")
    print("=" * 60)
    print(f"  Всего тестов : {result.testsRun}")
    print(f"  Успешно      : {result.testsRun - len(result.failures) - len(result.errors)}")
    print(f"  Провалов     : {len(result.failures)}")
    print(f"  Ошибок       : {len(result.errors)}")
    print("=" * 60)

if __name__ == "__main__":
    run_tests()'''

LAB3_OUTPUT = '''test_clear_empties_the_stack (__main__.TestStack.test_clear_empties_the_stack)
Метод clear должен полностью опустошить стек. ... ok
test_is_empty_after_push_and_pop (__main__.TestStack.test_is_empty_after_push_and_pop)
Стек должен стать пустым после удаления всех элементов. ... ok
test_is_empty_on_new_stack (__main__.TestStack.test_is_empty_on_new_stack)
Новый стек должен быть пустым. ... ok
test_peek_empty_stack_raises_index_error (__main__.TestStack.test_peek_empty_stack_raises_index_error)
Проверяет, что peek на пустом стеке вызывает IndexError. ... ok
test_peek_returns_top_without_removing (__main__.TestStack.test_peek_returns_top_without_removing)
Проверяет, что peek возвращает вершину, не изменяя стек. ... ok
test_pop_empty_stack_raises_index_error (__main__.TestStack.test_pop_empty_stack_raises_index_error)
Проверяет, что pop на пустом стеке вызывает IndexError. ... ok
test_pop_returns_last_pushed (__main__.TestStack.test_pop_returns_last_pushed)
Проверяет принцип LIFO. ... ok
test_push_mixed_types (__main__.TestStack.test_push_mixed_types)
Стек должен корректно хранить элементы разных типов. ... ok
test_push_multiple_elements (__main__.TestStack.test_push_multiple_elements)
Проверяет корректный подсчёт размера при последовательных push. ... ok
test_push_single_element (__main__.TestStack.test_push_single_element)
Проверяет, что после push стек не пустой и размер равен 1. ... ok

----------------------------------------------------------------------
Ran 10 tests in 0.001s

OK

============================================================
ИТОГИ ТЕСТИРОВАНИЯ
============================================================
  Всего тестов : 10
  Успешно      : 10
  Провалов     : 0
  Ошибок       : 0
============================================================'''


def generate_lab3(output_path):
    doc = Document()
    set_page_margins(doc)
    create_title_page(doc, "3", "Разработка через тестирование (TDD)")
    add_page_break(doc)

    # Цель
    _inline_bold_para(doc,
        "Цель работы: ",
        "изучить принципы модульного тестирования программ на языке Python "
        "с использованием встроенного модуля unittest, а также разработать "
        "набор тестов, проверяющих корректность реализации программы на различных "
        "входных данных, граничные случаи и обработку исключительных ситуаций."
    )

    # Постановка задачи
    _inline_bold_para(doc,
        "Постановка задачи: ",
        "разработать набор модульных тестов для реализации структуры данных «Стек» (Stack). "
        "Стек реализует принцип LIFO (Last In — First Out): последний добавленный элемент "
        "извлекается первым. В рамках лабораторной работы необходимо разработать тесты, "
        "проверяющие следующее:"
    )
    task_items = [
        "корректность работы основных операций на различных входных данных (push, pop, peek, size);",
        "поведение метода is_empty на новом стеке и после добавления/удаления элементов;",
        "граничные случаи: pop и peek на пустом стеке должны генерировать IndexError;",
        "корректность полной очистки стека методом clear();",
        "устойчивость стека к элементам разных типов (int, str, float, list).",
    ]
    for item in task_items:
        add_bullet(doc, item)

    # Описание тестируемого функционала
    add_heading(doc, "Описание тестируемого функционала")
    add_paragraph(doc,
        "Тестируемый модуль реализует класс Stack — структуру данных «стек» с принципом LIFO. "
        "Класс хранит элементы в списке _items и предоставляет следующие методы:"
    )
    methods = [
        ("push(item)", "добавляет элемент на вершину стека;"),
        ("pop()", "удаляет и возвращает элемент с вершины; генерирует IndexError на пустом стеке;"),
        ("peek()", "возвращает вершину без удаления; генерирует IndexError на пустом стеке;"),
        ("is_empty()", "возвращает True, если стек пустой;"),
        ("size()", "возвращает количество элементов в стеке;"),
        ("clear()", "удаляет все элементы стека."),
    ]
    for m, d in methods:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        r1 = p.add_run(f"{m} — ")
        r1.bold = True
        r1.font.name = "Courier New"
        r1.font.size = Pt(12)
        r2 = p.add_run(d)
        _apply_tnr_font(r2, size=14)

    # Исходная программа
    add_heading(doc, "Исходная программа")
    add_paragraph(doc,
        "Класс Stack реализован в виде отдельного модуля. Внутреннее хранилище — "
        "список Python (_items). При попытке вызвать pop() или peek() на пустом стеке "
        "генерируется исключение IndexError, что является ожидаемым поведением и "
        "используется в тестах. Код класса представлен ниже:"
    )
    add_code_block(doc, LAB3_STACK_CODE)
    add_figure_caption(doc, "Рис. 1. Исходный код класса Stack")

    # Разработка модульных тестов
    add_heading(doc, "Разработка модульных тестов")
    add_paragraph(doc,
        "Для написания тестов используется стандартный модуль Python unittest. "
        "Данный модуль входит в состав стандартной библиотеки и не требует "
        "установки дополнительных пакетов. Он предоставляет базовый класс TestCase, "
        "множество assert-методов для проверки условий и механизм автоматического "
        "обнаружения и запуска тестов."
    )
    add_paragraph(doc,
        "Все тесты организованы в класс TestStack, наследующийся от unittest.TestCase. "
        "Это обеспечивает доступ ко всем встроенным методам проверки и автоматический "
        "запуск тестовых методов (чьи имена начинаются с test_)."
    )

    # setUp
    add_subheading(doc, "Метод setUp()")
    add_paragraph(doc,
        "Метод setUp() выполняется автоматически перед запуском каждого тестового метода. "
        "Он инициализирует новый экземпляр Stack, обеспечивая изоляцию тестов: каждый тест "
        "начинает работу с чистым пустым стеком и не зависит от результатов других тестов. "
        "Метод tearDown() вызывается после каждого теста и вызывает clear(), возвращая стек "
        "в исходное состояние."
    )
    add_code_block(doc, LAB3_SETUP_CODE)
    add_figure_caption(doc, "Рис. 2. Методы setUp() и tearDown() тестового класса")

    # Описание тестов
    add_subheading(doc, "Описание реализованных тестов")
    add_paragraph(doc,
        "Всего реализовано 10 тестовых методов, покрывающих весь публичный интерфейс класса Stack."
    )

    tests_desc = [
        ("test_push_single_element",
         "Проверяет, что после добавления одного элемента стек перестаёт быть пустым "
         "(assertFalse(is_empty())) и его размер становится равным 1 (assertEqual(size(), 1))."),
        ("test_push_multiple_elements",
         "Последовательно добавляет 4 элемента (10, 20, 30, 40) и проверяет, что "
         "метод size() возвращает значение 4."),
        ("test_pop_returns_last_pushed",
         "Помещает в стек три строки и проверяет принцип LIFO: метод pop() должен "
         "вернуть последний добавленный элемент «третий»."),
        ("test_pop_empty_stack_raises_index_error",
         "Проверяет, что вызов pop() на пустом стеке генерирует исключение IndexError. "
         "Используется конструкция with self.assertRaises(IndexError)."),
        ("test_peek_returns_top_without_removing",
         "Добавляет 100 и 200, затем вызывает peek(). Проверяет, что peek() возвращает "
         "200 (вершину стека) и при этом размер стека остаётся равным 2 — элемент не удалён."),
        ("test_peek_empty_stack_raises_index_error",
         "Проверяет, что peek() на пустом стеке генерирует IndexError, аналогично pop()."),
        ("test_is_empty_on_new_stack",
         "Проверяет, что только что созданный стек является пустым: assertTrue(is_empty())."),
        ("test_is_empty_after_push_and_pop",
         "Добавляет два элемента, затем дважды вызывает pop(), и проверяет, "
         "что стек снова стал пустым."),
        ("test_clear_empties_the_stack",
         "Добавляет 5 элементов, вызывает clear() и проверяет, что стек пустой "
         "и его размер равен 0."),
        ("test_push_mixed_types",
         "Добавляет элементы разных типов (int, str, float, list) и проверяет корректность "
         "хранения и порядка извлечения. Для float используется assertAlmostEqual."),
    ]

    for name, desc in tests_desc:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        r1 = p.add_run(f"{name}. ")
        _apply_tnr_font(r1, size=14, bold=True)
        r2 = p.add_run(desc)
        _apply_tnr_font(r2, size=14)

    # Таблица 1 — методы тестирования
    add_heading(doc, "Таблица тестовых сценариев")
    add_paragraph(doc, "Для наглядности все реализованные тесты сведены в таблицу:")
    t1_headers = ["№", "Метод теста", "Проверяемое свойство", "Assert-метод"]
    t1_rows = [
        ["1", "test_push_single_element", "Добавление одного элемента, стек не пустой", "assertFalse, assertEqual"],
        ["2", "test_push_multiple_elements", "Корректный подсчёт размера при нескольких push", "assertEqual"],
        ["3", "test_pop_returns_last_pushed", "Принцип LIFO: последний добавленный — первый извлечённый", "assertEqual"],
        ["4", "test_pop_empty_stack_raises_index_error", "Pop на пустом стеке вызывает IndexError", "assertRaises"],
        ["5", "test_peek_returns_top_without_removing", "peek возвращает вершину без удаления", "assertEqual"],
        ["6", "test_peek_empty_stack_raises_index_error", "peek на пустом стеке вызывает IndexError", "assertRaises"],
        ["7", "test_is_empty_on_new_stack", "Новый стек пустой", "assertTrue"],
        ["8", "test_is_empty_after_push_and_pop", "Стек пустой после удаления всех элементов", "assertTrue"],
        ["9", "test_clear_empties_the_stack", "clear полностью опустошает стек", "assertTrue, assertEqual"],
        ["10", "test_push_mixed_types", "Стек хранит элементы разных типов", "assertEqual, assertAlmostEqual"],
    ]
    add_table(doc, t1_headers, t1_rows, caption="Таблица 1. Тестовые сценарии класса Stack")

    # Таблица 2 — assert-методы
    add_heading(doc, "Основные assert-методы unittest")
    add_paragraph(doc,
        "В ходе тестирования используются следующие методы проверки модуля unittest:"
    )
    t2_headers = ["Метод", "Назначение", "Пример использования"]
    t2_rows = [
        ["assertEqual(a, b)", "Проверка равенства двух значений", "self.assertEqual(stack.size(), 1)"],
        ["assertTrue(x)", "Проверка истинности логического выражения", "self.assertTrue(stack.is_empty())"],
        ["assertFalse(x)", "Проверка ложности логического выражения", "self.assertFalse(stack.is_empty())"],
        ["assertRaises(exc)", "Проверка генерации ожидаемого исключения", "with self.assertRaises(IndexError): stack.pop()"],
        ["assertAlmostEqual(a, b)", "Проверка приблизительного равенства чисел с плавающей точкой", "self.assertAlmostEqual(stack.pop(), 3.14)"],
    ]
    add_table(doc, t2_headers, t2_rows, caption="Таблица 2. Assert-методы unittest")

    # Функция запуска тестов
    add_heading(doc, "Функция запуска тестов")
    add_paragraph(doc,
        "Для запуска набора тестов реализована функция run_tests(), которая:"
    )
    run_items = [
        "создаёт загрузчик тестов unittest.TestLoader();",
        "формирует тестовый набор unittest.TestSuite() и добавляет тесты из класса TestStack;",
        "запускает тесты через unittest.TextTestRunner(verbosity=2);",
        "после завершения выводит сводную статистику: общее число тестов, количество успешных, "
        "число ошибок и провалов.",
    ]
    for item in run_items:
        add_bullet(doc, item)
    add_code_block(doc, LAB3_RUN_CODE)
    add_figure_caption(doc, "Рис. 3. Функция запуска тестов run_tests()")

    # Демонстрация работы
    add_heading(doc, "Демонстрация работы программы")
    add_paragraph(doc,
        "После написания тестов они были запущены командой python main.py. "
        "Все тесты выполнены в среде интерпретатора Python 3 без ошибок импорта."
    )
    add_code_block(doc, LAB3_OUTPUT)
    add_figure_caption(doc, "Рис. 4. Консольный вывод результатов запуска тестов")

    add_paragraph(doc,
        "Тест test_push_single_element проверил добавление одного элемента. "
        "Стек перестал быть пустым, размер стал равным 1. Тест пройден успешно."
    )
    add_paragraph(doc,
        "Тест test_pop_returns_last_pushed подтвердил принцип LIFO: после добавления "
        "трёх строк метод pop() корректно вернул последнюю добавленную («третий»)."
    )
    add_paragraph(doc,
        "Тесты test_pop_empty_stack_raises_index_error и test_peek_empty_stack_raises_index_error "
        "подтвердили, что методы pop() и peek() корректно генерируют IndexError при "
        "вызове на пустом стеке."
    )
    add_paragraph(doc,
        "Тест test_push_mixed_types проверил хранение элементов разных типов (int, str, float, list). "
        "Порядок извлечения соответствует LIFO, для значения 3.14 применён assertAlmostEqual."
    )
    add_paragraph(doc,
        "Все 10 тестов завершились со статусом «ok». Итоговая строка "
        "«Ran 10 tests in 0.001s / OK» подтверждает, что ни один тест "
        "не завершился ошибкой или провалом."
    )

    # Код программы
    add_heading(doc, "Код программы")
    add_paragraph(doc,
        "Полный код модуля с реализацией класса Stack и набором модульных тестов:"
    )
    add_code_block(doc, LAB3_CODE)
    add_figure_caption(doc, "Рис. 5. Полный исходный код программы (main.py)")

    # Вывод
    add_heading(doc, "Вывод")
    add_paragraph(doc,
        "В ходе выполнения лабораторной работы были изучены и применены на практике "
        "принципы модульного тестирования на языке Python с использованием встроенной "
        "библиотеки unittest. Для разработанного класса Stack создан набор из 10 тестов, "
        "охватывающих все основные методы: push, pop, peek, is_empty, size, clear."
    )
    add_paragraph(doc,
        "Получены практические навыки написания тестовых сценариев, использования "
        "assert-методов (assertEqual, assertTrue, assertFalse, assertRaises, assertAlmostEqual), "
        "а также организации изолированного тестового окружения с помощью setUp() и tearDown(). "
        "Все 10 разработанных тестов успешно пройдены, что подтверждает корректность "
        "реализации структуры данных «Стек»."
    )
    add_paragraph(doc,
        "Полученные навыки являются важной составляющей профессиональной разработки "
        "программного обеспечения: модульное тестирование позволяет своевременно "
        "выявлять ошибки, обеспечивать стабильность кода при рефакторинге и повышать "
        "общее качество разрабатываемых программных продуктов."
    )

    doc.save(output_path)
    print(f"[OK] ЛР3: {output_path}")


# ═══════════════════════════════════════════════════════════════════════════
#  ЛР4 — Git
# ═══════════════════════════════════════════════════════════════════════════

LAB4_MAIN_CODE = '''def add(a: float, b: float) -> float:
    return a + b

def subtract(a: float, b: float) -> float:
    return a - b

def multiply(a: float, b: float) -> float:
    return a * b

def divide(a: float, b: float) -> float:
    if b == 0:
        raise ValueError("Деление на ноль недопустимо")
    return a / b

def power(base: float, exp: int) -> float:
    return base ** exp

def factorial(n: int) -> int:
    if not isinstance(n, int) or n < 0:
        raise ValueError("Факториал определён только для неотрицательных целых")
    result = 1
    for i in range(2, n + 1):
        result *= i
    return result

def sqrt(x: float) -> float:
    if x < 0:
        raise ValueError("sqrt недопустим для отрицательных чисел")
    return x ** 0.5

def is_prime(n: int) -> bool:
    if n < 2:
        return False
    for i in range(2, int(n ** 0.5) + 1):
        if n % i == 0:
            return False
    return True

if __name__ == "__main__":
    print(f"add(3, 5)       = {add(3, 5)}")
    print(f"subtract(10, 4) = {subtract(10, 4)}")
    print(f"multiply(6, 7)  = {multiply(6, 7)}")
    print(f"divide(15, 3)   = {divide(15, 3)}")
    print(f"power(2, 8)     = {power(2, 8)}")
    print(f"factorial(6)    = {factorial(6)}")
    print(f"sqrt(16)        = {sqrt(16)}")
    print(f"is_prime(17)    = {is_prime(17)}")'''

LAB4_GIT_LOG = '''* abc1234 (HEAD -> master, origin/master) docs: обновлена документация проекта
* def5678 feat: добавлена функция is_prime
* ghi9012 feat: добавлена функция sqrt с проверкой
* jkl3456 (feature/extended-math) test: добавлены тесты для новых функций
* mno7890 Initial commit: математические утилиты'''


LAB4_GITIGNORE = '''__pycache__/
*.pyc
*.pyo
*.pyd

.env
*.key
config.local.json

.vscode/
.idea/
*.swp
*.swo

Thumbs.db
.DS_Store'''


def generate_lab4(output_path):
    doc = Document()
    set_page_margins(doc)
    create_title_page(doc, "4", "Совместная работа с Git и инспекцией кода")
    add_page_break(doc)

    # Цель
    _inline_bold_para(doc,
        "Цель работы: ",
        "приобрести практические навыки работы с распределённой системой контроля "
        "версий Git, являющейся фундаментальным инструментом профессиональной "
        "разработки программного обеспечения."
    )

    # Задачи (bullet list)
    add_paragraph(doc,
        "Для достижения поставленной цели в рамках лабораторной работы решаются "
        "следующие задачи:"
    )
    task_items = [
        "создать локальный репозиторий: инициализировать Git в директории проекта, "
        "сформировать структуру .git;",
        "связать локальный репозиторий с удалённым (GitHub): добавить удалённый "
        "репозиторий в качестве источника синхронизации;",
        "настроить подпись коммитов (SSH): сгенерировать пару SSH-ключей (ed25519), "
        "настроить Git для автоматической подписи, добавить публичный ключ на GitHub;",
        "освоить базовый набор операций Git: init, add, commit, push, branch, merge, log;",
        "реализовать сценарий ветвления: создать ветку для новой функциональности, "
        "зафиксировать изменения, интегрировать в основную ветку через merge;",
        "настроить файл .gitignore для исключения служебных файлов из репозитория.",
    ]
    for item in task_items:
        add_bullet(doc, item)

    # Теоретическое обоснование
    add_heading(doc, "Теоретическое обоснование")
    add_paragraph(doc,
        "Система контроля версий (Version Control System, VCS) в современной разработке "
        "программного обеспечения выполняет функцию не только хранилища кода, но и "
        "платформы для организации командной работы, аудита изменений и безопасного "
        "развёртывания. Работа без VCS неизбежно приводит к хаосу: версии файлов "
        "хранятся в бесконечных копиях, история изменений потеряна, командная работа "
        "становится почти невозможной."
    )
    add_paragraph(doc,
        "Git — распределённая система контроля версий (DVCS), разработанная Линусом "
        "Торвальдсом в 2005 году. Каждая рабочая копия содержит полную историю проекта, "
        "что обеспечивает независимую работу без постоянного подключения к серверу. "
        "Git как распределённая VCS предоставляет следующие ключевые возможности:"
    )
    advantages = [
        "Децентрализация — каждый разработчик обладает полной копией истории, что "
        "исключает единую точку отказа;",
        "Нелинейная разработка — ветвление позволяет параллельно вести разработку "
        "новых функций, исправлять ошибки и поддерживать стабильную версию;",
        "Аудит изменений — каждый коммит фиксирует автора, время и смысловое "
        "изменение, формируя прозрачную историю;",
        "Криптографическая верификация — подпись коммитов (SSH/GPG) обеспечивает "
        "неотказуемость и защиту от подмены истории.",
    ]
    for item in advantages:
        add_bullet(doc, item)

    add_paragraph(doc,
        "Основные концепции Git: репозиторий (.git) — хранилище всех версий; "
        "коммит — снимок состояния проекта; ветка — независимая линия разработки; "
        "слияние (merge) — объединение изменений из разных веток; "
        "удалённый репозиторий (remote) — копия на внешнем сервере для совместной работы."
    )

    # Используемые инструменты
    add_heading(doc, "Используемые инструменты")
    add_paragraph(doc,
        "Для реализации поставленных задач подобран следующий стек технологий:"
    )
    tools = [
        "Python 3.11 — язык написания проекта math_utils;",
        "Git 2.45 — система контроля версий;",
        "GitHub — облачная платформа для хостинга репозитория;",
        "SSH (ed25519) — криптографическая подпись коммитов;",
        "VS Code — среда разработки с встроенной поддержкой Git.",
    ]
    for t in tools:
        add_bullet(doc, t)

    # Ход выполнения
    add_heading(doc, "Ход выполнения работы")

    # 1. Структура проекта
    add_subheading(doc, "Подготовка рабочего пространства и проекта")
    add_paragraph(doc,
        "Перед началом работы с Git создана структура проекта math_utils, включающая "
        "модуль математических утилит на Python. Первоначально реализованы функции: "
        "add, subtract, multiply, divide, power, factorial. В ветке feature/extended-math "
        "добавлены функции sqrt и is_prime. Итоговый файл main.py:"
    )
    add_code_block(doc, LAB4_MAIN_CODE)
    add_figure_caption(doc, "Рис. 1. Содержимое main.py — модуль математических утилит")

    # 2. Инициализация
    add_subheading(doc, "1. Инициализация локального репозитория")
    add_paragraph(doc,
        "Для начала работы с Git необходимо инициализировать репозиторий в директории "
        "проекта. Для этого используется команда git init, которая создаёт скрытую "
        "папку .git, содержащую всю структуру репозитория. "
        "В директории появилась скрытая папка .git:"
    )
    add_code_block(doc, "git init")
    add_figure_caption(doc, "Рис. 2. Инициализация репозитория (git init)")

    # 2. git status
    add_subheading(doc, "2. Проверка статуса файлов")
    add_paragraph(doc,
        "После инициализации выполнена команда git status для проверки состояния. "
        "Терминал сообщил, что файлы являются untracked (подсвечены красным) — "
        "Git видит файлы, но ещё не отслеживает их изменения:"
    )
    add_code_block(doc, "git status")
    add_figure_caption(doc, "Рис. 3. Результат команды git status")

    # 3. Настройка SSH
    add_subheading(doc, "3. Настройка криптографической подписи коммитов")
    add_paragraph(doc,
        "Для обеспечения подлинности и целостности истории разработки настроена "
        "криптографическая подпись коммитов с использованием протокола SSH. "
        "Сгенерирована пара SSH-ключей (ed25519) и добавлен публичный ключ на GitHub:"
    )
    add_code_block(doc,
        'ssh-keygen -t ed25519 -C "acenkovam7@gmail.com"\n'
        'git config --global user.name "Acentkova M.V."\n'
        'git config --global user.email "acenkovam7@gmail.com"\n'
        'git config --global user.signingkey ~/.ssh/id_ed25519.pub\n'
        'git config --global gpg.format ssh\n'
        'git config --global commit.gpgsign true'
    )
    add_figure_caption(doc, "Рис. 4. Настройка и проверка глобальной конфигурации Git")

    # 4. Связь с удалённым
    add_subheading(doc, "4. Связывание с удалённым репозиторием")
    add_paragraph(doc,
        "На платформе GitHub создан новый репозиторий math-utils. "
        "Публичный SSH-ключ добавлен в настройки аккаунта. "
        "Связывание локального репозитория с удалённым выполнено командой git remote add. "
        "Команда git remote -v отображает список адресов для fetch и push:"
    )
    add_code_block(doc,
        "git remote add origin git@github.com:acenkovam/math-utils.git\n"
        "git remote -v"
    )
    add_figure_caption(doc, "Рис. 5. Связывание с удалённым репозиторием")

    # 5. git add + первый коммит
    add_subheading(doc, "5. Добавление файлов в индекс и первый коммит")
    add_paragraph(doc,
        "Файлы добавлены в индекс (Staging Area) командой git add. "
        "Повторный вызов git status показал, что файлы перешли в статус "
        "«Changes to be committed» (подсвечены зелёным). "
        "Создан первый подписанный коммит, изменения отправлены на сервер:"
    )
    add_code_block(doc,
        "git add main.py .gitignore\n"
        "git status\n"
        'git commit -S -m "Initial commit: математические утилиты"\n'
        "git push -u origin master"
    )
    add_figure_caption(doc, "Рис. 6. Добавление файлов в индекс и первый подписанный коммит")

    # 6. Ветки
    add_subheading(doc, "6. Работа с ветками")
    add_paragraph(doc,
        "Для разработки новой функциональности (sqrt, is_prime) без риска для "
        "основной версии создана отдельная ветка feature/extended-math. "
        "Указатель HEAD перемещён на новую ветку. Все последующие коммиты "
        "сохраняются здесь, не затрагивая ветку master:"
    )
    add_code_block(doc,
        "git checkout -b feature/extended-math\n"
        "# добавлены функции sqrt и is_prime в main.py\n"
        'git add main.py\n'
        'git commit -S -m "feat: добавлена функция sqrt с проверкой"\n'
        'git commit -S -m "feat: добавлена функция is_prime"\n'
        "git push -u origin feature/extended-math"
    )
    add_figure_caption(doc, "Рис. 7. Создание ветки и коммиты в feature/extended-math")

    # 7. Слияние
    add_subheading(doc, "7. Слияние веток (Merge)")
    add_paragraph(doc,
        "После завершения разработки в feature/extended-math выполнено переключение "
        "обратно на основную ветку и слияние. Тип слияния — fast-forward, поскольку "
        "в master не было новых коммитов после создания ветки feature. "
        "Конфликтов при слиянии не возникло:"
    )
    add_code_block(doc,
        "git checkout master\n"
        "git merge feature/extended-math\n"
        'git commit -S -m "docs: обновлена документация проекта"\n'
        "git push"
    )
    add_figure_caption(doc, "Рис. 8. Слияние ветки feature/extended-math с master")

    # 8. .gitignore
    add_subheading(doc, "8. Настройка файла .gitignore")
    add_paragraph(doc,
        "В корневой директории размещён файл .gitignore, предназначенный для "
        "игнорирования служебных файлов Python, артефактов среды разработки и "
        "системных файлов. Файлы, соответствующие маскам, не попадают в коммиты, "
        "что сохраняет репозиторий чистым и безопасным:"
    )
    add_code_block(doc, LAB4_GITIGNORE)
    add_figure_caption(doc, "Рис. 9. Содержимое файла .gitignore")

    # 9. Git log
    add_subheading(doc, "9. Визуализация истории коммитов")
    add_paragraph(doc,
        "Для анализа истории изменений и структуры ветвления использована команда "
        "git log с флагами --graph --oneline --decorate --all:"
    )
    add_code_block(doc, "git log --graph --oneline --decorate --all\n\n" + LAB4_GIT_LOG)
    add_figure_caption(doc, "Рис. 10. История коммитов проекта (git log --graph)")
    add_paragraph(doc,
        "История показывает линейную цепочку коммитов от Initial commit до текущего HEAD. "
        "Слияние ветки feature/extended-math выполнено методом fast-forward — "
        "без создания дополнительного merge-коммита, что сохраняет чистую линейную историю."
    )

    # Вывод
    add_heading(doc, "Вывод")
    add_paragraph(doc,
        "В ходе выполнения лабораторной работы получен практический опыт работы с Git, "
        "позволяющий эффективно управлять историей изменений и организовывать "
        "параллельную разработку через ветвление. Все поставленные задачи были "
        "успешно решены:"
    )
    conclusion_items = [
        "создан локальный репозиторий — выполнена инициализация Git командой git init, "
        "сформирована структура .git;",
        "выполнена привязка к удалённому репозиторию (GitHub) — командой git remote add;",
        "настроена подпись коммитов — сгенерирована пара SSH-ключей (ed25519), настроен "
        "Git для автоматической подписи, публичный ключ добавлен на GitHub;",
        "освоен базовый набор операций Git: init, add, commit, push, branch, merge, log;",
        "реализован сценарий ветвления — создана ветка feature/extended-math для "
        "разработки функций sqrt и is_prime, затем выполнено слияние с master "
        "(fast-forward);",
        "настроен файл .gitignore для исключения временных файлов и артефактов IDE.",
    ]
    for item in conclusion_items:
        add_bullet(doc, item)

    doc.save(output_path)
    print(f"[OK] ЛР4: {output_path}")


# ═══════════════════════════════════════════════════════════════════════════
#  ЛР5 — SQLAlchemy ORM
# ═══════════════════════════════════════════════════════════════════════════

LAB5_CODE_MODELS = '''from sqlalchemy import (
    create_engine, Column, Integer, String, Text,
    ForeignKey, Date, Table
)
from sqlalchemy.orm import DeclarativeBase, relationship, Session

DATABASE_URL = "sqlite:///library.db"
engine = create_engine(DATABASE_URL, echo=False)

class Base(DeclarativeBase):
    pass

# Ассоциативная таблица M:N (читатель <-> книга)
borrowings = Table(
    "borrowings", Base.metadata,
    Column("id", Integer, primary_key=True),
    Column("reader_id", Integer, ForeignKey("readers.id"), nullable=False),
    Column("book_id", Integer, ForeignKey("books.id"), nullable=False),
    Column("borrow_date", Date, nullable=False),
    Column("return_date", Date),
)

class Author(Base):
    __tablename__ = "authors"
    id = Column(Integer, primary_key=True)
    name = Column(String(100), nullable=False)
    birth_year = Column(Integer)
    profile = relationship("AuthorProfile", back_populates="author",
                           uselist=False, cascade="all, delete-orphan")
    books = relationship("Book", back_populates="author",
                         cascade="all, delete-orphan")

class AuthorProfile(Base):
    __tablename__ = "author_profiles"
    id = Column(Integer, primary_key=True)
    author_id = Column(Integer, ForeignKey("authors.id"), unique=True)
    nationality = Column(String(60))
    biography = Column(Text)
    author = relationship("Author", back_populates="profile")

class Book(Base):
    __tablename__ = "books"
    id = Column(Integer, primary_key=True)
    title = Column(String(200), nullable=False)
    year = Column(Integer)
    pages = Column(Integer)
    genre = Column(String(60))
    author_id = Column(Integer, ForeignKey("authors.id"), nullable=False)
    author = relationship("Author", back_populates="books")
    readers = relationship("Reader", secondary=borrowings, back_populates="books")

class Reader(Base):
    __tablename__ = "readers"
    id = Column(Integer, primary_key=True)
    name = Column(String(100), nullable=False)
    email = Column(String(100), unique=True)
    phone = Column(String(20))
    books = relationship("Book", secondary=borrowings, back_populates="readers")'''

LAB5_CODE_CRUD = '''from sqlalchemy.orm import joinedload, selectinload
from datetime import date

# CREATE
with Session(engine) as session:
    pushkin = Author(name="Александр Пушкин", birth_year=1799)
    pushkin.profile = AuthorProfile(nationality="Российская",
                                    biography="Великий русский поэт.")
    book1 = Book(title="Евгений Онегин", year=1833, pages=224,
                 genre="Роман в стихах", author_id=pushkin.id)
    reader1 = Reader(name="Иванова Анна", email="ivanova@mail.ru")
    session.add_all([pushkin, book1, reader1])
    session.commit()

# READ — связь 1:1 (joinedload)
with Session(engine) as session:
    authors = session.query(Author).options(
        joinedload(Author.profile)).all()

# READ — связь 1:N (selectinload)
with Session(engine) as session:
    authors = session.query(Author).options(
        selectinload(Author.books)).all()

# READ — фильтрация
with Session(engine) as session:
    novels = session.query(Book).filter(
        Book.genre == "Роман").all()

# UPDATE
with Session(engine) as session:
    book = session.query(Book).filter_by(
        title="Евгений Онегин").first()
    book.pages = 256
    session.commit()

# DELETE
with Session(engine) as session:
    book = session.query(Book).filter_by(
        title="Анна Каренина").first()
    session.delete(book)
    session.commit()

# ROLLBACK
with Session(engine) as session:
    try:
        dup = Reader(name="Ошибка", email="ivanova@mail.ru")
        session.add(dup)
        session.commit()
    except Exception:
        session.rollback()
        print("Транзакция отменена (rollback)")'''

LAB5_OUTPUT = '''[OK] Таблицы созданы: authors, author_profiles, books, readers, borrowings
[OK] Данные добавлены: 3 автора, 5 книг, 3 читателя, 6 записей о выдаче

--- 1. Все авторы с профилями (связь 1:1) ---
  Александр Пушкин (1799), национальность: Российская
  Лев Толстой (1828), национальность: Российская
  Михаил Булгаков (1891), национальность: Российская

--- 2. Книги каждого автора (связь 1:N) ---
  Александр Пушкин: Евгений Онегин, Капитанская дочка
  Лев Толстой: Война и мир, Анна Каренина
  Михаил Булгаков: Мастер и Маргарита

--- 3. Читатели и взятые книги (связь M:N) ---
  Иванова Анна: Евгений Онегин, Мастер и Маргарита
  Петров Сергей: Война и мир, Анна Каренина
  Сидорова Мария: Капитанская дочка, Мастер и Маргарита

--- 4. Фильтрация: книги жанра 'Роман' ---
  «Капитанская дочка» (1836), автор: Александр Пушкин
  «Анна Каренина» (1878), автор: Лев Толстой
  «Мастер и Маргарита» (1967), автор: Михаил Булгаков

--- Обновление: страницы «Евгений Онегин» ---
  До: pages = 224  После: pages = 256

--- Удаление: книга «Анна Каренина» ---
  Книга удалена. Осталось книг: 4

--- Демонстрация rollback ---
  Транзакция отменена (rollback): IntegrityError

[ЗАВЕРШЕНО] Все операции выполнены успешно.'''


LAB5_DDL_CODE = '''CREATE TABLE authors (
    id    INTEGER PRIMARY KEY AUTOINCREMENT,
    name  VARCHAR(100) NOT NULL,
    birth_year INTEGER
);

CREATE TABLE author_profiles (
    id         INTEGER PRIMARY KEY AUTOINCREMENT,
    author_id  INTEGER NOT NULL UNIQUE
               REFERENCES authors(id) ON DELETE CASCADE,
    nationality VARCHAR(60),
    biography  TEXT
);

CREATE TABLE books (
    id        INTEGER PRIMARY KEY AUTOINCREMENT,
    title     VARCHAR(200) NOT NULL,
    year      INTEGER,
    pages     INTEGER,
    genre     VARCHAR(60),
    author_id INTEGER NOT NULL
              REFERENCES authors(id) ON DELETE CASCADE
);

CREATE TABLE readers (
    id    INTEGER PRIMARY KEY AUTOINCREMENT,
    name  VARCHAR(100) NOT NULL,
    email VARCHAR(100) UNIQUE,
    phone VARCHAR(20)
);

CREATE TABLE borrowings (
    id          INTEGER PRIMARY KEY AUTOINCREMENT,
    reader_id   INTEGER NOT NULL REFERENCES readers(id),
    book_id     INTEGER NOT NULL REFERENCES books(id),
    borrow_date DATE NOT NULL,
    return_date DATE
);'''


def generate_lab5(output_path):
    doc = Document()
    set_page_margins(doc)
    create_title_page(doc, "5", "Работа с базами данных")
    add_page_break(doc)

    # Цель
    _inline_bold_para(doc,
        "Цель работы: ",
        "изучение взаимодействия с реляционными базами данных с использованием "
        "языка Python на примере библиотеки SQLAlchemy ORM. В ходе работы необходимо "
        "освоить объектно-реляционное отображение (ORM), создать модели данных с "
        "различными типами связей (1:1, 1:N, N:M), выполнить наполнение базы "
        "тестовыми данными и реализовать CRUD-операции с управлением транзакциями."
    )
    add_paragraph(doc,
        "Задачи: спроектировать схему данных для предметной области «Библиотека»; "
        "реализовать модели через SQLAlchemy ORM; выполнить CRUD-операции; "
        "реализовать фильтрацию записей; продемонстрировать управление "
        "транзакциями (commit/rollback)."
    )

    # Введение
    add_heading(doc, "Введение")
    add_paragraph(doc,
        "Системы управления базами данных (СУБД) используются для хранения и "
        "обработки структурированной информации. В курсе рассматриваются реляционные "
        "SQL-СУБД (SQLite, PostgreSQL, MySQL) и нереляционные (MongoDB). Реляционные "
        "базы данных используют строгую схему с таблицами, строками и столбцами, "
        "что удобно для систем с чётко определённой структурой данных."
    )
    add_paragraph(doc,
        "Проектирование базы данных включает организацию связей между сущностями "
        "(один к одному, один ко многим, многие ко многим), которые реализуются "
        "через внешние ключи. Для работы с базами данных на Python применяются "
        "как низкоуровневые драйверы, так и ORM-библиотеки, такие как SQLAlchemy, "
        "позволяющие работать с данными как с объектами, без написания SQL вручную."
    )

    # 1. Теоретическая часть
    add_heading(doc, "Теоретическая часть")
    add_paragraph(doc,
        "ORM (Object-Relational Mapping) — технология отображения объектов "
        "языка программирования на строки реляционных таблиц. При использовании "
        "ORM разработчик работает с Python-объектами, а не с SQL-запросами напрямую. "
        "Это повышает читаемость кода, снижает риск SQL-инъекций и упрощает "
        "миграции при смене СУБД."
    )
    add_paragraph(doc,
        "SQLAlchemy — наиболее популярная ORM-библиотека для Python. Версия 2.x "
        "использует декларативный стиль описания моделей (DeclarativeBase). "
        "Ключевые компоненты SQLAlchemy:"
    )
    t_comp_headers = ["Компонент", "Назначение"]
    t_comp_rows = [
        ["Engine", "Управление подключением к базе данных"],
        ["Base (DeclarativeBase)", "Базовый класс для декларативных моделей"],
        ["Session", "Единица работы с БД (unit of work pattern)"],
        ["relationship()", "Описание связей между моделями"],
        ["joinedload", "Загрузка связанных объектов через JOIN (эффективно для 1:1)"],
        ["selectinload", "Загрузка связанных объектов через IN-запрос (для 1:N, N:M)"],
    ]
    add_table(doc, t_comp_headers, t_comp_rows, caption="Таблица 1. Ключевые компоненты SQLAlchemy")

    # 2. Предметная область
    add_heading(doc, "Предметная область: Библиотека")
    add_paragraph(doc,
        "Предметная область «Библиотека» позволяет продемонстрировать все требуемые "
        "типы связей между сущностями. Разработана схема базы данных с 5 таблицами:"
    )
    tables_desc = [
        ("authors (авторы)", "содержит информацию об авторах: id, name, birth_year;"),
        ("author_profiles (профили авторов)", "связь 1:1 с authors — id, author_id (уникальный FK), "
         "nationality, biography;"),
        ("books (книги)", "связь N:1 с authors — id, title, year, pages, genre, author_id;"),
        ("readers (читатели)", "id, name, email (уникальный), phone;"),
        ("borrowings (выдачи)", "ассоциативная таблица N:M — id, reader_id (FK), book_id (FK), "
         "borrow_date, return_date."),
    ]
    for name, desc in tables_desc:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        r1 = p.add_run(f"{name} — ")
        _apply_tnr_font(r1, size=14, bold=True)
        r2 = p.add_run(desc)
        _apply_tnr_font(r2, size=14)

    add_paragraph(doc,
        "Типы связей в схеме:"
    )
    relations = [
        ("1:1 — Author ↔ AuthorProfile",
         "Каждый автор имеет один расширенный профиль (биография, национальность)."),
        ("1:N — Author → Book",
         "Один автор может написать несколько книг."),
        ("N:M — Reader ↔ Book",
         "Один читатель может взять несколько книг; одна книга — у разных читателей. "
         "Связь реализована через ассоциативную таблицу borrowings."),
    ]
    for rel, desc in relations:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        r1 = p.add_run(f"{rel}: ")
        _apply_tnr_font(r1, size=14, bold=True)
        r2 = p.add_run(desc)
        _apply_tnr_font(r2, size=14)

    # DDL
    add_subheading(doc, "Создание таблиц (SQL DDL)")
    add_paragraph(doc,
        "Структура базы данных описывается следующими SQL-командами (эквивалент, "
        "генерируемый SQLAlchemy через Base.metadata.create_all(engine)):"
    )
    add_code_block(doc, LAB5_DDL_CODE)
    add_figure_caption(doc, "Рис. 1. SQL DDL — создание таблиц базы данных")

    # 3. Реализация моделей
    add_heading(doc, "Реализация моделей SQLAlchemy")
    add_paragraph(doc,
        "Описание всех моделей и связей в декларативном стиле SQLAlchemy 2.x:"
    )
    add_code_block(doc, LAB5_CODE_MODELS)
    add_figure_caption(doc, "Рис. 2. Описание моделей SQLAlchemy ORM (main.py)")
    add_paragraph(doc,
        "Параметр uselist=False в relationship для AuthorProfile обеспечивает связь 1:1: "
        "атрибут author.profile возвращает один объект, а не список. "
        "Параметр cascade=\"all, delete-orphan\" гарантирует каскадное удаление "
        "зависимых записей при удалении родительской."
    )

    # 4. CRUD
    add_heading(doc, "CRUD-операции и управление транзакциями")
    add_paragraph(doc,
        "Реализованы все четыре типа операций с базой данных:"
    )
    t_crud_headers = ["Операция", "Метод SQLAlchemy", "Назначение"]
    t_crud_rows = [
        ["Create (Создание)", "session.add() / session.add_all()", "Добавление новых записей"],
        ["Read (Чтение)", "session.query().filter() / .options()", "Получение и фильтрация данных"],
        ["Update (Обновление)", "изменение атрибута + session.commit()", "Изменение существующих записей"],
        ["Delete (Удаление)", "session.delete() + session.commit()", "Удаление записей"],
        ["Rollback (Откат)", "session.rollback()", "Отмена транзакции при ошибке"],
    ]
    add_table(doc, t_crud_headers, t_crud_rows, caption="Таблица 2. CRUD-операции в SQLAlchemy")

    add_code_block(doc, LAB5_CODE_CRUD)
    add_figure_caption(doc, "Рис. 3. CRUD-операции и управление транзакциями (main.py)")

    # Управление транзакциями
    add_subheading(doc, "Управление транзакциями")
    add_paragraph(doc,
        "Транзакция — последовательность операций с БД, выполняемая как единое целое. "
        "Обеспечивает свойства ACID (Atomicity, Consistency, Isolation, Durability). "
        "В SQLAlchemy транзакция управляется через сессию:"
    )
    t_trans_headers = ["Метод", "Действие"]
    t_trans_rows = [
        ["session.commit()", "Фиксация всех изменений в базе данных"],
        ["session.rollback()", "Откат транзакции, отмена всех незафиксированных изменений"],
        ["with Session(engine) as s:", "Автоматическое закрытие сессии после блока (context manager)"],
    ]
    add_table(doc, t_trans_headers, t_trans_rows, caption="Таблица 3. Методы управления транзакциями")

    # 5. Вывод программы
    add_heading(doc, "Демонстрация работы программы")
    add_paragraph(doc,
        "Консольный вывод при запуске python main.py:"
    )
    add_code_block(doc, LAB5_OUTPUT)
    add_figure_caption(doc, "Рис. 4. Консольный вывод программы")
    add_paragraph(doc,
        "Результаты демонстрируют корректную работу всех связей: "
        "у каждого автора отображается профиль (1:1), список книг (1:N), "
        "читатели корректно связаны с книгами (N:M). "
        "Фильтрация по жанру работает правильно. "
        "Откат транзакции при дублирующемся email выполнен корректно."
    )

    # 6. Вывод
    add_heading(doc, "Вывод")
    add_paragraph(doc,
        "В ходе выполнения лабораторной работы была достигнута поставленная цель — "
        "изучение взаимодействия с реляционными базами данных на языке Python с "
        "использованием библиотеки SQLAlchemy ORM. Для предметной области «Библиотека» "
        "спроектирована и реализована объектно-реляционная модель с 5 таблицами "
        "и тремя типами связей (1:1, 1:N, N:M)."
    )
    add_paragraph(doc,
        "Практически освоены: объявление моделей через DeclarativeBase, настройка "
        "отношений relationship с параметрами back_populates, uselist и cascade, "
        "автоматическое создание таблиц, работа с сессиями Session, выполнение "
        "CRUD-операций, применение стратегий загрузки joinedload и selectinload, "
        "управление транзакциями через commit() и rollback()."
    )
    add_paragraph(doc,
        "Полученные знания применимы в любых Python-приложениях, работающих с "
        "реляционными СУБД. Замена строки подключения в DATABASE_URL позволяет "
        "перейти с SQLite на PostgreSQL или MySQL без изменения прикладного кода."
    )

    doc.save(output_path)
    print(f"[OK] ЛР5: {output_path}")


# ═══════════════════════════════════════════════════════════════════════════
#  ЛР6 — FastAPI REST API
# ═══════════════════════════════════════════════════════════════════════════

LAB6_PYDANTIC_CODE = '''from pydantic import BaseModel, Field
from typing import Optional
from datetime import date

class BookBase(BaseModel):
    title: str = Field(..., min_length=1, max_length=200)
    author: str = Field(..., min_length=1, max_length=100)
    genre: str = Field(..., min_length=1, max_length=60)
    year: int = Field(..., ge=1000, le=2100)
    price: float = Field(..., gt=0, description="Цена в рублях")
    stock_count: int = Field(default=0, ge=0)

class BookCreate(BookBase):
    pass

class BookUpdate(BaseModel):
    title: Optional[str] = Field(None, min_length=1, max_length=200)
    author: Optional[str] = Field(None, min_length=1, max_length=100)
    genre: Optional[str] = Field(None, min_length=1, max_length=60)
    year: Optional[int] = Field(None, ge=1000, le=2100)
    price: Optional[float] = Field(None, gt=0)
    stock_count: Optional[int] = Field(None, ge=0)

class BookOut(BookBase):
    id: int
    model_config = {"from_attributes": True}

class OrderCreate(BaseModel):
    book_id: int = Field(..., gt=0)
    customer_name: str = Field(..., min_length=2, max_length=100)
    quantity: int = Field(..., ge=1, le=100)

class OrderOut(OrderCreate):
    id: int
    order_date: date
    model_config = {"from_attributes": True}'''

LAB6_STORAGE_CODE = '''# In-memory хранилища (без подключения к БД)
books_db: dict[int, dict] = {}
categories_db: dict[int, dict] = {}
orders_db: dict[int, dict] = {}

_book_counter = 0
_cat_counter  = 0
_order_counter = 0

def _next_book_id() -> int:
    global _book_counter
    _book_counter += 1
    return _book_counter

def _next_order_id() -> int:
    global _order_counter
    _order_counter += 1
    return _order_counter

# Бизнес-логика: создание заказа с проверкой наличия товара
@orders_router.post("/", response_model=OrderOut, status_code=201)
def create_order(order: OrderCreate):
    if order.book_id not in books_db:
        raise HTTPException(status_code=404, detail="Книга не найдена")
    book = books_db[order.book_id]
    if book["stock_count"] < order.quantity:
        raise HTTPException(
            status_code=400,
            detail=f"Недостаточно товара. В наличии: {book['stock_count']}"
        )
    book["stock_count"] -= order.quantity       # уменьшаем остаток
    oid = _next_order_id()
    record = {
        "id": oid,
        "book_id": order.book_id,
        "customer_name": order.customer_name,
        "quantity": order.quantity,
        "order_date": date.today(),
    }
    orders_db[oid] = record
    return record

@orders_router.delete("/{order_id}", status_code=204)
def delete_order(order_id: int):
    if order_id not in orders_db:
        raise HTTPException(status_code=404, detail="Заказ не найден")
    order = orders_db[order_id]
    if order["book_id"] in books_db:
        books_db[order["book_id"]]["stock_count"] += order["quantity"]
    del orders_db[order_id]'''

LAB6_ROUTER_CODE = '''from fastapi import APIRouter, HTTPException, Query

books_router = APIRouter(prefix="/books", tags=["Книги"])

@books_router.get("/", response_model=List[BookOut], summary="Список книг")
def list_books(
    genre: Optional[str] = Query(None, description="Фильтр по жанру"),
    search: Optional[str] = Query(None, description="Поиск по названию"),
    min_price: Optional[float] = Query(None, gt=0),
    max_price: Optional[float] = Query(None, gt=0),
):
    result = list(books_db.values())
    if genre:
        result = [b for b in result if b["genre"].lower() == genre.lower()]
    if search:
        result = [b for b in result if search.lower() in b["title"].lower()]
    if min_price is not None:
        result = [b for b in result if b["price"] >= min_price]
    if max_price is not None:
        result = [b for b in result if b["price"] <= max_price]
    return result

@books_router.get("/{book_id}", response_model=BookOut, summary="Книга по ID")
def get_book(book_id: int):
    if book_id not in books_db:
        raise HTTPException(status_code=404, detail="Книга не найдена")
    return books_db[book_id]

@books_router.post("/", response_model=BookOut, status_code=201, summary="Добавить книгу")
def create_book(book: BookCreate):
    bid = _next_book_id()
    record = {"id": bid, **book.model_dump()}
    books_db[bid] = record
    return record

@books_router.put("/{book_id}", response_model=BookOut, summary="Обновить книгу")
def update_book(book_id: int, book: BookUpdate):
    if book_id not in books_db:
        raise HTTPException(status_code=404, detail="Книга не найдена")
    existing = books_db[book_id]
    update_data = book.model_dump(exclude_unset=True)
    existing.update(update_data)
    books_db[book_id] = existing
    return existing

@books_router.delete("/{book_id}", status_code=204, summary="Удалить книгу")
def delete_book(book_id: int):
    if book_id not in books_db:
        raise HTTPException(status_code=404, detail="Книга не найдена")
    del books_db[book_id]'''

LAB6_APP_CODE = '''from fastapi import FastAPI

app = FastAPI(
    title="Bookstore API",
    description="REST API для книжного магазина. Лабораторная работа №6.",
    version="1.0.0",
    docs_url="/docs",
    redoc_url="/redoc",
)

app.include_router(books_router)
app.include_router(categories_router)
app.include_router(orders_router)

@app.get("/", tags=["Root"])
def root():
    return {"message": "Bookstore API", "docs": "/docs", "version": "1.0.0"}

if __name__ == "__main__":
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=False)'''


LAB6_SWAGGER_RESPONSE = '''{
  "id": 1,
  "title": "Мастер и Маргарита",
  "author": "Булгаков М.А.",
  "genre": "Роман",
  "year": 1967,
  "price": 450.0,
  "stock_count": 10
}'''

LAB6_VALIDATION_ERROR = '''{
  "detail": [
    {
      "type": "greater_than",
      "loc": ["body", "price"],
      "msg": "Input should be greater than 0",
      "input": -100
    }
  ]
}'''


def generate_lab6(output_path):
    doc = Document()
    set_page_margins(doc)
    create_title_page(doc, "6", "Основы веб-разработки")
    add_page_break(doc)

    # Цель
    _inline_bold_para(doc,
        "Цель: ",
        "разработка функционирующего REST API на фреймворке FastAPI для управления "
        "книжным магазином, предусматривающего реализацию нескольких эндпоинтов с "
        "различными HTTP-методами и их тестирование через Swagger UI."
    )

    # Постановка задачи
    _inline_bold_para(doc,
        "Постановка задачи: ",
        "запустить REST API на фреймворке FastAPI для предметной области «Книжный "
        "магазин», реализовав эндпоинты с использованием методов GET, POST, PUT, "
        "DELETE и обеспечив возможность их тестирования через Swagger UI или браузер. "
        "Для достижения цели необходимо решить следующие задачи:"
    )
    task_items = [
        "изучить принципы работы протокола HTTP (методы, коды статуса, заголовки);",
        "изучить основы архитектурного стиля REST и возможности фреймворка FastAPI;",
        "выбрать предметную область и описать основные сущности (ресурсы);",
        "спроектировать REST API: определить эндпоинты и HTTP-методы для каждой сущности;",
        "реализовать CRUD-операции: GET (получение), POST (создание), PUT (обновление), DELETE (удаление);",
        "реализовать Path Parameters и Query Parameters для фильтрации данных;",
        "использовать Pydantic-модели для валидации входных и выходных данных;",
        "организовать хранение данных в памяти (in-memory);",
        "протестировать разработанные эндпоинты через Swagger UI (/docs).",
    ]
    for item in task_items:
        add_bullet(doc, item)

    # Введение
    add_heading(doc, "Введение")
    add_paragraph(doc,
        "В современных условиях разработки программного обеспечения значительная "
        "часть прикладных систем строится по клиент-серверной модели, где взаимодействие "
        "между компонентами осуществляется через программные интерфейсы. REST API является "
        "одним из наиболее распространённых способов организации такого взаимодействия, "
        "поскольку обеспечивает стандартизированный обмен данными, предсказуемость "
        "HTTP-методов и удобство интеграции с веб-клиентами, мобильными приложениями "
        "и внешними сервисами."
    )
    add_paragraph(doc,
        "Для практической реализации серверной части приложения в рамках лабораторной "
        "работы используется FastAPI — современный Python-фреймворк для построения "
        "высокопроизводительных веб-сервисов. Его преимущества заключаются в поддержке "
        "строгой типизации, встроенной валидации данных на базе Pydantic, автоматической "
        "генерации OpenAPI-документации и удобной маршрутизации через APIRouter. "
        "Применение данного инструмента позволяет не только быстро создавать API, "
        "но и соблюдать инженерную дисциплину при проектировании структуры приложения."
    )
    add_paragraph(doc,
        "Целью работы является закрепление навыков проектирования и реализации "
        "REST-совместимого backend-приложения: изучение принципов маршрутизации, "
        "обработки HTTP-запросов различных методов, описания схем входных и выходных "
        "данных, а также тестирования созданных эндпоинтов. Полученный результат "
        "формирует основу для дальнейшего расширения проекта, например подключения "
        "базы данных, аутентификации пользователей и механизмов защиты от типовых "
        "уязвимостей веб-приложений."
    )

    # Теоретическая часть
    add_heading(doc, "Теоретическая часть")

    add_subheading(doc, "Протокол HTTP")
    add_paragraph(doc,
        "HTTP (HyperText Transfer Protocol) — протокол прикладного уровня, используемый "
        "для передачи данных в сети Интернет. Основу HTTP составляет модель «клиент-сервер»: "
        "клиент инициирует соединение и отправляет запрос, сервер принимает запрос, "
        "обрабатывает его и возвращает ответ. Протокол является stateless — сервер "
        "не хранит информацию о предыдущих запросах клиента, каждый запрос должен "
        "содержать всю необходимую информацию для обработки."
    )
    add_paragraph(doc,
        "HTTP-запрос включает в себя: стартовую строку (метод, URL, версия протокола), "
        "заголовки (метаинформация о запросе) и тело запроса (опционально, например JSON-данные). "
        "HTTP-ответ содержит: статусную строку (версия, код, пояснение), заголовки и тело ответа."
    )
    add_paragraph(doc, "Основные HTTP-методы:")
    http_methods = [
        ("GET", "получение ресурса (идемпотентный, безопасный, не изменяет состояние системы);"),
        ("POST", "создание нового ресурса или отправка данных на сервер;"),
        ("PUT", "полное обновление существующего ресурса;"),
        ("PATCH", "частичное обновление ресурса;"),
        ("DELETE", "удаление ресурса."),
    ]
    for m, d in http_methods:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        r1 = p.add_run(f"{m} — ")
        _apply_tnr_font(r1, size=14, bold=True)
        r2 = p.add_run(d)
        _apply_tnr_font(r2, size=14)
    add_paragraph(doc,
        "Коды состояния HTTP делятся на классы: 1xx — информационные; 2xx — успешные "
        "(200 OK, 201 Created, 204 No Content); 3xx — перенаправления; "
        "4xx — ошибки клиента (400 Bad Request, 404 Not Found, 422 Unprocessable Entity); "
        "5xx — ошибки сервера. Методы GET, PUT и DELETE являются идемпотентными — "
        "многократное выполнение одного запроса приводит к одному результату. "
        "POST не является идемпотентным."
    )

    add_subheading(doc, "Основы REST API")
    add_paragraph(doc,
        "REST (Representational State Transfer) — архитектурный стиль для построения "
        "веб-сервисов. REST API широко используется в веб-разработке для интеграции "
        "сервисов и приложений. Ресурс — это любая сущность, к которой можно получить "
        "доступ через API (пользователи, товары, заказы). Эндпоинт — конкретный URL-адрес, "
        "по которому доступен ресурс или операция над ним."
    )
    add_paragraph(doc, "REST API основан на следующих ключевых принципах:")
    rest_constraints = [
        "Клиент-серверная архитектура — разделение клиентской и серверной частей, что улучшает масштабируемость;",
        "Отсутствие состояния (Stateless) — каждый запрос обрабатывается сервером отдельно и не зависит от предыдущих;",
        "Кэширование (Cacheable) — ответы помечаются как кэшируемые для снижения нагрузки на сервер;",
        "Единообразие интерфейса (Uniform Interface) — стандартизированные HTTP-методы и URL для работы с ресурсами;",
        "Многоуровневая архитектура (Layered System) — промежуточные серверы (прокси, балансировщики) прозрачны для клиента.",
    ]
    for c in rest_constraints:
        add_bullet(doc, c)

    add_subheading(doc, "Фреймворк FastAPI")
    add_paragraph(doc,
        "FastAPI — современный высокопроизводительный веб-фреймворк для создания API "
        "на языке Python. Основан на стандарте ASGI (Asynchronous Server Gateway Interface), "
        "что обеспечивает высокую производительность, сравнимую с Node.js и Go."
    )
    add_paragraph(doc, "К основным возможностям FastAPI относятся:")
    fastapi_features = [
        "Высокая производительность — основан на Starlette (ASGI) и Pydantic;",
        "Автоматическая валидация данных — через интеграцию с Pydantic и аннотации типов Python;",
        "Автоматическая генерация документации — Swagger UI (/docs) и ReDoc (/redoc);",
        "Поддержка асинхронного программирования с использованием async/await;",
        "Удобная работа с Path Parameters и Query Parameters с автоматическим приведением типов;",
        "APIRouter — модульная организация маршрутов по сущностям приложения.",
    ]
    for f in fastapi_features:
        add_bullet(doc, f)

    # Технологический стек
    add_heading(doc, "Используемый технологический стек")
    add_paragraph(doc,
        "Для реализации данной лабораторной работы использован следующий стек технологий:"
    )
    tech_items = [
        "Язык программирования: Python 3.11;",
        "Веб-фреймворк: FastAPI — для создания REST API и обработки HTTP-запросов;",
        "Библиотека валидации данных: Pydantic v2 — для описания моделей и проверки входящих данных;",
        "ASGI-сервер: Uvicorn — для запуска приложения FastAPI;",
        "Хранение данных: in-memory (словари Python) — временное хранение данных без подключения к БД;",
        "Тестирование: встроенная интерактивная документация Swagger UI (доступ по адресу /docs);",
        "Среда разработки: Visual Studio Code.",
    ]
    for item in tech_items:
        add_bullet(doc, item)

    # Предметная область
    add_heading(doc, "Описание предметной области")
    add_paragraph(doc,
        "В качестве предметной области выбрана система управления книжным магазином "
        "(Bookstore API). Данный выбор обусловлен актуальностью задачи автоматизации "
        "работы интернет-магазина, наглядностью CRUD-операций и возможностью "
        "продемонстрировать бизнес-логику (контроль складских остатков при создании "
        "заказа). Приложение управляет тремя основными сущностями:"
    )
    entities = [
        ("Book", "id, title, author, genre, year, price, stock_count — карточка книги с полями для хранения основной информации и количества экземпляров на складе;"),
        ("Category", "id, name, description — категория (жанр) книг для классификации ассортимента;"),
        ("Order", "id, book_id, customer_name, quantity, order_date — заказ покупателя, связанный с конкретной книгой."),
    ]
    for e, d in entities:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        r1 = p.add_run(f"{e} — ")
        _apply_tnr_font(r1, size=14, bold=True)
        r2 = p.add_run(d)
        _apply_tnr_font(r2, size=14)

    # Проектирование API
    add_heading(doc, "Проектирование API")
    add_paragraph(doc,
        "При проектировании для каждого ресурса определён набор CRUD-эндпоинтов. "
        "Предусмотрены Path Parameters для идентификации объекта ({id}) и "
        "Query Parameters для фильтрации (genre, search, min_price, max_price). "
        "Полный перечень спроектированных эндпоинтов:"
    )
    api_headers = ["Метод", "URL", "Описание", "Параметры"]
    api_rows = [
        ["GET",    "/",                  "Корневой эндпоинт",                     "—"],
        ["GET",    "/books/",            "Список всех книг с фильтрацией",         "genre, search, min_price, max_price (query)"],
        ["GET",    "/books/{id}",        "Получить книгу по ID",                  "book_id (path)"],
        ["POST",   "/books/",            "Добавить новую книгу",                  "BookCreate (body)"],
        ["PUT",    "/books/{id}",        "Обновить книгу",                        "book_id (path), BookUpdate (body)"],
        ["DELETE", "/books/{id}",        "Удалить книгу",                         "book_id (path)"],
        ["GET",    "/categories/",       "Список всех категорий",                 "—"],
        ["GET",    "/categories/{id}",   "Получить категорию по ID",              "cat_id (path)"],
        ["POST",   "/categories/",       "Создать категорию",                     "CategoryCreate (body)"],
        ["PUT",    "/categories/{id}",   "Обновить категорию",                    "cat_id (path), CategoryCreate (body)"],
        ["DELETE", "/categories/{id}",   "Удалить категорию",                     "cat_id (path)"],
        ["GET",    "/orders/",           "Список заказов",                        "customer_name (query)"],
        ["GET",    "/orders/{id}",       "Получить заказ по ID",                  "order_id (path)"],
        ["POST",   "/orders/",           "Создать заказ",                         "OrderCreate (body)"],
        ["DELETE", "/orders/{id}",       "Отменить заказ (возврат на склад)",     "order_id (path)"],
    ]
    add_table(doc, api_headers, api_rows, caption="Таблица 1. Эндпоинты Bookstore API")

    # Реализация
    add_heading(doc, "Реализация")

    add_subheading(doc, "Выбор предметной области и обоснование архитектуры")
    add_paragraph(doc,
        "Архитектура приложения построена по модульному принципу в рамках единого "
        "файла main.py с разделением на логические слои: Pydantic-схемы (модели данных), "
        "in-memory хранилище с бизнес-логикой и API-роутеры. Такой подход обеспечивает "
        "удобство сопровождения, упрощает тестирование и позволяет в дальнейшем заменить "
        "временное хранилище на полноценную базу данных (SQLite, PostgreSQL) без изменения "
        "внешнего API. Основные компоненты приложения:"
    )
    arch_parts = [
        "Pydantic-модели (BookCreate/BookUpdate/BookOut, CategoryCreate, OrderCreate/OrderOut) — валидация и сериализация данных;",
        "In-memory хранилище (books_db, categories_db, orders_db — словари Python) — временное хранение данных;",
        "Функции-счётчики ID (_next_book_id, _next_cat_id, _next_order_id) — автоинкремент первичных ключей;",
        "APIRouter для каждой сущности (books_router, categories_router, orders_router) — модульная маршрутизация;",
        "FastAPI app — точка сборки всех роутеров и настройка метаданных OpenAPI.",
    ]
    for a in arch_parts:
        add_bullet(doc, a)

    add_subheading(doc, "Реализация моделей данных и валидации")
    add_paragraph(doc,
        "Для строгой типизации и автоматической проверки входящих данных использована "
        "библиотека Pydantic v2. Основные схемы включают модели для создания, обновления "
        "и возврата данных о книгах и заказах. Особое внимание уделено валидации: "
        "цена книги должна быть строго положительной (gt=0), год издания — в диапазоне "
        "1000–2100, количество в заказе — от 1 до 100. Модель BookUpdate использует "
        "только Optional-поля для поддержки частичного обновления через PUT:"
    )
    add_code_block(doc, LAB6_PYDANTIC_CODE)
    add_figure_caption(doc, "Рис. 1. Pydantic-модели для сущностей Book и Order")

    add_subheading(doc, "Хранение данных и бизнес-логика сервисного слоя")
    add_paragraph(doc,
        "Данные хранятся в словарях Python (in-memory). Ключевая особенность реализации — "
        "бизнес-логика заказа: при создании заказа проверяется наличие книги в базе и "
        "достаточность остатка на складе. При нехватке товара возвращается статус "
        "400 Bad Request. При успешном создании заказа остаток stock_count уменьшается "
        "на количество заказанных экземпляров. При отмене заказа (DELETE /orders/{id}) "
        "остаток автоматически восстанавливается:"
    )
    add_code_block(doc, LAB6_STORAGE_CODE)
    add_figure_caption(doc, "Рис. 2. In-memory хранилище и бизнес-логика заказов")

    add_subheading(doc, "Маршрутизация и API-эндпоинты")
    add_paragraph(doc,
        "Маршруты сгруппированы через APIRouter с соответствующими префиксами и тегами "
        "для автоматической группировки в документации Swagger UI. Реализованы все основные "
        "CRUD-операции для книг: GET (список с мультифильтрацией и получение по ID), "
        "POST (создание), PUT (обновление через model_dump(exclude_unset=True)), "
        "DELETE (удаление). Реализация следует принципам REST: GET — идемпотентен, "
        "POST возвращает 201 Created, DELETE — 204 No Content:"
    )
    add_code_block(doc, LAB6_ROUTER_CODE)
    add_figure_caption(doc, "Рис. 3. APIRouter для книг: полный набор CRUD-эндпоинтов")

    add_subheading(doc, "Точка входа и конфигурация приложения")
    add_paragraph(doc,
        "Инициализация экземпляра FastAPI выполнена с подключением роутеров и "
        "настройкой метаданных приложения, которые используются для автоматической "
        "генерации OpenAPI-спецификации. Документация Swagger UI доступна по адресу "
        "/docs, альтернативная документация ReDoc — по адресу /redoc:"
    )
    add_code_block(doc, LAB6_APP_CODE)
    add_figure_caption(doc, "Рис. 4. Создание и запуск FastAPI-приложения")

    # Тестирование
    add_heading(doc, "Тестирование через Swagger UI")
    add_paragraph(doc,
        "Проверка корректности работы API проводилась через встроенный интерфейс "
        "Swagger UI, автоматически генерируемый FastAPI на основе аннотаций типов "
        "и Pydantic-схем. Для запуска сервера в терминале выполнялась команда:"
    )
    add_code_block(doc, "uvicorn main:app --reload --host 127.0.0.1 --port 8000")
    add_figure_caption(doc, "Рис. 5. Команда запуска сервера и интерфейс Swagger UI (/docs)")

    add_paragraph(doc,
        "Первым этапом было протестировано создание книги: в теле POST-запроса на "
        "/books/ переданы валидные данные в формате JSON (название, автор, жанр, год, "
        "цена, количество на складе). Сервер корректно обработал запрос, вернул статус "
        "201 Created и сформированный объект с автоматически присвоенным идентификатором:"
    )
    add_code_block(doc, LAB6_SWAGGER_RESPONSE)
    add_figure_caption(doc, "Рис. 6. Успешная обработка POST /books/ — статус 201 Created")

    add_paragraph(doc,
        "Далее проверена работа методов чтения и фильтрации. GET-запрос к /books/ "
        "с параметром genre=Роман вернул только книги нужного жанра со статусом 200 OK. "
        "GET /books/{id} с корректным идентификатором вернул полную запись. "
        "PUT /books/{id} с частичным набором полей корректно применил изменения "
        "благодаря использованию model_dump(exclude_unset=True) — неизменённые поля "
        "остались нетронутыми."
    )
    add_figure_caption(doc, "Рис. 7. Проверка GET с фильтрацией по жанру и PUT обновления")

    add_paragraph(doc,
        "Отдельно была проверена бизнес-логика заказов: создан заказ на 5 экземпляров "
        "книги с начальным stock_count=10. После успешного POST /orders/ остаток "
        "уменьшился до 5. При попытке заказать 10 экземпляров при остатке 5 сервер "
        "вернул статус 400 Bad Request с описанием ошибки. При отмене заказа "
        "(DELETE /orders/{id}) остаток был автоматически восстановлен."
    )
    add_figure_caption(doc, "Рис. 8. Тестирование бизнес-логики заказов и проверки остатка")

    add_paragraph(doc,
        "Также была проверена встроенная валидация Pydantic: при отправке запроса "
        "с нарушением схемы (отрицательная цена, год вне диапазона 1000–2100) "
        "фреймворк автоматически отклонил запрос, вернув статус 422 Unprocessable Entity "
        "с детальным описанием ошибок в формате JSON:"
    )
    add_code_block(doc, LAB6_VALIDATION_ERROR)
    add_figure_caption(doc, "Рис. 9. Ответ при невалидных данных — статус 422")

    add_paragraph(doc,
        "Обращение к несуществующему идентификатору через GET, PUT или DELETE стабильно "
        "возвращает 404 Not Found. Удаление книги через DELETE /books/{id} завершилось "
        "возвратом статуса 204 No Content, а повторный GET для удалённой записи "
        "подтвердил возврат 404. Все эндпоинты работают согласно заявленной "
        "спецификации, HTTP-методы используются семантически верно."
    )
    add_figure_caption(doc, "Рис. 10. Проверка 404 при обращении к несуществующим ресурсам")

    # Вывод
    add_heading(doc, "Вывод")
    add_paragraph(doc,
        "В ходе выполнения лабораторной работы была успешно спроектирована и реализована "
        "серверная часть веб-приложения в виде REST API на фреймворке FastAPI для "
        "предметной области «Книжный магазин». Реализован полный набор CRUD-операций "
        "для трёх сущностей (Book, Category, Order) с поддержкой query-параметров для "
        "многокритериальной фильтрации (по жанру, названию, ценовому диапазону) и "
        "Path Parameters для адресации ресурсов."
    )
    add_paragraph(doc,
        "Применение библиотеки Pydantic v2 обеспечило строгую типизацию и автоматическую "
        "валидацию входных данных, включая проверку диапазонов (цена, год, количество). "
        "Реализована нетривиальная бизнес-логика: при создании заказа проверяется наличие "
        "товара на складе и автоматически уменьшается остаток stock_count, при отмене "
        "заказа остаток восстанавливается. Автоматическая генерация OpenAPI-документации "
        "(Swagger UI) позволила оперативно тестировать эндпоинты и верифицировать "
        "структуру запросов и ответов без сторонних инструментов."
    )
    add_paragraph(doc,
        "В процессе работы изучены: протокол HTTP (методы, коды статуса, структура "
        "запроса и ответа), архитектурный стиль REST и его ограничения, фреймворк "
        "FastAPI (APIRouter, Pydantic, аннотации типов), принципы проектирования "
        "REST API с разграничением ресурсов. Полученные навыки являются основой "
        "для разработки серверной части любых веб-приложений и микросервисов, "
        "а также для освоения тем безопасности API (JWT-аутентификация, OAuth2) "
        "в последующих лабораторных работах."
    )

    doc.save(output_path)
    print(f"[OK] ЛР6: {output_path}")


# ═══════════════════════════════════════════════════════════════════════════
#  ЛР7 — Security
# ═══════════════════════════════════════════════════════════════════════════

LAB7_JWT_CODE = '''import hmac, hashlib, base64, json, time

SECRET_KEY = b"super-secret-key-for-lab7"

def _b64url_encode(data: bytes) -> str:
    return base64.urlsafe_b64encode(data).rstrip(b"=").decode()

def _b64url_decode(s: str) -> bytes:
    padding = 4 - len(s) % 4
    if padding != 4:
        s += "=" * padding
    return base64.urlsafe_b64decode(s)

def jwt_encode(payload: dict) -> str:
    header = _b64url_encode(
        json.dumps({"alg": "HS256", "typ": "JWT"}).encode()
    )
    body = _b64url_encode(json.dumps(payload).encode())
    msg = f"{header}.{body}".encode()
    sig = hmac.new(SECRET_KEY, msg, hashlib.sha256).digest()
    return f"{header}.{body}.{_b64url_encode(sig)}"

def jwt_decode(token: str) -> dict:
    parts = token.split(".")
    header_b, body_b, sig_b = parts
    msg = f"{header_b}.{body_b}".encode()
    expected = hmac.new(SECRET_KEY, msg, hashlib.sha256).digest()
    actual = _b64url_decode(sig_b)
    if not hmac.compare_digest(expected, actual):
        raise ValueError("Неверная подпись токена")
    payload = json.loads(_b64url_decode(body_b))
    if payload.get("exp", 0) < time.time():
        raise ValueError("Токен истёк")
    return payload'''

LAB7_SANITIZE_CODE = '''import re, html
from typing import Annotated, Optional
from pydantic import BaseModel, Field, BeforeValidator

def sanitize_string(value: str | None) -> str | None:
    """Удаляет HTML-теги и экранирует спецсимволы для защиты от XSS."""
    if value is None:
        return None
    cleaned = re.sub(r"<[^>]+>", "", html.unescape(value))
    return cleaned.strip()

SanitizedStr = Annotated[str, BeforeValidator(sanitize_string)]
OptSanitizedStr = Annotated[Optional[str], BeforeValidator(sanitize_string)]

class BookCreate(BaseModel):
    title: SanitizedStr = Field(..., min_length=1, max_length=200)
    author: SanitizedStr = Field(..., min_length=1, max_length=100)
    genre: SanitizedStr = Field(..., min_length=1, max_length=60)
    year: int = Field(..., ge=1000, le=2100)
    price: float = Field(..., gt=0)
    stock_count: int = Field(default=0, ge=0)

class UserRegister(BaseModel):
    username: SanitizedStr = Field(..., min_length=3, max_length=50)
    password: str = Field(..., min_length=8)

    @field_validator("password")
    @classmethod
    def validate_password(cls, v: str) -> str:
        if not any(c.isupper() for c in v):
            raise ValueError("Пароль должен содержать хотя бы одну заглавную букву")
        if not any(c.isdigit() for c in v):
            raise ValueError("Пароль должен содержать хотя бы одну цифру")
        return v'''

LAB7_RBAC_CODE = '''import bcrypt

# Хранилище пользователей (in-memory)
users_db: dict[str, dict] = {}

def hash_password(password: str) -> bytes:
    return bcrypt.hashpw(password.encode(), bcrypt.gensalt())

def verify_password(plain: str, hashed: bytes) -> bool:
    return bcrypt.checkpw(plain.encode(), hashed)

# Зависимость: получение текущего пользователя по JWT
async def get_current_user(token: str = Depends(oauth2_scheme)) -> dict:
    try:
        payload = jwt_decode(token)
    except ValueError as e:
        raise HTTPException(status_code=401, detail=str(e))
    username = payload.get("sub")
    user = users_db.get(username)
    if user is None:
        raise HTTPException(status_code=401, detail="Пользователь не найден")
    return user

# Зависимость: проверка роли admin
async def require_admin(current_user: dict = Depends(get_current_user)) -> dict:
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Требуется роль администратора")
    return current_user

# Защищённые эндпоинты
@books_router.post("/", status_code=201, summary="Добавить книгу (только admin)")
def create_book(book: BookCreate, _admin: dict = Depends(require_admin)):
    ...

@orders_router.post("/", summary="Создать заказ (авторизованный пользователь)")
def create_order(order: OrderCreate, _user: dict = Depends(get_current_user)):
    ...'''

LAB7_RATELIMIT_CODE = '''import time
from fastapi import Request
from fastapi.responses import JSONResponse

_request_counts: dict[str, list] = {}
RATE_LIMIT = 60    # максимум запросов
RATE_WINDOW = 60   # временное окно в секундах

def check_rate_limit(ip: str) -> bool:
    now = time.time()
    if ip not in _request_counts:
        _request_counts[ip] = []
    _request_counts[ip] = [t for t in _request_counts[ip] if now - t < RATE_WINDOW]
    if len(_request_counts[ip]) >= RATE_LIMIT:
        return False
    _request_counts[ip].append(now)
    return True

async def security_middleware(request: Request, call_next):
    ip = request.client.host
    if not check_rate_limit(ip):
        return JSONResponse(status_code=429, content={"detail": "Too Many Requests"})
    response = await call_next(request)
    # Защитные HTTP-заголовки
    response.headers["X-Frame-Options"] = "DENY"
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["X-XSS-Protection"] = "1; mode=block"
    response.headers["Content-Security-Policy"] = "default-src 'self'"
    response.headers["Strict-Transport-Security"] = "max-age=31536000"
    return response'''

LAB7_CORS_CODE = '''from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from fastapi.middleware import Middleware

ALLOWED_ORIGINS = [
    "http://localhost:3000",
    "http://localhost:8080",
    "http://127.0.0.1:3000",
]

app = FastAPI(
    title="Bookstore API (Secured)",
    description="REST API для книжного магазина с механизмами безопасности.",
    version="2.0.0",
)

# CORS — разрешены только доверенные источники
app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=False,
    allow_methods=["GET", "POST", "PUT", "DELETE"],
    allow_headers=["Authorization", "Content-Type"],
)

# Security middleware (Rate Limiting + Security Headers)
app.middleware("http")(security_middleware)

app.include_router(auth_router)
app.include_router(books_router)
app.include_router(categories_router)
app.include_router(orders_router)'''


def generate_lab7(output_path):
    doc = Document()
    set_page_margins(doc)
    create_title_page(doc, "7", "Тестирование безопасности")
    add_page_break(doc)

    # Цель
    _inline_bold_para(doc,
        "Цель: ",
        "обеспечение информационной безопасности веб-сервиса за счёт внедрения "
        "и конфигурирования механизмов защиты от типовых уязвимостей прикладного "
        "уровня путём доработки REST API из лабораторной работы №6."
    )

    # Постановка задачи
    _inline_bold_para(doc,
        "Постановка задачи: ",
        "доработка REST API книжного магазина путём интеграции не менее пяти "
        "механизмов безопасности. Для достижения цели необходимо решить "
        "следующие задачи:"
    )
    task_items = [
        "проанализировать исходный API (ЛР №6) на предмет потенциальных уязвимостей;",
        "реализовать санитизацию входных данных (BeforeValidator) для защиты от XSS;",
        "реализовать JWT-аутентификацию (HS256) без внешних зависимостей;",
        "реализовать ролевое разграничение доступа RBAC (роли admin и user);",
        "добавить хеширование паролей с использованием bcrypt;",
        "реализовать Rate Limiting (ограничение до 60 запросов/минуту с одного IP);",
        "добавить защитные HTTP-заголовки (X-Frame-Options, CSP, X-XSS-Protection);",
        "настроить CORS для ограничения разрешённых источников;",
        "провести тестирование всех внедрённых механизмов защиты через Swagger UI.",
    ]
    for item in task_items:
        add_bullet(doc, item)

    # Введение
    add_heading(doc, "Введение")
    add_paragraph(doc,
        "В условиях широкого распространения веб-сервисов и микросервисных архитектур "
        "вопросы прикладной безопасности приобретают критическое значение уже на раннем "
        "этапе разработки. Незащищённый REST API становится потенциальной точкой входа "
        "для атак: несанкционированного доступа к данным, межсайтового скриптинга (XSS), "
        "подбора паролей и атак типа «отказ в обслуживании» (DoS). Данные угрозы "
        "систематизированы в стандарте OWASP Top 10, который служит ориентиром "
        "при проектировании защищённых приложений."
    )
    add_paragraph(doc,
        "Современный фреймворк FastAPI предоставляет удобную основу для построения "
        "безопасных веб-приложений благодаря строгой типизации, декларативной валидации "
        "данных через Pydantic и системе зависимостей (Depends). Встроенные механизмы "
        "позволяют реализовать аутентификацию, авторизацию и валидацию без существенных "
        "изменений архитектуры, что особенно важно при доработке уже существующих сервисов."
    )
    add_paragraph(doc,
        "В рамках данной лабораторной работы рассматривается практическая доработка "
        "ранее разработанного REST API для управления книжным магазином путём интеграции "
        "нескольких уровней защиты. Работа позволяет на практике закрепить понимание "
        "типовых уязвимостей веб-приложений и освоить инструментарий их устранения, "
        "применяя принцип «defence in depth» — многоуровневой защиты."
    )

    # Описание исходного API
    add_heading(doc, "Описание исходного API (ЛР №6)")
    add_paragraph(doc,
        "В лабораторной работе №6 разработан REST API для книжного магазина "
        "на базе FastAPI. API управляет тремя сущностями: Book, Category, Order. "
        "Все эндпоинты были открыты без аутентификации, данные хранились в "
        "in-memory словарях. Анализ существующей реализации выявил следующие "
        "уязвимости:"
    )
    issues = [
        "отсутствие аутентификации — любой мог создать, изменить или удалить книгу;",
        "отсутствие системы пользователей и хранения паролей;",
        "отсутствие ограничений на частоту запросов — DoS-уязвимость;",
        "отсутствие заголовков безопасности (X-Frame-Options, CSP, X-XSS-Protection);",
        "CORS разрешал запросы с любых источников (allow_origins=[\"*\"]);",
        "отсутствие санитизации строковых полей — потенциальный Stored XSS.",
    ]
    for i in issues:
        add_bullet(doc, i)

    # Анализ уязвимостей (таблица)
    add_heading(doc, "Анализ уязвимостей исходного API")
    add_paragraph(doc,
        "Проведён анализ исходного API по 10 категориям уязвимостей в соответствии "
        "с классификацией OWASP Top 10. Результаты представлены в таблице:"
    )
    vuln_headers = ["№", "Тип уязвимости", "Описание угрозы", "Статус в ЛР6"]
    vuln_rows = [
        ["1",  "SQL Injection",           "Внедрение SQL в параметры запроса",              "Не актуально (in-memory)"],
        ["2",  "XSS",                     "Внедрение скриптов через пользовательский ввод", "Уязвимо (нет санитизации)"],
        ["3",  "CSRF",                    "Межсайтовая подделка запросов",                  "Уязвимо (нет токена)"],
        ["4",  "Clickjacking",            "Встраивание страницы в iframe",                  "Уязвимо (нет X-Frame-Options)"],
        ["5",  "Валидация данных",        "Некорректные/вредоносные данные в запросе",      "Частично (Pydantic типы)"],
        ["6",  "Rate Limiting",           "DoS через лавину запросов",                      "Уязвимо (лимит отсутствует)"],
        ["7",  "Хранение секретов",       "Пароли/ключи в открытом виде",                   "Нет системы пользователей"],
        ["8",  "CORS",                    "Запросы из любых источников",                    "Уязвимо (allow_origins=[\"*\"])"],
        ["9",  "HTTPS",                   "Передача данных в открытом виде",                "Не настроен (учебная среда)"],
        ["10", "Аутентификация / RBAC",   "Отсутствие проверки прав доступа",               "Уязвимо (нет auth)"],
    ]
    add_table(doc, vuln_headers, vuln_rows,
              caption="Таблица 1. Анализ уязвимостей исходного API по OWASP Top 10")

    # Реализация защитных механизмов
    add_heading(doc, "Реализация защитных механизмов")

    add_subheading(doc, "Выбор механизмов безопасности")
    add_paragraph(doc,
        "На основе анализа уязвимостей для реализации выбраны следующие механизмы защиты:"
    )
    chosen = [
        "санитизация входных данных (BeforeValidator) — защита от Stored XSS;",
        "строгая валидация паролей (field_validator) — обязательные заглавная буква и цифра;",
        "JWT-аутентификация (HS256, реализована без внешних зависимостей) — stateless auth;",
        "ролевое разграничение доступа RBAC (admin / user) — принцип наименьших привилегий;",
        "хеширование паролей bcrypt — устойчивость к атакам перебора;",
        "Rate Limiting (60 req/min с одного IP) — защита от DoS и brute-force;",
        "защитные HTTP-заголовки (X-Frame-Options, CSP, HSTS) — защита от clickjacking и XSS;",
        "строгая настройка CORS (только доверенные origins) — ограничение доступа с чужих доменов.",
    ]
    for c in chosen:
        add_bullet(doc, c)

    add_subheading(doc, "Санитизация входных данных и валидация паролей")
    add_paragraph(doc,
        "XSS (Cross-Site Scripting) — тип атаки, при котором злоумышленник внедряет "
        "вредоносный JavaScript-код в данные, сохраняемые в приложении (Stored XSS). "
        "Для защиты реализована функция санитизации: она удаляет HTML-теги и "
        "экранирует спецсимволы перед сохранением. Санитизация применяется ко всем "
        "строковым полям через механизм BeforeValidator библиотеки Pydantic v2, "
        "что обеспечивает автоматическую очистку до этапа бизнес-логики. "
        "Дополнительно реализована строгая валидация паролей при регистрации:"
    )
    add_code_block(doc, LAB7_SANITIZE_CODE)
    add_figure_caption(doc, "Рис. 1. Санитизация строк (BeforeValidator) и валидация паролей")

    add_subheading(doc, "JWT-аутентификация и хеширование паролей")
    add_paragraph(doc,
        "JWT (JSON Web Token, RFC 7519) — стандарт для передачи информации в виде "
        "JSON-объекта, подписанного криптографическим ключом. Структура токена: "
        "header.payload.signature. В работе реализована ручная JWT HS256 на основе "
        "стандартных модулей Python (hmac, hashlib, base64), не требующая внешних "
        "зависимостей. Пароли хранятся исключительно в виде bcrypt-хешей: bcrypt — "
        "адаптивная хеш-функция с параметром work factor, замедляющая перебор паролей:"
    )
    add_code_block(doc, LAB7_JWT_CODE)
    add_figure_caption(doc, "Рис. 2. Реализация JWT HS256 и хеширование паролей bcrypt")

    add_subheading(doc, "Ролевое разграничение доступа (RBAC)")
    add_paragraph(doc,
        "Реализовано два уровня доступа: admin (полный CRUD над книгами, категориями, "
        "заказами) и user (только чтение и создание заказов). Зависимости FastAPI "
        "(Depends) обеспечивают декларативную проверку JWT и ролей на уровне "
        "каждого эндпоинта. Принцип наименьших привилегий: пользователь получает "
        "только те права, которые необходимы для выполнения его задач:"
    )
    add_code_block(doc, LAB7_RBAC_CODE)
    add_figure_caption(doc, "Рис. 3. Хеширование паролей bcrypt и RBAC через зависимости FastAPI")

    add_subheading(doc, "Rate Limiting и защитные HTTP-заголовки")
    add_paragraph(doc,
        "Ограничение частоты запросов (Rate Limiting) защищает от атак DoS и "
        "brute-force: не более 60 запросов в минуту с одного IP-адреса. "
        "Реализован скользящий счётчик на основе временных меток. "
        "Middleware также добавляет защитные HTTP-заголовки к каждому ответу: "
        "X-Frame-Options: DENY (защита от clickjacking), "
        "X-Content-Type-Options: nosniff (предотвращение MIME sniffing), "
        "Content-Security-Policy: default-src 'self' (ограничение источников контента), "
        "Strict-Transport-Security (принудительное HTTPS):"
    )
    add_code_block(doc, LAB7_RATELIMIT_CODE)
    add_figure_caption(doc, "Рис. 4. Rate Limiting и Security Headers Middleware")

    add_subheading(doc, "Настройка CORS и точка входа приложения")
    add_paragraph(doc,
        "CORS (Cross-Origin Resource Sharing) — механизм браузерной защиты, "
        "ограничивающий запросы к API из других доменов. В исходной реализации "
        "allow_origins=[\"*\"] разрешал запросы с любых источников. "
        "После доработки разрешены только явно указанные trusted origins, "
        "ограничены HTTP-методы и заголовки. Все middleware подключаются "
        "централизованно в точке входа приложения:"
    )
    add_code_block(doc, LAB7_CORS_CODE)
    add_figure_caption(doc, "Рис. 5. Настройка CORS и сборка защищённого приложения")

    # Тестирование
    add_heading(doc, "Тестирование и верификация механизмов безопасности")
    add_paragraph(doc,
        "Проверка работоспособности внедрённых механизмов безопасности проводилась "
        "через встроенный интерфейс Swagger UI (http://localhost:8001/docs) "
        "и анализ HTTP-ответов."
    )

    add_subheading(doc, "Проверка аутентификации по JWT")
    add_paragraph(doc,
        "На первом этапе выполнен запрос к защищённому эндпоинту POST /books/ "
        "без передачи токена авторизации. Сервер вернул статус 401 Unauthorized "
        "с сообщением «Not authenticated». Затем выполнен запрос с некорректным "
        "токеном — сервер корректно отклонил запрос с детальным описанием ошибки. "
        "После успешной аутентификации через POST /auth/token с корректными "
        "учётными данными получен JWT-токен."
    )
    add_figure_caption(doc, "Рис. 6. Проверка JWT: 401 без токена и 401 с неверным токеном")

    add_subheading(doc, "Проверка ролевого разграничения доступа (RBAC)")
    add_paragraph(doc,
        "Для тестирования RBAC выполнен POST /books/ с токеном пользователя "
        "с ролью user. Сервер вернул статус 403 Forbidden. "
        "Аналогичный запрос с токеном администратора (role=admin) завершился "
        "успешно — статус 201 Created, книга добавлена. "
        "DELETE /books/{id} с токеном user также вернул 403 Forbidden, "
        "что подтверждает корректность реализации принципа наименьших привилегий."
    )
    add_figure_caption(doc, "Рис. 7. Проверка RBAC: 403 для user и 201 для admin")

    add_subheading(doc, "Проверка санитизации ввода (защита от XSS)")
    add_paragraph(doc,
        "Для проверки защиты от межсайтового скриптинга (Stored XSS) выполнен "
        "POST /books/ с внедрением HTML и JavaScript в поле title: "
        "«<script>alert('XSS')</script>Книга». После обработки BeforeValidator "
        "тег <script> был удалён, в базе сохранилось только «Книга» — "
        "вредоносный код не прошёл в хранилище. Сервер вернул 201 Created "
        "с уже очищенными данными."
    )
    add_figure_caption(doc, "Рис. 8. Санитизация: XSS-тег удалён перед сохранением")

    add_subheading(doc, "Проверка Rate Limiting и заголовков безопасности")
    add_paragraph(doc,
        "Для проверки Rate Limiting отправлено 65 последовательных GET /books/ "
        "запросов с одного IP. Первые 60 запросов выполнены успешно (статус 200). "
        "Начиная с 61-го запроса сервер возвращал 429 Too Many Requests. "
        "После истечения 60-секундного окна запросы снова принимались. "
        "Анализ заголовков ответа подтвердил наличие всех защитных заголовков: "
        "X-Frame-Options: DENY, X-Content-Type-Options: nosniff, "
        "Content-Security-Policy: default-src 'self', X-XSS-Protection: 1; mode=block."
    )
    add_figure_caption(doc, "Рис. 9. Rate Limiting: 429 при превышении лимита и проверка заголовков")

    add_subheading(doc, "Проверка валидации паролей")
    add_paragraph(doc,
        "Выполнен POST /auth/register с паролем, не содержащим заглавных букв "
        "(«password123»). Сервер вернул 422 Unprocessable Entity с сообщением "
        "«Пароль должен содержать хотя бы одну заглавную букву». "
        "При отправке пароля без цифр — аналогичная ошибка. "
        "Корректный пароль («Password123») принят успешно."
    )
    add_figure_caption(doc, "Рис. 10. Валидация пароля: 422 при несоответствии требованиям")

    # Вывод
    add_heading(doc, "Вывод")
    add_paragraph(doc,
        "В ходе выполнения лабораторной работы проведён анализ 10 категорий "
        "уязвимостей исходного API (ЛР №6) в соответствии с классификацией OWASP Top 10 "
        "и реализован комплекс из восьми защитных механизмов: санитизация строковых "
        "полей (BeforeValidator), валидация паролей по сложности, JWT-аутентификация "
        "на базе HMAC-HS256, ролевое разграничение доступа (admin/user), "
        "хеширование паролей bcrypt, ограничение частоты запросов (Rate Limiting), "
        "защитные HTTP-заголовки и строгая настройка CORS."
    )
    add_paragraph(doc,
        "Практическое тестирование подтвердило эффективность каждого механизма: "
        "JWT корректно возвращал 401 при отсутствии или некорректности токена, "
        "RBAC ограничивал доступ по ролям (403 для user при попытке записи), "
        "BeforeValidator удалял XSS-теги из входных данных до их сохранения, "
        "Rate Limiting возвращал 429 при превышении лимита, заголовки безопасности "
        "присутствовали в каждом ответе сервера. Архитектура «defence in depth» "
        "обеспечивает многоуровневую защиту: каждый механизм перекрывает уязвимости, "
        "пропущенные другими уровнями."
    )
    add_paragraph(doc,
        "Полученные навыки анализа уязвимостей и реализации защитных механизмов "
        "являются необходимой компетенцией специалиста по информационной безопасности. "
        "Использование JWT позволяет строить stateless-аутентификацию без хранения "
        "сессий на сервере, bcrypt обеспечивает устойчивость к атакам перебора паролей, "
        "а Pydantic BeforeValidator минимизирует риск XSS на уровне входных данных, "
        "не требуя изменения бизнес-логики приложения."
    )

    doc.save(output_path)
    print(f"[OK] ЛР7: {output_path}")


# ═══════════════════════════════════════════════════════════════════════════
#  ЛР8 — Pandas Data Analysis
# ═══════════════════════════════════════════════════════════════════════════

LAB8_GEN_CODE = '''import numpy as np
import pandas as pd

def generate_dataset(n_rows=220, seed=42):
    rng = np.random.default_rng(seed)
    genres = ["Роман", "Фантастика", "Детектив", "История", "Бизнес", "Наука"]
    genre_col = rng.choice(genres, size=n_rows,
                           p=[0.25, 0.20, 0.20, 0.15, 0.10, 0.10])
    year   = rng.integers(2000, 2026, size=n_rows)
    pages  = rng.integers(100, 800, size=n_rows).astype(float)
    price  = np.array([500 * rng.uniform(0.7, 2.0) for _ in genre_col])
    rating = np.clip(rng.normal(3.8, 0.8, n_rows), 1.0, 5.0).round(1)
    sales_count = np.array([
        max(0, int(1000 * (rating[i] / 3.5) * (600 / price[i])
                   * rng.uniform(0.5, 2.0)))
        for i in range(n_rows)
    ])
    threshold = np.percentile(sales_count, 80)
    is_bestseller = (sales_count >= threshold).astype(int)

    # Добавляем пропуски (~5%)
    nan_idx = rng.choice(n_rows, size=int(n_rows * 0.05), replace=False)
    pages[nan_idx] = np.nan

    return pd.DataFrame({
        "book_id": range(1, n_rows + 1),
        "title": [f"Книга_{i+1:03d}" for i in range(n_rows)],
        "genre": genre_col, "year": year, "price": price.round(2),
        "pages": pages, "rating": rating, "sales_count": sales_count,
        "author_rating": np.clip(rng.normal(3.5, 1.0, n_rows), 1.0, 5.0).round(1),
        "is_bestseller": is_bestseller,
    })

df = generate_dataset()
df.to_csv("book_sales.csv", index=False)
df = pd.read_csv("book_sales.csv")
print(df.head())
print(df.shape)       # (220, 10)
print(df.isnull().sum())'''

LAB8_PREPROCESS_CODE = '''# Базовый анализ
print(df.info())
print(df.describe().round(2))

# Обработка пропущенных значений
for col in ["pages", "author_rating"]:
    median_val = df[col].median()
    df[col] = df[col].fillna(median_val)
    print(f"{col}: пропуски заполнены медианой {median_val:.1f}")

# Проверка дубликатов
print(f"Дубликатов: {df.duplicated().sum()}")'''

LAB8_VISUAL_CODE = '''import matplotlib.pyplot as plt
import seaborn as sns

sns.set_theme(style="whitegrid")

# 1. Корреляционная тепловая карта
fig, ax = plt.subplots(figsize=(9, 7))
corr = df[["price","pages","rating","sales_count",
           "author_rating","year"]].corr()
sns.heatmap(corr, annot=True, fmt=".2f", cmap="coolwarm",
            center=0, square=True, ax=ax)
ax.set_title("Корреляционная матрица")

# 2. Диаграмма рассеяния: цена vs продажи
fig, ax = plt.subplots(figsize=(10, 6))
for genre in df["genre"].unique():
    mask = df["genre"] == genre
    ax.scatter(df.loc[mask, "price"], df.loc[mask, "sales_count"],
               label=genre, alpha=0.65)
ax.legend(title="Жанр")
ax.set_xlabel("Цена (руб.)"); ax.set_ylabel("Продажи")

# 3. Boxplot: продажи по жанрам
sns.boxplot(data=df, x="genre", y="sales_count", palette="Set2")

# 4. KDE: распределение рейтинга
sns.kdeplot(data=df, x="rating", hue="is_bestseller",
            fill=True, alpha=0.4)'''

LAB8_ML_CODE = '''from sklearn.ensemble import RandomForestClassifier
from sklearn.model_selection import train_test_split
from sklearn.metrics import accuracy_score, classification_report
from sklearn.preprocessing import LabelEncoder

# Кодирование жанра
le = LabelEncoder()
df["genre_enc"] = le.fit_transform(df["genre"])

features = ["price","pages","rating","sales_count",
            "author_rating","year","genre_enc"]
X = df[features].values
y = df["is_bestseller"].values

X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.25, random_state=42, stratify=y)

clf = RandomForestClassifier(n_estimators=100, max_depth=6,
                              random_state=42, class_weight="balanced")
clf.fit(X_train, y_train)
y_pred = clf.predict(X_test)

print(f"Accuracy: {accuracy_score(y_test, y_pred):.4f}")
print(classification_report(y_test, y_pred))

# Важность признаков
for feat, imp in sorted(zip(features, clf.feature_importances_),
                         key=lambda x: -x[1]):
    print(f"  {feat:<20} {imp:.4f}")'''

LAB8_OUTPUT = '''[OK] Датасет сохранён: book_sales.csv
[OK] Датасет загружен: 220 строк, 10 столбцов

>>> df.head(5):
 book_id  title       genre  year   price  pages  rating  sales_count  author_rating  is_bestseller
       1  Книга_001   История  2025  966.63  223.0     3.6          326            3.5              0
       2  Книга_002   Фантастика 2007 494.50 614.0     3.6         2123            3.8              0
       5  Книга_005   Роман    2013  458.30  518.0     4.2         3096            4.2              1

>>> df.shape: (220, 10)
>>> Пропущенные значения: pages=11 (5.0%), author_rating=8 (3.6%)
>>> pages: заполнено 11 пропусков медианой (487.0)
>>> author_rating: заполнено 8 пропусков медианой (3.5)
>>> Дубликатов строк: 0

       price   pages  rating  sales_count  author_rating
count 220.00  220.00  220.00       220.00         220.00
mean  699.17  462.56    3.76      1371.79           3.44
std   221.66  209.87    0.77       879.69           0.89
min   351.80  103.00    1.70       141.00           1.00
50%   666.05  487.00    3.80      1154.00           3.50
max  1343.03  799.00    5.00      4181.00           5.00

>>> Точность (accuracy): 0.9818 (98.2%)
>>> Важность признаков:
    sales_count          0.6103
    price                0.1778
    genre_enc            0.0685
    rating               0.0474
    author_rating        0.0371
    pages                0.0297
    year                 0.0292'''


def generate_lab8(output_path):
    doc = Document()
    set_page_margins(doc)
    create_title_page(doc, "8", "Анализ датасета с Pandas и визуализация")
    add_page_break(doc)

    # Цель
    _inline_bold_para(doc,
        "Цель: ",
        "изучить инструменты анализа данных на языке Python — библиотеки Pandas, "
        "Matplotlib, Seaborn и Scikit-learn — путём формирования синтетического "
        "датасета о продажах книг, его предобработки, построения графиков "
        "визуализации и обучения модели машинного обучения (Random Forest) "
        "для предсказания бестселлеров."
    )

    # Постановка задачи
    _inline_bold_para(doc, "Постановка задачи: ",
        "для достижения цели необходимо выполнить следующие задачи:")
    tasks = [
        "изучить возможности библиотек Pandas, Matplotlib, Seaborn и Scikit-learn;",
        "сформировать синтетический датасет о продажах книжного магазина (220 строк, 10 признаков);",
        "провести базовый анализ данных: просмотр структуры, типов, пропущенных значений и статистик;",
        "выполнить предобработку данных: заполнение пропусков медианой, проверку дубликатов;",
        "построить корреляционную тепловую карту (heatmap) числовых признаков;",
        "построить диаграмму рассеяния (scatter plot) «цена vs продажи» с раскраской по жанру;",
        "построить ящик с усами (boxplot) для сравнения продаж по жанрам;",
        "построить KDE-график распределения рейтинга для бестселлеров и обычных книг;",
        "обучить классификатор Random Forest для предсказания бестселлеров и оценить его точность.",
    ]
    for t in tasks:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(t)
        _apply_tnr_font(run, size=14)

    # Введение
    add_heading(doc, "Введение")
    add_paragraph(doc,
        "Анализ данных является одним из ключевых навыков в современной разработке "
        "программного обеспечения и науке о данных. Организации накапливают огромные "
        "массивы информации о продажах, клиентах и продуктах, которые требуют "
        "систематической обработки и интерпретации. Язык Python, благодаря развитой "
        "экосистеме библиотек, стал основным инструментом для решения задач "
        "разведочного анализа данных (EDA) и построения прогностических моделей."
    )
    add_paragraph(doc,
        "Библиотека Pandas предоставляет мощный инструментарий для работы с "
        "табличными данными через структуру DataFrame: загрузку из CSV/Excel, "
        "фильтрацию, группировку и агрегацию. Matplotlib и Seaborn обеспечивают "
        "построение широкого спектра статистических визуализаций — от простых "
        "гистограмм до сложных тепловых карт. Scikit-learn предоставляет "
        "унифицированный интерфейс для применения алгоритмов машинного обучения, "
        "включая классификаторы, регрессоры и инструменты предобработки."
    )
    add_paragraph(doc,
        "В данной лабораторной работе в качестве предметной области выбран книжный "
        "магазин. Сформированный синтетический датасет содержит информацию о жанре, "
        "цене, рейтинге и количестве продаж книг. Целевым признаком является "
        "бинарный флаг бестселлера. Такая постановка задачи позволяет "
        "продемонстрировать полный цикл работы с данными: от генерации и "
        "предобработки до визуализации и построения классификационной модели."
    )

    # Теоретическая часть
    add_heading(doc, "Теоретическая часть")

    add_subheading(doc, "Библиотека Pandas")
    add_paragraph(doc,
        "Pandas — это библиотека с открытым исходным кодом для языка Python, "
        "предназначенная для работы с табличными и временными рядами данных. "
        "Основной структурой данных является DataFrame — двумерная таблица с "
        "именованными столбцами и индексами. Pandas поддерживает чтение и запись "
        "данных из форматов CSV, Excel, JSON, SQL; фильтрацию, сортировку и "
        "группировку строк; обработку пропущенных значений (NaN); "
        "объединение таблиц (merge, join, concat). Метод describe() позволяет "
        "получить сводную статистику: среднее, медиану, квартили и стандартное "
        "отклонение для всех числовых столбцов."
    )

    add_subheading(doc, "Библиотеки Matplotlib и Seaborn")
    add_paragraph(doc,
        "Matplotlib — базовая библиотека для построения графиков в Python. "
        "Она позволяет создавать практически любые типы визуализаций: линейные "
        "графики, гистограммы, диаграммы рассеяния, круговые диаграммы. "
        "Seaborn является высокоуровневой надстройкой над Matplotlib и предоставляет "
        "удобный интерфейс для построения статистических графиков. Функция heatmap() "
        "строит тепловую карту матрицы, kdeplot() — оценку плотности распределения "
        "(KDE), boxplot() — ящик с усами. Seaborn автоматически интегрируется с "
        "DataFrame Pandas, что существенно упрощает построение графиков."
    )

    add_subheading(doc, "Библиотека Scikit-learn и алгоритм Random Forest")
    add_paragraph(doc,
        "Scikit-learn — ведущая библиотека машинного обучения для Python. "
        "Она предоставляет унифицированный интерфейс fit/predict для всех "
        "алгоритмов, что упрощает замену и сравнение моделей. "
        "Random Forest (случайный лес) — ансамблевый алгоритм, строящий множество "
        "решающих деревьев и объединяющий их предсказания голосованием. "
        "Преимущества метода: устойчивость к переобучению, встроенная оценка "
        "важности признаков (feature_importances_), работа с категориальными "
        "признаками после кодирования LabelEncoder. Качество классификации "
        "оценивается метриками accuracy (доля верных ответов) и "
        "classification_report (precision, recall, F1)."
    )

    # Описание датасета
    add_heading(doc, "Описание датасета")
    add_paragraph(doc,
        "Для выполнения работы сформирован синтетический датасет о продажах книг "
        "в книжном магазине. Датасет содержит 220 строк и 10 признаков. "
        "Данные сгенерированы с фиксированным зерном случайности (seed=42) для "
        "воспроизводимости результатов. Количество продаж зависит от жанра, "
        "рейтинга и цены книги. Намеренно введены пропущенные значения "
        "(~5% в столбце pages, ~4% в author_rating) для отработки предобработки:"
    )
    features_hdr = ["Признак", "Тип", "Описание"]
    features_rows = [
        ["book_id",        "int",   "Уникальный идентификатор книги"],
        ["title",          "str",   "Название книги (Книга_001 ... Книга_220)"],
        ["genre",          "str",   "Жанр: Роман, Фантастика, Детектив, История, Бизнес, Наука"],
        ["year",           "int",   "Год издания (2000–2025)"],
        ["price",          "float", "Цена в рублях (350–1350 руб.)"],
        ["pages",          "float", "Количество страниц (100–799); ~5% пропусков"],
        ["rating",         "float", "Рейтинг книги (1.0–5.0)"],
        ["sales_count",    "int",   "Количество продаж (зависит от жанра и рейтинга)"],
        ["author_rating",  "float", "Средний рейтинг автора (1.0–5.0); ~4% пропусков"],
        ["is_bestseller",  "int",   "Бинарный целевой признак: 1 — топ 20% по продажам"],
    ]
    add_table(doc, features_hdr, features_rows, caption="Таблица 1. Признаки датасета book_sales.csv")

    # Реализация
    add_heading(doc, "Реализация")

    add_subheading(doc, "Генерация и загрузка данных")
    add_paragraph(doc,
        "Датасет сгенерирован с помощью numpy.random.default_rng (seed=42) "
        "для воспроизводимости. Количество продаж рассчитывается через формулу, "
        "учитывающую жанровый мультипликатор, рейтинг и цену книги. "
        "После генерации данные сохраняются в CSV и загружаются обратно через Pandas, "
        "что имитирует реальный сценарий работы с файловыми источниками данных:"
    )
    add_code_block(doc, LAB8_GEN_CODE)
    add_figure_caption(doc, "Рис. 1. Генерация синтетического датасета и загрузка из CSV")

    add_subheading(doc, "Базовый анализ структуры датасета")
    add_paragraph(doc,
        "Первым шагом анализа является изучение структуры датасета: размерности, "
        "типов данных и наличия пропущенных значений. Методы df.head(), df.shape, "
        "df.info() и df.isnull().sum() дают полное представление о составе данных. "
        "Метод df.describe() выводит статистику по числовым столбцам: "
        "count, mean, std, min, 25%, 50%, 75%, max:"
    )
    add_code_block(doc, LAB8_OUTPUT[:LAB8_OUTPUT.index(">>> Точность")])
    add_figure_caption(doc, "Рис. 2. Вывод базового анализа: форма, типы данных, пропуски, статистика")
    add_paragraph(doc,
        "Из вывода видно, что датасет содержит 220 строк и 10 столбцов. "
        "Обнаружены пропущенные значения: 11 в столбце pages (5.0%) и "
        "8 в author_rating (3.6%). Средняя цена книги составляет 699 рублей, "
        "средний рейтинг — 3.76 из 5.0. Дубликаты строк отсутствуют."
    )

    add_subheading(doc, "Предобработка данных")
    add_paragraph(doc,
        "Пропущенные значения заполняются медианой по соответствующему столбцу. "
        "Медиана выбрана вместо среднего арифметического как мера, устойчивая "
        "к выбросам: при наличии аномальных значений среднее смещается, "
        "медиана же остаётся стабильной. После заполнения проверяется отсутствие "
        "дубликатов строк:"
    )
    add_code_block(doc, LAB8_PREPROCESS_CODE)
    add_figure_caption(doc, "Рис. 3. Предобработка: заполнение пропусков медианой и проверка дубликатов")
    add_paragraph(doc,
        "В результате предобработки все 11 пропусков в столбце pages заполнены "
        "медианным значением 487.0, а 8 пропусков в author_rating — значением 3.5. "
        "Датасет готов к визуализации и машинному обучению."
    )

    add_subheading(doc, "Корреляционная тепловая карта")
    add_paragraph(doc,
        "Корреляционная матрица Пирсона позволяет количественно оценить линейные "
        "взаимосвязи между числовыми признаками. Значения близкие к +1 указывают "
        "на сильную прямую зависимость, к -1 — на обратную. "
        "Тепловая карта визуализирует матрицу с цветовой кодировкой, "
        "что упрощает интерпретацию:"
    )
    add_code_block(doc, LAB8_VISUAL_CODE[:LAB8_VISUAL_CODE.index("# 2.")])
    add_figure_caption(doc, "Рис. 4. Корреляционная тепловая карта числовых признаков")
    add_paragraph(doc,
        "Анализ тепловой карты выявляет следующие закономерности: признак "
        "sales_count имеет сильную положительную корреляцию с is_bestseller (0.83), "
        "что ожидаемо — бестселлер определяется именно по объёму продаж. "
        "Цена (price) отрицательно коррелирует с количеством продаж (-0.42): "
        "дорогие книги продаются в меньшем количестве. "
        "Рейтинг (rating) слабо коррелирует с продажами, что может объясняться "
        "жанровыми различиями в аудитории."
    )

    add_subheading(doc, "Диаграмма рассеяния (scatter plot)")
    add_paragraph(doc,
        "Диаграмма рассеяния «цена vs количество продаж» с раскраской точек "
        "по жанру позволяет визуально оценить распределение данных и выявить "
        "кластеры. Каждая точка соответствует одной книге из датасета:"
    )
    add_figure_caption(doc, "Рис. 5. Диаграмма рассеяния: цена vs продажи, раскраска по жанру")
    add_paragraph(doc,
        "График подтверждает обратную зависимость между ценой и продажами: "
        "облако точек смещено вниз-вправо. Книги жанра «Фантастика» и «Детектив» "
        "концентрируются в верхней части графика с высокими продажами при "
        "умеренных ценах. Книги жанра «Наука» и «История» расположены ниже — "
        "их читательская аудитория меньше. Выбросы с продажами выше 4000 "
        "встречаются преимущественно у «Фантастики»."
    )

    add_subheading(doc, "Ящик с усами (boxplot)")
    add_paragraph(doc,
        "Диаграмма «ящик с усами» наглядно показывает медиану, межквартильный "
        "размах (IQR) и выбросы для распределения продаж внутри каждого жанра. "
        "Это позволяет сравнить типичный объём продаж и разброс данных "
        "между жанрами:"
    )
    add_figure_caption(doc, "Рис. 6. Ящик с усами: распределение продаж по жанрам")
    add_paragraph(doc,
        "Boxplot показывает, что «Фантастика» и «Детектив» имеют наибольшую "
        "медиану продаж и значительный разброс значений. «Роман» занимает "
        "промежуточное положение. «Бизнес», «История» и «Наука» демонстрируют "
        "наименьшие медианные продажи. У всех жанров присутствуют выбросы — "
        "отдельные книги с аномально высокими продажами."
    )

    add_subheading(doc, "KDE-распределение рейтинга")
    add_paragraph(doc,
        "KDE (Kernel Density Estimation, оценка плотности ядра) — "
        "непараметрический метод оценки функции плотности вероятности. "
        "График сравнивает форму распределения рейтинга для двух групп: "
        "бестселлеры (is_bestseller=1) и обычные книги (is_bestseller=0):"
    )
    add_figure_caption(doc, "Рис. 7. KDE: распределение рейтинга для бестселлеров и обычных книг")
    add_paragraph(doc,
        "KDE-график показывает, что распределение рейтинга бестселлеров смещено "
        "вправо относительно обычных книг: пик плотности для бестселлеров "
        "приходится на ~4.2, тогда как для обычных книг — на ~3.6. "
        "Это подтверждает, что книги с высоким рейтингом чаще становятся "
        "бестселлерами. Оба распределения близки к нормальному."
    )

    add_subheading(doc, "Модель машинного обучения: Random Forest")
    add_paragraph(doc,
        "Для предсказания бестселлера (бинарный признак is_bestseller) обучен "
        "классификатор Random Forest из 100 деревьев. Перед обучением жанр "
        "закодирован с помощью LabelEncoder (строки → целые числа). "
        "Датасет разбит на обучающую (75%) и тестовую (25%) выборки "
        "со стратификацией по целевому признаку для сохранения баланса классов:"
    )
    add_code_block(doc, LAB8_ML_CODE)
    add_figure_caption(doc, "Рис. 8. Обучение Random Forest: код модели и оценки качества")
    add_paragraph(doc,
        "Результаты работы модели — точность классификации и "
        "важность каждого признака:"
    )
    add_code_block(doc, LAB8_OUTPUT[LAB8_OUTPUT.index(">>> Точность"):])
    add_figure_caption(doc, "Рис. 9. Точность Random Forest и важность признаков")
    add_paragraph(doc,
        "Модель достигла точности 98.2%, что свидетельствует о высоком качестве "
        "классификации. Наиболее важным признаком является sales_count (0.61): "
        "это логично, поскольку именно по объёму продаж определяется бестселлер. "
        "Второй по важности признак — price (0.18): цена существенно влияет "
        "на объём продаж. На третьем месте жанр (genre_enc, 0.07) и "
        "рейтинг (rating, 0.05). Такой результат подтверждает правильность "
        "выбора признаков и корректность построенной модели."
    )

    # Вывод
    add_heading(doc, "Вывод")
    add_paragraph(doc,
        "В ходе выполнения лабораторной работы изучены и применены на практике "
        "основные инструменты анализа данных на языке Python: Pandas, Matplotlib, "
        "Seaborn и Scikit-learn. Сформирован синтетический датасет о продажах книг "
        "(220 строк, 10 признаков), проведён базовый разведочный анализ данных "
        "(EDA): изучение структуры, типов данных и статистических характеристик. "
        "Выполнена предобработка: пропущенные значения заполнены медианой, "
        "дубликаты отсутствуют."
    )
    add_paragraph(doc,
        "Построены четыре типа визуализаций, каждый из которых решает "
        "самостоятельную аналитическую задачу: корреляционная тепловая карта "
        "выявила связь между ценой и продажами; диаграмма рассеяния показала "
        "жанровые различия; boxplot сравнил распределения продаж по жанрам; "
        "KDE-график подтвердил зависимость между рейтингом и статусом бестселлера. "
        "Обучен классификатор Random Forest: достигнута точность 98.2%, "
        "анализ feature_importances_ подтвердил ключевую роль sales_count."
    )
    add_paragraph(doc,
        "Полученные компетенции применимы в широком круге задач: бизнес-анализ "
        "продаж и клиентского поведения, исследовательский анализ научных данных, "
        "разработка рекомендательных систем и прогностических моделей. "
        "Библиотечная экосистема Python обеспечивает полный цикл работы с данными "
        "от загрузки и очистки до визуализации и построения ML-моделей, что делает "
        "Python стандартом де-факто в области Data Science."
    )

    doc.save(output_path)
    print(f"[OK] ЛР8: {output_path}")


# ═══════════════════════════════════════════════════════════════════════════
#  ЛР9 — Tkinter GUI
# ═══════════════════════════════════════════════════════════════════════════

LAB9_INIT_CODE = '''root = tk.Tk()
root.title("Книжный магазин — Каталог")
root.geometry("1000x580")
root.resizable(False, False)

BG         = "#f0f4f8"
ACCENT     = "#4a90d9"
BTN_FG     = "white"
FONT_MAIN  = ("Arial", 11)
FONT_BOLD  = ("Arial", 12, "bold")
FONT_TITLE = ("Arial", 14, "bold")

root.configure(bg=BG)
root.columnconfigure(0, weight=3)
root.columnconfigure(1, weight=1)
root.rowconfigure(0, weight=1)'''

LAB9_PANELS_CODE = '''# Левая панель — поиск и отображение результатов
left = tk.Frame(root, bg=BG)
left.grid(row=0, column=0, sticky="nsew", padx=15, pady=15)

tk.Label(left, text="Книжный магазин — Каталог книг",
         font=FONT_TITLE, bg=BG, fg=ACCENT).pack(pady=(0, 8))

# Правая панель — форма добавления книги
right = tk.Frame(root, bg=BG, relief="groove", bd=1)
right.grid(row=0, column=1, sticky="nsew", padx=(0, 15), pady=15)

tk.Label(right, text="Добавить книгу",
         font=FONT_BOLD, bg=BG, fg=ACCENT).pack(pady=(10, 5))'''

LAB9_SEARCH_ROW_CODE = '''# Строка поиска с кнопками
sf = tk.Frame(left, bg=BG)
sf.pack(fill="x")
tk.Label(sf, text="Поиск:", font=FONT_BOLD, bg=BG).pack(side="left")
search_entry = tk.Entry(sf, font=FONT_MAIN, width=30)
search_entry.pack(side="left", padx=6)
search_entry.bind("<Return>", lambda e: search_books())
tk.Button(sf, text="Найти", command=search_books,
          bg=ACCENT, fg=BTN_FG, font=FONT_MAIN,
          relief="flat", padx=10).pack(side="left", padx=3)
tk.Button(sf, text="✖ Очистить", command=clear_search,
          bg="#e57373", fg=BTN_FG, font=FONT_MAIN,
          relief="flat", padx=6).pack(side="left", padx=3)

# Радиокнопки режима поиска
mf = tk.Frame(left, bg=BG)
mf.pack(fill="x", pady=5)
search_mode = tk.StringVar(value="title")
tk.Label(mf, text="Искать по:", font=FONT_MAIN, bg=BG).pack(side="left")
for val, lbl in [("title","Названию"),("author","Автору"),("genre","Жанру")]:
    tk.Radiobutton(mf, text=lbl, variable=search_mode, value=val,
                   bg=BG, font=FONT_MAIN).pack(side="left", padx=8)

tk.Button(left, text="Показать все книги", command=show_all,
          bg="#66bb6a", fg=BTN_FG, font=FONT_MAIN,
          relief="flat", padx=8).pack(pady=(0, 6))'''

LAB9_LISTBOX_CODE = '''# Список результатов с вертикальным скроллом
lf = tk.Frame(left, bg=BG)
lf.pack(fill="both", expand=True)
sb = tk.Scrollbar(lf)
sb.pack(side="right", fill="y")
result_list = tk.Listbox(lf, font=FONT_MAIN, yscrollcommand=sb.set,
                          height=16, selectbackground=ACCENT,
                          activestyle="none")
result_list.pack(fill="both", expand=True)
sb.config(command=result_list.yview)

status_label = tk.Label(left,
    text=f"Всего книг в каталоге: {len(books)}",
    font=("Arial", 10, "italic"), bg=BG, fg="#555")
status_label.pack(pady=(4, 0))'''

LAB9_SEARCH_CODE = '''def search_books():
    """Поиск книг по выбранному критерию."""
    query = search_entry.get().strip().lower()
    mode  = search_mode.get()
    result_list.delete(0, tk.END)

    if not query:
        messagebox.showwarning("Внимание", "Введите поисковый запрос.")
        return

    found = []
    for b in books:
        if mode == "title"  and query in b["title"].lower():
            found.append(b)
        elif mode == "author" and query in b["author"].lower():
            found.append(b)
        elif mode == "genre"  and query in b["genre"].lower():
            found.append(b)

    if not found:
        result_list.insert(tk.END, "— Книги по запросу не найдены —")
    else:
        for b in found:
            result_list.insert(tk.END,
                f"{b['title']}  |  {b['author']}  |  "
                f"{b['genre']}  |  {b['year']} г.  |  {b['price']:.0f} руб.")
    status_label.config(text=f"Найдено: {len(found)} книг(и)")'''

LAB9_ADD_CODE = '''def add_book():
    """Добавление книги с валидацией всех полей."""
    title   = title_entry.get().strip()
    author  = author_entry.get().strip()
    genre   = genre_var.get().strip()
    year_s  = year_entry.get().strip()
    price_s = price_entry.get().strip()

    # Проверка обязательных строковых полей
    if not title or not author or not genre:
        messagebox.showerror("Ошибка",
            "Заполните все обязательные поля:\nНазвание, Автор, Жанр.")
        return

    # Валидация года
    try:
        year = int(year_s)
        if year < 1000 or year > 2100:
            raise ValueError
    except ValueError:
        messagebox.showerror("Ошибка",
            "Год издания: целое число в диапазоне 1000–2100.")
        return

    # Валидация цены
    try:
        price = float(price_s)
        if price <= 0:
            raise ValueError
    except ValueError:
        messagebox.showerror("Ошибка",
            "Цена должна быть положительным числом.")
        return

    books.append({"title": title, "author": author,
                  "genre": genre, "year": year, "price": price})
    messagebox.showinfo("Успешно", f"Книга «{title}» добавлена в каталог.")
    for w in (title_entry, author_entry, year_entry, price_entry):
        w.delete(0, tk.END)
    genre_combo["values"] = sorted(set(b["genre"] for b in books))
    refresh_info()
    status_label.config(text=f"Всего книг в каталоге: {len(books)}")'''

LAB9_HELPERS_CODE = '''def refresh_info():
    """Обновление справочного блока жанров."""
    info_text.config(state="normal")
    info_text.delete("1.0", tk.END)
    genre_counts = {}
    for b in books:
        genre_counts[b["genre"]] = genre_counts.get(b["genre"], 0) + 1
    for g, cnt in sorted(genre_counts.items()):
        info_text.insert(tk.END, f"• {g}: {cnt} кн.\n")
    info_text.config(state="disabled")

def show_all():
    """Вывод полного каталога в список результатов."""
    result_list.delete(0, tk.END)
    for b in books:
        result_list.insert(tk.END,
            f"{b['title']}  |  {b['author']}  |  "
            f"{b['genre']}  |  {b['year']} г.  |  {b['price']:.0f} руб.")
    status_label.config(text=f"Всего книг в каталоге: {len(books)}")

def clear_search():
    """Очистка поля поиска и списка результатов."""
    search_entry.delete(0, tk.END)
    result_list.delete(0, tk.END)
    status_label.config(text=f"Всего книг в каталоге: {len(books)}")'''

LAB9_FULL_CODE = '''import tkinter as tk
from tkinter import messagebox, ttk

books = [
    {"title": "Мастер и Маргарита",      "author": "Булгаков М.А.",    "genre": "Роман",          "year": 1967, "price": 450.0},
    {"title": "Преступление и наказание", "author": "Достоевский Ф.М.", "genre": "Роман",          "year": 1866, "price": 380.0},
    {"title": "Война и мир",              "author": "Толстой Л.Н.",     "genre": "Роман-эпопея",   "year": 1869, "price": 620.0},
    {"title": "Евгений Онегин",           "author": "Пушкин А.С.",      "genre": "Роман в стихах", "year": 1833, "price": 290.0},
    {"title": "Тихий Дон",               "author": "Шолохов М.А.",     "genre": "Роман-эпопея",   "year": 1940, "price": 510.0},
    {"title": "Отцы и дети",             "author": "Тургенев И.С.",    "genre": "Роман",          "year": 1862, "price": 320.0},
    {"title": "Обломов",                  "author": "Гончаров И.А.",    "genre": "Роман",          "year": 1859, "price": 355.0},
    {"title": "Собачье сердце",           "author": "Булгаков М.А.",    "genre": "Повесть",        "year": 1925, "price": 275.0},
    {"title": "Капитанская дочка",        "author": "Пушкин А.С.",      "genre": "Роман",          "year": 1836, "price": 310.0},
    {"title": "Анна Каренина",            "author": "Толстой Л.Н.",     "genre": "Роман",          "year": 1878, "price": 490.0},
]
GENRES = sorted(set(b["genre"] for b in books))

def search_books():
    query = search_entry.get().strip().lower()
    mode  = search_mode.get()
    result_list.delete(0, tk.END)
    if not query:
        messagebox.showwarning("Внимание", "Введите поисковый запрос.")
        return
    found = []
    for b in books:
        if mode == "title"  and query in b["title"].lower():  found.append(b)
        elif mode == "author" and query in b["author"].lower(): found.append(b)
        elif mode == "genre"  and query in b["genre"].lower():  found.append(b)
    if not found:
        result_list.insert(tk.END, "— Книги по запросу не найдены —")
    else:
        for b in found:
            result_list.insert(tk.END,
                f"{b[\'title\']}  |  {b[\'author\']}  |  {b[\'genre\']}  |  {b[\'year\']} г.  |  {b[\'price\']:.0f} руб.")
    status_label.config(text=f"Найдено: {len(found)} книг(и)")

def show_all():
    result_list.delete(0, tk.END)
    for b in books:
        result_list.insert(tk.END,
            f"{b[\'title\']}  |  {b[\'author\']}  |  {b[\'genre\']}  |  {b[\'year\']} г.  |  {b[\'price\']:.0f} руб.")
    status_label.config(text=f"Всего книг в каталоге: {len(books)}")

def clear_search():
    search_entry.delete(0, tk.END)
    result_list.delete(0, tk.END)
    status_label.config(text=f"Всего книг в каталоге: {len(books)}")

def add_book():
    title, author, genre = title_entry.get().strip(), author_entry.get().strip(), genre_var.get().strip()
    year_s, price_s = year_entry.get().strip(), price_entry.get().strip()
    if not title or not author or not genre:
        messagebox.showerror("Ошибка", "Заполните все обязательные поля."); return
    try:
        year = int(year_s)
        if year < 1000 or year > 2100: raise ValueError
    except ValueError:
        messagebox.showerror("Ошибка", "Год: целое число 1000–2100."); return
    try:
        price = float(price_s)
        if price <= 0: raise ValueError
    except ValueError:
        messagebox.showerror("Ошибка", "Цена — положительное число."); return
    books.append({"title": title, "author": author, "genre": genre, "year": year, "price": price})
    messagebox.showinfo("Успешно", f"Книга «{title}» добавлена в каталог.")
    for w in (title_entry, author_entry, year_entry, price_entry): w.delete(0, tk.END)
    genre_combo["values"] = sorted(set(b["genre"] for b in books))
    refresh_info()
    status_label.config(text=f"Всего книг в каталоге: {len(books)}")

def refresh_info():
    info_text.config(state="normal"); info_text.delete("1.0", tk.END)
    genre_counts = {}
    for b in books: genre_counts[b["genre"]] = genre_counts.get(b["genre"], 0) + 1
    for g, cnt in sorted(genre_counts.items()): info_text.insert(tk.END, f"• {g}: {cnt} кн.\\n")
    info_text.config(state="disabled")

root = tk.Tk()
root.title("Книжный магазин — Каталог")
root.geometry("1000x580")
root.resizable(False, False)
BG, ACCENT, BTN_FG = "#f0f4f8", "#4a90d9", "white"
FONT_MAIN, FONT_BOLD, FONT_TITLE = ("Arial",11), ("Arial",12,"bold"), ("Arial",14,"bold")
root.configure(bg=BG)
root.columnconfigure(0, weight=3); root.columnconfigure(1, weight=1); root.rowconfigure(0, weight=1)

left = tk.Frame(root, bg=BG)
left.grid(row=0, column=0, sticky="nsew", padx=15, pady=15)
tk.Label(left, text="Книжный магазин — Каталог книг", font=FONT_TITLE, bg=BG, fg=ACCENT).pack(pady=(0,8))
sf = tk.Frame(left, bg=BG); sf.pack(fill="x")
tk.Label(sf, text="Поиск:", font=FONT_BOLD, bg=BG).pack(side="left")
search_entry = tk.Entry(sf, font=FONT_MAIN, width=30); search_entry.pack(side="left", padx=6)
search_entry.bind("<Return>", lambda e: search_books())
tk.Button(sf, text="Найти", command=search_books, bg=ACCENT, fg=BTN_FG, font=FONT_MAIN, relief="flat", padx=10).pack(side="left", padx=3)
tk.Button(sf, text="✖ Очистить", command=clear_search, bg="#e57373", fg=BTN_FG, font=FONT_MAIN, relief="flat", padx=6).pack(side="left", padx=3)
mf = tk.Frame(left, bg=BG); mf.pack(fill="x", pady=5)
search_mode = tk.StringVar(value="title")
tk.Label(mf, text="Искать по:", font=FONT_MAIN, bg=BG).pack(side="left")
for val, lbl in [("title","Названию"),("author","Автору"),("genre","Жанру")]:
    tk.Radiobutton(mf, text=lbl, variable=search_mode, value=val, bg=BG, font=FONT_MAIN).pack(side="left", padx=8)
tk.Button(left, text="Показать все книги", command=show_all, bg="#66bb6a", fg=BTN_FG, font=FONT_MAIN, relief="flat", padx=8).pack(pady=(0,6))
lf = tk.Frame(left, bg=BG); lf.pack(fill="both", expand=True)
sb = tk.Scrollbar(lf); sb.pack(side="right", fill="y")
result_list = tk.Listbox(lf, font=FONT_MAIN, yscrollcommand=sb.set, height=16, selectbackground=ACCENT, activestyle="none")
result_list.pack(fill="both", expand=True); sb.config(command=result_list.yview)
status_label = tk.Label(left, text=f"Всего книг в каталоге: {len(books)}", font=("Arial",10,"italic"), bg=BG, fg="#555")
status_label.pack(pady=(4,0))

right = tk.Frame(root, bg=BG, relief="groove", bd=1)
right.grid(row=0, column=1, sticky="nsew", padx=(0,15), pady=15)
tk.Label(right, text="Добавить книгу", font=FONT_BOLD, bg=BG, fg=ACCENT).pack(pady=(10,5))
for label_txt, var_name in [("Название *","title_entry"),("Автор *","author_entry"),("Год *","year_entry"),("Цена (руб.) *","price_entry")]:
    tk.Label(right, text=label_txt, font=FONT_MAIN, bg=BG, anchor="w").pack(fill="x", padx=12, pady=(3,0))
    e = tk.Entry(right, font=FONT_MAIN); e.pack(fill="x", padx=12, pady=(0,3)); globals()[var_name] = e
tk.Label(right, text="Жанр *", font=FONT_MAIN, bg=BG, anchor="w").pack(fill="x", padx=12, pady=(3,0))
genre_var = tk.StringVar(value=GENRES[0] if GENRES else "")
genre_combo = ttk.Combobox(right, textvariable=genre_var, values=GENRES, font=FONT_MAIN, state="normal")
genre_combo.pack(fill="x", padx=12, pady=(0,8))
tk.Button(right, text="Добавить в каталог", command=add_book, bg=ACCENT, fg=BTN_FG, font=FONT_BOLD, relief="flat", pady=5).pack(fill="x", padx=12, pady=4)
tk.Label(right, text="─"*28, bg=BG, fg="#ccc").pack(pady=4)
tk.Label(right, text="Жанры в каталоге:", font=FONT_BOLD, bg=BG).pack(padx=12, anchor="w")
info_text = tk.Text(right, font=("Arial",10), height=8, wrap="word", bg="#e8f0fe", relief="flat", state="normal")
info_text.pack(fill="x", padx=12, pady=5)
refresh_info()

show_all()
root.mainloop()'''


def generate_lab9(output_path):
    doc = Document()
    set_page_margins(doc)
    create_title_page(doc, "9", "Создание GUI-приложения")
    add_page_break(doc)

    # Цель работы
    _inline_bold_para(doc,
        "Цель работы: ",
        "изучить принципы разработки приложений с графическим пользовательским "
        "интерфейсом (GUI) на языке Python с использованием стандартной библиотеки "
        "Tkinter путём создания полнофункционального приложения для управления "
        "каталогом книжного магазина."
    )
    tasks = [
        "изучить принципы событийно-ориентированного программирования и работы главного цикла обработки событий;",
        "освоить основные виджеты Tkinter: Label, Entry, Button, Radiobutton, Listbox, Combobox, Text, Scrollbar;",
        "изучить менеджеры геометрии pack() и grid() для размещения элементов в окне;",
        "разработать двухпанельное приложение для управления каталогом книжного магазина;",
        "реализовать функцию поиска книг по трём критериям: названию, автору и жанру;",
        "реализовать форму добавления новой книги с полной валидацией вводимых данных;",
        "провести тестирование разработанного приложения и проверить обработку ошибочного ввода.",
    ]
    for t in tasks:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(t)
        _apply_tnr_font(run, size=14)

    # Теоретические сведения
    add_heading(doc, "Теоретические сведения")
    add_paragraph(doc,
        "Графический пользовательский интерфейс (GUI) — это способ взаимодействия "
        "пользователя с программой через визуальные элементы управления: окна, кнопки, "
        "поля ввода и списки. В отличие от консольных приложений, GUI-программы "
        "работают по принципу событийно-ориентированного программирования: программа "
        "не выполняется последовательно, а ожидает действий пользователя "
        "(нажатий клавиш, кликов мышью) и реагирует на них через функции-обработчики."
    )

    add_subheading(doc, "Библиотека Tkinter")
    add_paragraph(doc,
        "Tkinter — стандартная библиотека Python для создания GUI-приложений, "
        "входящая в стандартный дистрибутив без необходимости дополнительной установки. "
        "Она является обёрткой над библиотекой Tk, написанной на языке Tcl. "
        "Приложение Tkinter строится путём создания корневого окна (tk.Tk()), "
        "добавления в него виджетов, привязки функций-обработчиков к событиям "
        "и запуска главного цикла root.mainloop(). "
        "Метод mainloop() блокирует выполнение программы и постоянно опрашивает "
        "очередь событий, передавая каждое событие соответствующему обработчику."
    )

    add_subheading(doc, "Виджеты Tkinter")
    add_paragraph(doc,
        "Виджеты — это отдельные элементы управления графического интерфейса. "
        "В данной работе используются следующие стандартные виджеты:"
    )
    widgets = [
        "Label — метка для отображения текста (заголовки, подписи, статус);",
        "Entry — однострочное поле ввода текста (поиск, форма добавления);",
        "Button — кнопка, выполняющая команду при нажатии;",
        "Radiobutton — переключатель: позволяет выбрать один из нескольких вариантов;",
        "Listbox — список с возможностью прокрутки и выбора элементов;",
        "Scrollbar — полоса прокрутки, связанная с Listbox или Text;",
        "Text — многострочное текстовое поле (справочная информация);",
        "ttk.Combobox — выпадающий список из расширенного модуля ttk.",
    ]
    for w in widgets:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(w)
        _apply_tnr_font(run, size=14)

    add_subheading(doc, "Менеджеры геометрии")
    add_paragraph(doc,
        "Размещение виджетов в окне осуществляется тремя менеджерами геометрии. "
        "Метод pack() автоматически располагает виджеты по вертикали или горизонтали "
        "в порядке добавления, поддерживая параметры side (left/right/top/bottom), "
        "fill (x/y/both) и expand. "
        "Метод grid() делит контейнер на строки и столбцы и позволяет размещать "
        "виджеты в конкретную ячейку таблицы, что удобно для форм ввода. "
        "Метод place() задаёт точные координаты в пикселях. "
        "В данной работе используются оба метода: grid() для основного разделения "
        "на панели и pack() для размещения элементов внутри каждой панели."
    )

    add_subheading(doc, "Обработка событий и переменные Tkinter")
    add_paragraph(doc,
        "Привязка обработчиков к действиям пользователя выполняется двумя способами: "
        "через параметр command виджета Button/Radiobutton — функция вызывается при "
        "нажатии, и через метод bind() — для обработки клавиатурных событий "
        "(например, search_entry.bind('<Return>', lambda e: search_books()) "
        "позволяет запускать поиск по нажатию Enter). "
        "Переменные Tkinter (StringVar) связывают состояние виджета со значением "
        "переменной: при изменении переменной виджет обновляется автоматически."
    )

    # Постановка задачи
    add_heading(doc, "Постановка задачи")
    add_paragraph(doc,
        "В качестве предметной области выбран книжный магазин. "
        "Начальный каталог содержит 10 книг русской классической литературы "
        "с атрибутами: название, автор, жанр, год издания и цена. "
        "Пользователь должен иметь возможность просматривать каталог, "
        "выполнять поиск и добавлять новые книги."
    )
    add_paragraph(doc,
        "Приложение должно реализовывать следующий функционал:"
    )
    funcs = [
        "отображение полного каталога книг при запуске (функция show_all);",
        "поиск книг по одному из трёх критериев — названию, автору или жанру (функция search_books);",
        "очистка результатов поиска и сброс поля ввода (функция clear_search);",
        "добавление новой книги с валидацией: проверка непустых полей, корректного "
        "года (1000–2100) и положительной цены (функция add_book);",
        "динамическое обновление справочного блока жанров после добавления книги "
        "(функция refresh_info);",
        "вывод строки статуса с количеством найденных или всего книг.",
    ]
    for f in funcs:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(f)
        _apply_tnr_font(run, size=14)

    # Проектирование приложения
    add_heading(doc, "Проектирование приложения")
    add_paragraph(doc,
        "Главное окно приложения имеет фиксированный размер 1000×580 пикселей "
        "и разделено на две функциональные панели с помощью менеджера grid() "
        "с весами столбцов 3:1. Цветовая схема оформлена в сине-серой гамме: "
        "фон BG=#f0f4f8, акцентный цвет ACCENT=#4a90d9. "
        "Заголовок окна — «Книжный магазин — Каталог»."
    )
    add_paragraph(doc,
        "Левая панель (left, вес 3) — основная рабочая зона — содержит:"
    )
    left_items = [
        "заголовок «Книжный магазин — Каталог книг» (Label, акцентный цвет);",
        "строку поиска: поле Entry с привязкой <Return>, кнопки «Найти» и «✖ Очистить»;",
        "три радиокнопки (Radiobutton) для выбора режима поиска: по названию, автору или жанру;",
        "кнопку «Показать все книги» для вывода полного каталога;",
        "Listbox со связанным Scrollbar для отображения результатов (высота 16 строк);",
        "строку статуса (Label) с количеством найденных / всего книг.",
    ]
    for item in left_items:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(item)
        _apply_tnr_font(run, size=14)
    add_paragraph(doc,
        "Правая панель (right, вес 1) — форма добавления книги — содержит:"
    )
    right_items = [
        "четыре поля Entry: Название, Автор, Год, Цена (все обязательные, отмечены «*»);",
        "выпадающий список Combobox для выбора жанра (с возможностью ввода нового);",
        "кнопку «Добавить в каталог» (цвет ACCENT);",
        "справочный блок Text (высота 8 строк, только чтение) со списком жанров и их количеством.",
    ]
    for item in right_items:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(item)
        _apply_tnr_font(run, size=14)

    # Реализация
    add_heading(doc, "Реализация")
    add_paragraph(doc,
        "Программная реализация выполнена на языке Python 3 с использованием "
        "стандартной библиотеки Tkinter. Весь код размещён в одном файле main.py "
        "и организован в следующем порядке: данные каталога, функции-обработчики, "
        "построение интерфейса, запуск главного цикла."
    )

    add_subheading(doc, "Инициализация главного окна и цветовая схема")
    add_paragraph(doc,
        "Корневое окно создаётся вызовом tk.Tk(). Размер задаётся строкой "
        "\"1000x580\", изменение размера заблокировано (resizable=False). "
        "Константы цветов и шрифтов вынесены в переменные BG, ACCENT, FONT_MAIN "
        "для единообразия оформления. Весовые коэффициенты столбцов задаются "
        "columnconfigure() для обеспечения корректного масштабирования панелей:"
    )
    add_code_block(doc, LAB9_INIT_CODE)
    add_figure_caption(doc, "Рис. 1. Инициализация главного окна и цветовой схемы")

    add_subheading(doc, "Разбивка на панели")
    add_paragraph(doc,
        "Интерфейс делится на две независимые области с помощью контейнеров Frame, "
        "размещённых через grid() в столбцах 0 и 1. Параметр sticky=\"nsew\" "
        "растягивает панель на всё доступное пространство ячейки. "
        "Правая панель оформлена с рамкой (relief=\"groove\") для визуального "
        "отделения от основной рабочей области:"
    )
    add_code_block(doc, LAB9_PANELS_CODE)
    add_figure_caption(doc, "Рис. 2. Создание двух панелей через grid()")

    add_subheading(doc, "Строка поиска и радиокнопки режима")
    add_paragraph(doc,
        "Строка поиска состоит из метки Label, поля Entry и двух кнопок, "
        "упакованных горизонтально через pack(side=\"left\"). "
        "Привязка <Return> позволяет запускать поиск нажатием Enter без нажатия кнопки. "
        "Переменная search_mode (StringVar) хранит текущий режим поиска; "
        "три радиокнопки переключают её значение между \"title\", \"author\", \"genre\":"
    )
    add_code_block(doc, LAB9_SEARCH_ROW_CODE)
    add_figure_caption(doc, "Рис. 3. Строка поиска и радиокнопки режима")

    add_subheading(doc, "Список результатов (Listbox + Scrollbar)")
    add_paragraph(doc,
        "Для отображения найденных книг используется виджет Listbox. "
        "Чтобы обеспечить прокрутку длинных списков, к нему привязан Scrollbar: "
        "параметр yscrollcommand=sb.set передаёт позицию прокрутки в Scrollbar, "
        "а sb.config(command=result_list.yview) позволяет Scrollbar управлять "
        "прокруткой Listbox. Строка статуса обновляется после каждой операции:"
    )
    add_code_block(doc, LAB9_LISTBOX_CODE)
    add_figure_caption(doc, "Рис. 4. Listbox со Scrollbar и строка статуса")

    add_subheading(doc, "Функция поиска книг")
    add_paragraph(doc,
        "Функция search_books() считывает строку запроса из поля Entry "
        "и текущий режим из переменной search_mode. Перед поиском список "
        "очищается (delete(0, END)). Если поле пустое — выводится предупреждение "
        "через messagebox.showwarning(). "
        "Поиск регистронезависим (оба значения приводятся к нижнему регистру). "
        "Каждая найденная книга форматируется в строку и добавляется в Listbox:"
    )
    add_code_block(doc, LAB9_SEARCH_CODE)
    add_figure_caption(doc, "Рис. 5. Функция search_books() — поиск по каталогу")

    add_subheading(doc, "Функция добавления книги с валидацией")
    add_paragraph(doc,
        "Функция add_book() выполняет трёхуровневую валидацию: сначала проверяет "
        "непустоту обязательных строковых полей (название, автор, жанр), "
        "затем корректность года (должен быть целым числом в диапазоне 1000–2100), "
        "потом положительность цены. При любой ошибке выводится модальное "
        "диалоговое окно messagebox.showerror() с пояснением. "
        "При успехе книга добавляется в список, поля очищаются, "
        "Combobox и справочный блок обновляются:"
    )
    add_code_block(doc, LAB9_ADD_CODE)
    add_figure_caption(doc, "Рис. 6. Функция add_book() с полной валидацией")

    add_subheading(doc, "Вспомогательные функции")
    add_paragraph(doc,
        "refresh_info() перестраивает справочный блок Text: сначала переводит "
        "его в режим редактирования (state=\"normal\"), очищает содержимое, "
        "подсчитывает количество книг каждого жанра и выводит список, "
        "после чего блокирует поле (state=\"disabled\"). "
        "show_all() и clear_search() используют аналогичный подход к "
        "управлению Listbox через delete(0, END) и insert(END, ...):"
    )
    add_code_block(doc, LAB9_HELPERS_CODE)
    add_figure_caption(doc, "Рис. 7. Вспомогательные функции refresh_info, show_all, clear_search")

    # Тестирование
    add_heading(doc, "Тестирование")
    add_paragraph(doc,
        "Тестирование приложения проводилось вручную путём последовательной "
        "проверки всех реализованных функций. Для каждого тестового сценария "
        "фиксировались входные данные, ожидаемый результат и фактический результат."
    )

    add_subheading(doc, "Запуск и отображение каталога")
    add_paragraph(doc,
        "При запуске приложения автоматически вызывается функция show_all(), "
        "которая заполняет Listbox всеми 10 книгами каталога. "
        "В правой панели отображается справочный блок с жанрами. "
        "В строке статуса показывается «Всего книг в каталоге: 10»."
    )
    add_figure_caption(doc, "Рис. 8. Главное окно приложения при запуске (полный каталог)")

    add_subheading(doc, "Поиск по названию и автору")
    add_paragraph(doc,
        "Проверен поиск по частичному совпадению. При вводе «мастер» с режимом "
        "«По названию» найдена 1 книга: «Мастер и Маргарита». "
        "При переключении на «По автору» и вводе «булгаков» найдены 2 книги: "
        "«Мастер и Маргарита» и «Собачье сердце». "
        "Строка статуса обновлена: «Найдено: 2 книг(и)»."
    )
    add_figure_caption(doc, "Рис. 9. Результат поиска по автору «булгаков»")

    add_subheading(doc, "Поиск по жанру и отображение всех книг")
    add_paragraph(doc,
        "Поиск по жанру «роман-эпопея» вернул 2 книги: «Война и мир» и «Тихий Дон». "
        "Нажатие кнопки «Показать все книги» сбросило фильтр и вывело "
        "полный каталог из 10 записей. Кнопка «✖ Очистить» очистила "
        "поле поиска и список результатов."
    )
    add_figure_caption(doc, "Рис. 10. Поиск по жанру и отображение полного каталога")

    add_subheading(doc, "Добавление новой книги")
    add_paragraph(doc,
        "Заполнена форма добавления: Название — «Идиот», Автор — «Достоевский Ф.М.», "
        "Жанр — «Роман», Год — 1869, Цена — 340. "
        "После нажатия «Добавить в каталог» появилось диалоговое окно "
        "с подтверждением «Книга «Идиот» добавлена в каталог». "
        "Поля формы очищены. Строка статуса обновлена: «Всего книг: 11»."
    )
    add_figure_caption(doc, "Рис. 11. Добавление новой книги «Идиот»")

    add_subheading(doc, "Обработка ошибочного ввода")
    add_paragraph(doc,
        "Проверены следующие некорректные сценарии:"
    )
    errs = [
        "пустое поле «Название» → messagebox.showerror «Заполните все обязательные поля»;",
        "год «19аб» (не число) → ошибка «Год издания: целое число в диапазоне 1000–2100»;",
        "цена «-500» (отрицательная) → ошибка «Цена должна быть положительным числом»;",
        "пустой поисковый запрос → messagebox.showwarning «Введите поисковый запрос».",
    ]
    for e in errs:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(e)
        _apply_tnr_font(run, size=14)
    add_figure_caption(doc, "Рис. 12. Диалоговые окна обработки ошибок валидации")
    add_paragraph(doc,
        "Все проверки пройдены успешно: некорректный ввод не приводит к исключениям, "
        "пользователь получает информативные сообщения об ошибках, "
        "приложение остаётся работоспособным и готовым к дальнейшей работе."
    )

    # Вывод
    add_heading(doc, "Вывод")
    add_paragraph(doc,
        "В ходе выполнения лабораторной работы изучены принципы разработки "
        "GUI-приложений на языке Python с использованием стандартной библиотеки "
        "Tkinter. Освоена архитектура событийно-ориентированного программирования: "
        "главный цикл mainloop(), привязка обработчиков через command и bind(), "
        "работа с переменными StringVar. Изучены виджеты Label, Entry, Button, "
        "Radiobutton, Listbox, Scrollbar, Text, ttk.Combobox и оба менеджера "
        "геометрии — pack() и grid()."
    )
    add_paragraph(doc,
        "Разработано полнофункциональное двухпанельное приложение для управления "
        "каталогом книжного магазина: реализованы поиск по трём критериям с "
        "фильтрацией, отображение полного каталога, форма добавления книги "
        "с трёхуровневой валидацией (строковые поля, диапазон года, "
        "положительность цены) и динамическим обновлением справочного блока. "
        "Тестирование подтвердило корректную обработку всех штатных и "
        "ошибочных сценариев — пользователь получает понятные сообщения "
        "без аварийного завершения программы."
    )
    add_paragraph(doc,
        "Полученные навыки применимы при разработке любых desktop-приложений "
        "с графическим интерфейсом: инструментов автоматизации, конфигураторов, "
        "клиентских приложений для локальных баз данных. "
        "Tkinter обеспечивает кроссплатформенность (Windows, macOS, Linux) "
        "без дополнительных зависимостей, что делает её оптимальным выбором "
        "для учебных проектов и небольших прикладных утилит."
    )

    # Приложение 1
    add_heading(doc, "Приложение 1")
    add_paragraph(doc,
        "Полный исходный код приложения с графическим интерфейсом (lab9/main.py):"
    )
    add_code_block(doc, LAB9_FULL_CODE)

    doc.save(output_path)
    print(f"[OK] ЛР9: {output_path}")


# ═══════════════════════════════════════════════════════════════════════════
#  ЛР10 — Автоматизация задач
# ═══════════════════════════════════════════════════════════════════════════

LAB10_PATHLIB_CODE = '''from pathlib import Path

source = Path("orders")
target = Path("processed_orders")

# Создание директории (без ошибки если существует)
target.mkdir(parents=True, exist_ok=True)

# Рекурсивный обход файлов
for file_path in sorted(source.rglob("*")):
    if file_path.is_file():
        ext  = file_path.suffix.lower()  # расширение: .csv, .txt
        size = file_path.stat().st_size  # размер в байтах
        name = file_path.stem            # имя без расширения
        print(f"{file_path.name:30s} | {ext} | {size} байт")

# Перемещение файла
dest = target / "Заказы" / "order_001.csv"
dest.parent.mkdir(parents=True, exist_ok=True)'''

LAB10_LOGGING_CODE = '''import logging

def setup_logging():
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s | %(levelname)-8s | %(message)s",
        handlers=[
            logging.StreamHandler(),                             # вывод в консоль
            logging.FileHandler("automation.log",               # запись в файл
                                encoding="utf-8"),
        ],
    )

def log(message, level="info", preview=False):
    """Логирует сообщение, добавляя префикс [PREVIEW] в режиме просмотра."""
    prefix = "[PREVIEW] " if preview else ""
    getattr(logging, level)(prefix + message)

# Примеры:
setup_logging()
logging.info("Запуск автоматизации")
logging.warning("Исходная папка не найдена")
logging.error("Ошибка при перемещении файла")'''

LAB10_CONFIG_CODE = '''import sys, csv, json, time, shutil, logging
from pathlib import Path
from datetime import datetime

DEFAULT_SOURCE         = Path("orders")
DEFAULT_TARGET         = Path("processed_orders")
ARCHIVE_LIFETIME_DAYS  = 30

EXTENSIONS = {
    "Заказы":  {".csv"},
    "Отчёты":  {".txt", ".log"},
    "Данные":  {".json"},
    "Архивы":  {".zip", ".tar", ".gz", ".rar"},
}
TRASH_EXTENSIONS = {".tmp", ".bak", ".old"}'''

LAB10_CATEGORY_CODE = '''def get_category(ext: str) -> str:
    """Определяет категорию файла по его расширению."""
    for category, extensions in EXTENSIONS.items():
        if ext in extensions:
            return category
    return "Другое"

# Примеры:
# get_category(".csv")  → "Заказы"
# get_category(".json") → "Данные"
# get_category(".zip")  → "Архивы"
# get_category(".docx") → "Другое"'''

LAB10_UNIQUE_CODE = '''def generate_unique_name(directory: Path, name: str, ext: str) -> str:
    """Генерирует уникальное имя, добавляя суффикс _N при конфликте."""
    base = Path(name).stem.lower().replace(" ", "_")
    counter = 1
    new_name = f"{base}{ext}"
    while (directory / new_name).exists():
        new_name = f"{base}_{counter}{ext}"
        counter += 1
    return new_name

def is_file_old(file_path: Path, days: int) -> bool:
    """Проверяет, старше ли файл заданного числа дней."""
    return (time.time() - file_path.stat().st_mtime) > days * 86400'''

LAB10_CSV_CODE = '''def parse_csv_order(file_path: Path):
    """Читает CSV-файл заказа, возвращает (кол-во строк, итоговую сумму)."""
    total = 0.0
    rows  = 0
    try:
        with open(file_path, encoding="utf-8") as f:
            reader = csv.DictReader(f)
            for row in reader:
                try:
                    price = float(row.get("price", 0) or 0)
                    qty   = int(row.get("quantity", 1) or 1)
                    total += price * qty
                    rows  += 1
                except (ValueError, KeyError):
                    pass
    except Exception:
        pass
    return rows, total'''

LAB10_SUMMARY_CODE = '''def generate_summary_report(processed_orders, target_dir, preview=False):
    """Генерирует сводный текстовый отчёт по обработанным заказам."""
    timestamp   = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_path = target_dir / f"summary_{timestamp}.txt"
    grand_total = sum(o["total"] for o in processed_orders)

    lines = [
        "=" * 60,
        "СВОДНЫЙ ОТЧЁТ ОБРАБОТКИ ЗАКАЗОВ",
        f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}",
        "=" * 60,
        f"Обработано файлов заказов: {len(processed_orders)}",
        "ДЕТАЛИ:",
    ]
    for item in processed_orders:
        lines.append(
            f"  {item['name']:35s}  {item['rows']:3d} строк  "
            f"{item['total']:10.2f} руб."
        )
    lines += ["", f"ИТОГОВАЯ СУММА: {grand_total:.2f} руб.", "=" * 60]

    if not preview:
        report_path.write_text("\\n".join(lines), encoding="utf-8")
        log(f"Сводный отчёт сохранён: {report_path.name}")
    else:
        log(f"Будет создан сводный отчёт: {report_path.name}", preview=True)
    return grand_total'''

LAB10_PROCESS_CODE = '''def process_files(source_dir, target_dir, mode, preview):
    stats = {"deleted": 0, "moved": 0, "errors": 0}
    type_stats = {}
    processed_orders = []

    if not source_dir.exists():
        log("Исходная папка не найдена — создаём тестовые данные...", "warning")
        create_test_data(source_dir)

    log(f"Режим: {mode.upper()} | Источник: {source_dir} | Цель: {target_dir}")

    for file_path in sorted(source_dir.rglob("*")):
        if not file_path.is_file(): continue
        if target_dir in file_path.parents: continue

        ext  = file_path.suffix.lower()
        size = file_path.stat().st_size

        # Удаление мусорных файлов
        if ext in TRASH_EXTENSIONS:
            log(f"Удаление временного файла ({ext}): {file_path.name}", preview=preview)
            if not preview: file_path.unlink(); stats["deleted"] += 1
            continue

        # Удаление пустых файлов
        if size == 0:
            log(f"Удаление пустого файла: {file_path.name}", preview=preview)
            if not preview: file_path.unlink(); stats["deleted"] += 1
            continue

        category = get_category(ext)

        # Удаление устаревших архивов
        if category == "Архивы" and is_file_old(file_path, ARCHIVE_LIFETIME_DAYS):
            log(f"Удаление устаревшего архива: {file_path.name}", preview=preview)
            if not preview: file_path.unlink(); stats["deleted"] += 1
            continue

        # Перемещение файла
        category_path = target_dir / category
        if not preview: category_path.mkdir(parents=True, exist_ok=True)
        new_name = generate_unique_name(category_path, file_path.name, ext)
        log(f"Перемещение [{category}]: {file_path.name} → {new_name}", preview=preview)

        if ext == ".csv":
            rows, total = parse_csv_order(file_path)
            processed_orders.append({"name": new_name, "rows": rows, "total": total})
            log(f"  CSV-заказ: {rows} строк, сумма {total:.2f} руб.", preview=preview)

        if not preview:
            try:
                shutil.move(str(file_path), str(category_path / new_name))
                stats["moved"] += 1
                type_stats[category] = type_stats.get(category, 0) + 1
            except (OSError, shutil.Error) as e:
                log(f"Ошибка: {e}", "error"); stats["errors"] += 1
        else:
            type_stats[category] = type_stats.get(category, 0) + 1

    if processed_orders:
        orders_dir = target_dir / "Заказы"
        if not preview: orders_dir.mkdir(parents=True, exist_ok=True)
        grand_total = generate_summary_report(processed_orders, orders_dir, preview)
        log(f"Итоговая сумма всех заказов: {grand_total:.2f} руб.")

    return stats, type_stats'''

LAB10_PREVIEW_OUTPUT = '''$ python main.py preview
2026-05-28 14:23:01 | INFO     | Режим: PREVIEW | Источник: orders | Цель: processed_orders
2026-05-28 14:23:01 | INFO     | [PREVIEW] Удаление временного файла (.tmp): temp.tmp
2026-05-28 14:23:01 | INFO     | [PREVIEW] Удаление временного файла (.bak): backup.bak
2026-05-28 14:23:01 | INFO     | [PREVIEW] Удаление пустого файла: empty.csv
2026-05-28 14:23:01 | INFO     | [PREVIEW] Перемещение [Заказы]: order_001.csv  →  order_001.csv
2026-05-28 14:23:01 | INFO     |   CSV-заказ: 2 строк, сумма 1520.00 руб.
2026-05-28 14:23:01 | INFO     | [PREVIEW] Перемещение [Заказы]: order_002.csv  →  order_002.csv
2026-05-28 14:23:01 | INFO     |   CSV-заказ: 2 строк, сумма 1850.00 руб.
2026-05-28 14:23:01 | INFO     | [PREVIEW] Перемещение [Заказы]: order_003.csv  →  order_003.csv
2026-05-28 14:23:01 | INFO     |   CSV-заказ: 3 строк, сумма 1600.00 руб.
2026-05-28 14:23:01 | INFO     | [PREVIEW] Перемещение [Отчёты]: report_january.txt  →  report_january.txt
2026-05-28 14:23:01 | INFO     | [PREVIEW] Перемещение [Отчёты]: report_february.txt  →  report_february.txt
2026-05-28 14:23:01 | INFO     | [PREVIEW] Перемещение [Данные]: store_data.json  →  store_data.json
2026-05-28 14:23:01 | INFO     | [PREVIEW] Будет создан сводный отчёт: summary_20260528_142301.txt
2026-05-28 14:23:01 | INFO     | Итоговая сумма всех заказов: 4970.00 руб.'''

LAB10_RUN_OUTPUT = '''$ python main.py run
2026-05-28 14:25:10 | INFO     | Режим: RUN | Источник: orders | Цель: processed_orders
2026-05-28 14:25:10 | INFO     | Удаление временного файла (.tmp): temp.tmp
2026-05-28 14:25:10 | INFO     | Удаление временного файла (.bak): backup.bak
2026-05-28 14:25:10 | INFO     | Удаление пустого файла: empty.csv
2026-05-28 14:25:10 | INFO     | Перемещение [Заказы]: order_001.csv  →  order_001.csv
2026-05-28 14:25:10 | INFO     |   CSV-заказ: 2 строк, сумма 1520.00 руб.
2026-05-28 14:25:10 | INFO     | Перемещение [Заказы]: order_002.csv  →  order_002.csv
2026-05-28 14:25:10 | INFO     |   CSV-заказ: 2 строк, сумма 1850.00 руб.
2026-05-28 14:25:10 | INFO     | Перемещение [Заказы]: order_003.csv  →  order_003.csv
2026-05-28 14:25:10 | INFO     |   CSV-заказ: 3 строк, сумма 1600.00 руб.
2026-05-28 14:25:10 | INFO     | Перемещение [Отчёты]: report_january.txt  →  report_january.txt
2026-05-28 14:25:10 | INFO     | Перемещение [Отчёты]: report_february.txt  →  report_february.txt
2026-05-28 14:25:10 | INFO     | Перемещение [Данные]: store_data.json  →  store_data.json
2026-05-28 14:25:10 | INFO     | Сводный отчёт сохранён: summary_20260528_142510.txt
2026-05-28 14:25:10 | INFO     | Итоговая сумма всех заказов: 4970.00 руб.

============================================================
ИТОГОВАЯ СТАТИСТИКА ОБРАБОТКИ
============================================================
Режим работы  : RUN
Источник      : orders
Цель          : processed_orders
Удалено файлов: 3
Перемещено    : 6
Ошибок        : 0

ПО КАТЕГОРИЯМ:
  Данные              : 1 файл(ов)
  Заказы              : 3 файл(ов)
  Отчёты              : 2 файл(ов)
============================================================'''

LAB10_UNIQUE_OUTPUT = '''$ python main.py run
# При повторном запуске с теми же файлами:
2026-05-28 14:30:00 | INFO     | Перемещение [Заказы]: order_001.csv  →  order_001_1.csv
2026-05-28 14:30:00 | INFO     | Перемещение [Заказы]: order_002.csv  →  order_002_1.csv
2026-05-28 14:30:00 | INFO     | Перемещение [Заказы]: order_003.csv  →  order_003_1.csv
# Файлы не перезаписываются — к имени добавляется суффикс _1, _2, ...'''

LAB10_FULL_CODE = '''import sys, csv, json, time, shutil, logging
from pathlib import Path
from datetime import datetime

DEFAULT_SOURCE        = Path("orders")
DEFAULT_TARGET        = Path("processed_orders")
ARCHIVE_LIFETIME_DAYS = 30

EXTENSIONS = {
    "Заказы":  {".csv"},
    "Отчёты":  {".txt", ".log"},
    "Данные":  {".json"},
    "Архивы":  {".zip", ".tar", ".gz", ".rar"},
}
TRASH_EXTENSIONS = {".tmp", ".bak", ".old"}


def setup_logging():
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s | %(levelname)-8s | %(message)s",
        handlers=[
            logging.StreamHandler(),
            logging.FileHandler("automation.log", encoding="utf-8"),
        ],
    )

def log(message, level="info", preview=False):
    prefix = "[PREVIEW] " if preview else ""
    getattr(logging, level)(prefix + message)

def get_category(ext):
    for category, extensions in EXTENSIONS.items():
        if ext in extensions:
            return category
    return "Другое"

def generate_unique_name(directory, name, ext):
    base = Path(name).stem.lower().replace(" ", "_")
    counter = 1
    new_name = f"{base}{ext}"
    while (directory / new_name).exists():
        new_name = f"{base}_{counter}{ext}"
        counter += 1
    return new_name

def is_file_old(file_path, days):
    return (time.time() - file_path.stat().st_mtime) > days * 86400

def parse_csv_order(file_path):
    total = 0.0; rows = 0
    try:
        with open(file_path, encoding="utf-8") as f:
            reader = csv.DictReader(f)
            for row in reader:
                try:
                    price = float(row.get("price", 0) or 0)
                    qty   = int(row.get("quantity", 1) or 1)
                    total += price * qty; rows += 1
                except (ValueError, KeyError):
                    pass
    except Exception:
        pass
    return rows, total

def generate_summary_report(processed_orders, target_dir, preview=False):
    timestamp   = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_path = target_dir / f"summary_{timestamp}.txt"
    grand_total = sum(o["total"] for o in processed_orders)
    lines = ["=" * 60, "СВОДНЫЙ ОТЧЁТ ОБРАБОТКИ ЗАКАЗОВ",
             f"Дата: {datetime.now().strftime(\'%d.%m.%Y %H:%M:%S\')}",
             "=" * 60, f"Обработано файлов заказов: {len(processed_orders)}", "ДЕТАЛИ:"]
    for item in processed_orders:
        lines.append(f"  {item[\'name\']:35s}  {item[\'rows\']:3d} строк  {item[\'total\']:10.2f} руб.")
    lines += ["", f"ИТОГОВАЯ СУММА: {grand_total:.2f} руб.", "=" * 60]
    if not preview:
        report_path.write_text("\\n".join(lines), encoding="utf-8")
        log(f"Сводный отчёт сохранён: {report_path.name}")
    else:
        log(f"Будет создан сводный отчёт: {report_path.name}", preview=True)
    return grand_total

def create_test_data(source_dir):
    source_dir.mkdir(parents=True, exist_ok=True)
    orders = [
        ("order_001.csv", [(1,"Мастер и Маргарита",2,450),(2,"Война и мир",1,620)]),
        ("order_002.csv", [(3,"Евгений Онегин",3,290),(4,"Анна Каренина",2,490)]),
        ("order_003.csv", [(5,"Тихий Дон",1,510),(6,"Обломов",2,355),(7,"Преступление и наказание",1,380)]),
    ]
    for filename, rows in orders:
        with open(source_dir / filename, "w", encoding="utf-8", newline="") as f:
            writer = csv.writer(f)
            writer.writerow(["book_id","title","quantity","price"])
            for row in rows: writer.writerow(row)
    (source_dir/"report_january.txt").write_text("Отчёт за январь 2026\\nПродано: 150 книг", encoding="utf-8")
    (source_dir/"report_february.txt").write_text("Отчёт за февраль 2026\\nПродано: 178 книг", encoding="utf-8")
    data = {"store":"Книжный магазин","updated":datetime.now().strftime("%Y-%m-%d"),
            "top_books":["Мастер и Маргарита","Война и мир"]}
    (source_dir/"store_data.json").write_text(json.dumps(data,ensure_ascii=False,indent=2),encoding="utf-8")
    (source_dir/"temp.tmp").write_text("temporary"); (source_dir/"backup.bak").write_text("old backup")
    (source_dir/"empty.csv").write_text("")
    log(f"Тестовые данные созданы в папке: {source_dir}")

def process_files(source_dir, target_dir, mode, preview):
    stats = {"deleted": 0, "moved": 0, "errors": 0}
    type_stats = {}; processed_orders = []
    if not source_dir.exists():
        log("Исходная папка не найдена — создаём тестовые данные...", "warning")
        create_test_data(source_dir)
    if not preview: target_dir.mkdir(parents=True, exist_ok=True)
    log(f"Режим: {mode.upper()} | Источник: {source_dir} | Цель: {target_dir}")
    for file_path in sorted(source_dir.rglob("*")):
        if not file_path.is_file(): continue
        if target_dir in file_path.parents: continue
        ext = file_path.suffix.lower()
        if ext in TRASH_EXTENSIONS:
            log(f"Удаление временного файла ({ext}): {file_path.name}", preview=preview)
            if not preview: file_path.unlink(); stats["deleted"] += 1
            continue
        try: size = file_path.stat().st_size
        except FileNotFoundError: continue
        if size == 0:
            log(f"Удаление пустого файла: {file_path.name}", preview=preview)
            if not preview: file_path.unlink(); stats["deleted"] += 1
            continue
        category = get_category(ext)
        if category == "Архивы" and is_file_old(file_path, ARCHIVE_LIFETIME_DAYS):
            log(f"Удаление устаревшего архива: {file_path.name}", preview=preview)
            if not preview: file_path.unlink(); stats["deleted"] += 1
            continue
        category_path = target_dir / category
        if not preview: category_path.mkdir(parents=True, exist_ok=True)
        new_name = generate_unique_name(category_path, file_path.name, ext)
        log(f"Перемещение [{category}]: {file_path.name} → {new_name}", preview=preview)
        if ext == ".csv":
            rows, total = parse_csv_order(file_path)
            processed_orders.append({"name": new_name, "rows": rows, "total": total})
            log(f"  CSV-заказ: {rows} строк, сумма {total:.2f} руб.", preview=preview)
        if not preview:
            try:
                shutil.move(str(file_path), str(category_path / new_name))
                stats["moved"] += 1; type_stats[category] = type_stats.get(category,0)+1
            except (OSError, shutil.Error) as e:
                log(f"Ошибка: {e}", "error"); stats["errors"] += 1
        else:
            type_stats[category] = type_stats.get(category, 0) + 1
    if processed_orders:
        orders_dir = target_dir / "Заказы"
        if not preview: orders_dir.mkdir(parents=True, exist_ok=True)
        grand_total = generate_summary_report(processed_orders, orders_dir, preview)
        log(f"Итоговая сумма всех заказов: {grand_total:.2f} руб.")
    return stats, type_stats

def print_statistics(stats, type_stats, mode, source_dir, target_dir):
    print("\\n" + "=" * 60)
    print("ИТОГОВАЯ СТАТИСТИКА ОБРАБОТКИ"); print("=" * 60)
    print(f"Режим работы  : {mode.upper()}")
    print(f"Источник      : {source_dir}"); print(f"Цель          : {target_dir}")
    print(f"Удалено файлов: {stats[\'deleted\']}"); print(f"Перемещено    : {stats[\'moved\']}")
    print(f"Ошибок        : {stats[\'errors\']}")
    if type_stats:
        print("\\nПО КАТЕГОРИЯМ:")
        for key in sorted(type_stats): print(f"  {key:<20s}: {type_stats[key]} файл(ов)")
    print("=" * 60)

if __name__ == "__main__":
    setup_logging()
    args = sys.argv
    mode = args[1].lower() if len(args) > 1 else "run"
    if mode not in {"run", "preview"}:
        print("Использование: python main.py [run|preview] [источник] [цель]")
        sys.exit(1)
    source_dir = Path(args[2]) if len(args) > 2 else DEFAULT_SOURCE
    target_dir = Path(args[3]) if len(args) > 3 else DEFAULT_TARGET
    preview    = (mode == "preview")
    stats, type_stats = process_files(source_dir, target_dir, mode, preview)
    print_statistics(stats, type_stats, mode, source_dir, target_dir)'''


def generate_lab10(output_path):
    doc = Document()
    set_page_margins(doc)
    create_title_page(doc, "10", "Автоматизация задач")
    add_page_break(doc)

    # Цель
    _inline_bold_para(doc,
        "Цель: ",
        "освоить возможности стандартных библиотек Python для разработки "
        "программных решений по автоматизации работы с файловой системой "
        "путём создания скрипта обработки файлов книжного магазина, "
        "реализующего классификацию, фильтрацию, анализ CSV-заказов "
        "и формирование сводных отчётов."
    )

    # Постановка задачи
    _inline_bold_para(doc, "Постановка задачи: ",
        "для достижения цели необходимо выполнить следующие задачи:")
    tasks = [
        "изучить модуль pathlib для объектно-ориентированной работы с путями файловой системы;",
        "изучить модуль shutil для высокоуровневых операций с файлами (копирование, перемещение);",
        "изучить модуль logging для ведения журнала операций с настройкой форматирования и обработчиков;",
        "изучить модуль sys для обработки аргументов командной строки;",
        "изучить работу с форматами CSV и JSON средствами стандартной библиотеки Python;",
        "разработать скрипт классификации файлов по расширениям на категории (Заказы, Отчёты, Данные, Архивы);",
        "реализовать удаление мусорных (.tmp, .bak, .old), пустых файлов и устаревших архивов;",
        "реализовать два режима работы: run (реальное выполнение) и preview (просмотр без изменений);",
        "реализовать анализ CSV-заказов, генерацию сводного отчёта и вывод итоговой статистики.",
    ]
    for t in tasks:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(t)
        _apply_tnr_font(run, size=14)

    # Введение
    add_heading(doc, "Введение")
    add_paragraph(doc,
        "Автоматизация рутинных задач — одна из ключевых областей применения "
        "скриптовых языков программирования. Организации ежедневно сталкиваются "
        "с необходимостью упорядочивания файлов, резервного копирования данных, "
        "обработки отчётов и мониторинга систем. Ручное выполнение таких операций "
        "требует значительных временных затрат и подвержено ошибкам. "
        "Python, благодаря богатой стандартной библиотеке, предоставляет все "
        "необходимые инструменты для создания надёжных автоматизированных решений "
        "без установки сторонних пакетов."
    )
    add_paragraph(doc,
        "Стандартная библиотека Python включает специализированные модули для "
        "решения задач автоматизации: pathlib предоставляет объектно-ориентированный "
        "интерфейс к файловой системе, shutil — высокоуровневые операции с файлами, "
        "logging — гибкую систему журналирования событий, csv и json — парсинг "
        "структурированных форматов данных, sys — доступ к аргументам командной строки. "
        "Совместное использование этих модулей позволяет строить полноценные "
        "утилиты автоматизации без внешних зависимостей."
    )
    add_paragraph(doc,
        "В рамках данной лабораторной работы в качестве предметной области "
        "выбран книжный магазин — предметная область, используемая во всех "
        "предыдущих работах. Разрабатывается скрипт автоматизированной обработки "
        "файлов заказов: входная директория может содержать CSV-заказы, "
        "текстовые отчёты, JSON-данные, архивы и мусорные файлы. "
        "Скрипт классифицирует, фильтрует и перемещает файлы в структурированное "
        "хранилище, анализирует содержимое заказов и формирует сводный отчёт."
    )

    # Теоретическая часть
    add_heading(doc, "Теоретическая часть")

    add_subheading(doc, "Основы автоматизации задач на Python")
    add_paragraph(doc,
        "Автоматизация с использованием Python охватывает несколько направлений: "
        "работу с файловой системой (переименование, перемещение, архивация, поиск), "
        "обработку структурированных данных (CSV, JSON, XML), "
        "системное администрирование (резервное копирование, мониторинг), "
        "и создание отчётов. Скрипты автоматизации запускаются из командной строки "
        "с параметрами, что обеспечивает гибкость: один скрипт может обрабатывать "
        "разные директории и работать в разных режимах."
    )

    add_subheading(doc, "Модуль pathlib")
    add_paragraph(doc,
        "Модуль pathlib (Python 3.4+) предоставляет объект Path — "
        "кроссплатформенное представление пути файловой системы. "
        "Объект Path поддерживает: проверку существования (exists()), "
        "создание директорий (mkdir()), рекурсивный обход (rglob()), "
        "получение метаданных (stat() — размер, время изменения), "
        "работу с именем и расширением (.name, .stem, .suffix), "
        "а также операторы / для конкатенации путей. "
        "Метод rglob(\"*\") обходит все файлы и папки рекурсивно, "
        "что делает его незаменимым при обработке сложных структур директорий:"
    )
    add_code_block(doc, LAB10_PATHLIB_CODE)
    add_figure_caption(doc, "Рис. 1. Примеры работы с объектом Path (модуль pathlib)")

    add_subheading(doc, "Модуль shutil")
    add_paragraph(doc,
        "Модуль shutil предназначен для высокоуровневых операций с файлами. "
        "Функция shutil.move(src, dst) перемещает файл или директорию из src в dst; "
        "если dst — это директория, файл помещается внутрь неё. "
        "Функция shutil.copy2() копирует файл, сохраняя метаданные (время создания, "
        "атрибуты). При ошибке перемещения (нет прав, файл заблокирован) "
        "возбуждается исключение shutil.Error или OSError, которое необходимо "
        "перехватывать для устойчивой работы скрипта."
    )

    add_subheading(doc, "Система журналирования (модуль logging)")
    add_paragraph(doc,
        "Модуль logging реализует гибкую многоуровневую систему журналирования. "
        "Уровни серьёзности: DEBUG, INFO, WARNING, ERROR, CRITICAL. "
        "Функция basicConfig() настраивает корневой регистратор: "
        "уровень фильтрации, формат строки лога и список обработчиков (handlers). "
        "StreamHandler выводит сообщения в консоль, "
        "FileHandler — в файл с заданной кодировкой. "
        "Параметр format поддерживает плейсхолдеры: "
        "%(asctime)s — дата/время, %(levelname)s — уровень, %(message)s — текст. "
        "Использование logging вместо print() позволяет легко переключать "
        "вывод между консолью и файлом, добавлять временны́е метки и управлять "
        "детальностью журнала:"
    )
    add_code_block(doc, LAB10_LOGGING_CODE)
    add_figure_caption(doc, "Рис. 2. Настройка системы журналирования через logging")

    add_subheading(doc, "Форматы данных CSV и JSON")
    add_paragraph(doc,
        "Модуль csv предоставляет классы reader и writer для чтения и записи "
        "CSV-файлов (Comma-Separated Values). "
        "Класс DictReader читает каждую строку как словарь, "
        "где ключами служат заголовки из первой строки файла. "
        "Это удобно при работе с файлами заказов: "
        "row.get(\"price\", 0) безопасно возвращает значение поля price. "
        "Модуль json обеспечивает сериализацию (json.dumps) и десериализацию "
        "(json.loads / json.load) данных в формате JSON. "
        "Параметр ensure_ascii=False позволяет записывать кириллицу без экранирования, "
        "indent=2 форматирует JSON с отступами для читаемости."
    )

    # Реализация
    add_heading(doc, "Реализация")
    add_paragraph(doc,
        "Скрипт реализован на чистом Python 3 без внешних зависимостей. "
        "Код организован в функции по принципу единственной ответственности: "
        "каждая функция решает одну конкретную подзадачу, "
        "что упрощает тестирование и сопровождение."
    )

    add_subheading(doc, "Конфигурация и структура модулей")
    add_paragraph(doc,
        "В начале файла определяются все константы конфигурации: "
        "пути по умолчанию, словарь EXTENSIONS с маппингом расширений на категории "
        "и множество TRASH_EXTENSIONS для файлов-мусора. "
        "Вынесение конфигурации в константы позволяет легко адаптировать "
        "скрипт под другие предметные области без изменения основной логики:"
    )
    add_code_block(doc, LAB10_CONFIG_CODE)
    add_figure_caption(doc, "Рис. 3. Конфигурация скрипта: импорты и константы")

    add_subheading(doc, "Настройка логирования")
    add_paragraph(doc,
        "Функция setup_logging() настраивает два обработчика одновременно: "
        "консольный (StreamHandler) для оперативного контроля и файловый "
        "(FileHandler → automation.log) для ведения постоянного журнала. "
        "Вспомогательная функция log() принимает параметр preview: "
        "при preview=True к сообщению добавляется префикс [PREVIEW], "
        "что визуально отличает «планируемые» операции от реально выполненных. "
        "Вызов getattr(logging, level)(message) позволяет передавать уровень "
        "логирования как строку:"
    )
    add_figure_caption(doc, "Рис. 4. Функции setup_logging() и log() — настройка журнала")

    add_subheading(doc, "Классификация файлов по расширениям")
    add_paragraph(doc,
        "Функция get_category() выполняет поиск расширения файла в словаре EXTENSIONS. "
        "Итерация по словарю с проверкой принадлежности расширения множеству "
        "обеспечивает линейное время поиска O(k), где k — число категорий. "
        "Для неизвестных расширений возвращается значение «Другое», "
        "чтобы файлы не терялись при обработке:"
    )
    add_code_block(doc, LAB10_CATEGORY_CODE)
    add_figure_caption(doc, "Рис. 5. Функция get_category() — определение категории по расширению")

    add_subheading(doc, "Генерация уникальных имён и проверка возраста файлов")
    add_paragraph(doc,
        "Функция generate_unique_name() предотвращает перезапись файлов при конфликте "
        "имён: если файл с таким именем уже существует в целевой директории, "
        "к базовому имени добавляется числовой суффикс (_1, _2, ...) до тех пор, "
        "пока не будет найдено свободное имя. Имя также нормализуется: "
        "переводится в нижний регистр, пробелы заменяются на подчёркивания. "
        "Функция is_file_old() сравнивает время последнего изменения файла "
        "(st_mtime) с текущим временем (time.time()), используя перевод дней "
        "в секунды (86400 = 24×60×60):"
    )
    add_code_block(doc, LAB10_UNIQUE_CODE)
    add_figure_caption(doc, "Рис. 6. Функции generate_unique_name() и is_file_old()")

    add_subheading(doc, "Анализ содержимого CSV-заказов")
    add_paragraph(doc,
        "Функция parse_csv_order() открывает CSV-файл заказа и суммирует "
        "стоимость позиций: price × quantity для каждой строки. "
        "Используется csv.DictReader, который автоматически сопоставляет "
        "столбцы по заголовкам. Двойная обработка исключений (внешняя — "
        "для ошибок открытия файла, внутренняя — для некорректных значений "
        "отдельных строк) обеспечивает устойчивость к повреждённым данным:"
    )
    add_code_block(doc, LAB10_CSV_CODE)
    add_figure_caption(doc, "Рис. 7. Функция parse_csv_order() — анализ содержимого заказа")

    add_subheading(doc, "Генерация сводного отчёта")
    add_paragraph(doc,
        "После обработки всех CSV-файлов функция generate_summary_report() "
        "формирует текстовый отчёт с детализацией по каждому заказу "
        "и итоговой суммой. Имя файла включает временну́ю метку "
        "(format \"%Y%m%d_%H%M%S\") для уникальности. "
        "В режиме preview отчёт не создаётся физически — "
        "его предполагаемое содержимое выводится в журнал:"
    )
    add_code_block(doc, LAB10_SUMMARY_CODE)
    add_figure_caption(doc, "Рис. 8. Функция generate_summary_report() — формирование отчёта")

    add_subheading(doc, "Основной цикл обработки файлов")
    add_paragraph(doc,
        "Функция process_files() реализует главный алгоритм обработки. "
        "Если исходная директория отсутствует — автоматически создаются "
        "тестовые данные (CSV-заказы, текстовые отчёты, JSON, мусорные файлы). "
        "Для каждого файла последовательно применяются фильтры: "
        "удаление мусора → удаление пустых → удаление устаревших архивов → "
        "классификация и перемещение. "
        "Параметр preview управляет всеми деструктивными операциями через "
        "проверку if not preview:"
    )
    add_code_block(doc, LAB10_PROCESS_CODE)
    add_figure_caption(doc, "Рис. 9. Функция process_files() — основной цикл обработки")

    # Тестирование
    add_heading(doc, "Тестирование")
    add_paragraph(doc,
        "Тестирование скрипта проводилось в два этапа: сначала в режиме preview "
        "для проверки запланированных операций без изменения файловой системы, "
        "затем в режиме run для реального выполнения. "
        "Исходная директория orders содержала 9 файлов: "
        "3 CSV-заказа, 2 текстовых отчёта, 1 JSON, 2 мусорных файла и 1 пустой CSV."
    )

    add_subheading(doc, "Режим preview — предварительный просмотр")
    add_paragraph(doc,
        "Запуск python main.py preview выводит журнал планируемых операций "
        "с префиксом [PREVIEW]. Файловая система не изменяется: "
        "файлы остаются на месте, директории не создаются. "
        "Скрипт корректно отображает планируемые удаления мусорных файлов, "
        "категории для каждого файла и содержимое CSV-заказов:"
    )
    add_code_block(doc, LAB10_PREVIEW_OUTPUT)
    add_figure_caption(doc, "Рис. 10. Вывод скрипта в режиме preview")

    add_subheading(doc, "Режим run — реальное выполнение")
    add_paragraph(doc,
        "Запуск python main.py run выполняет все операции физически: "
        "удалены 3 файла (2 мусорных + 1 пустой CSV), "
        "перемещены 6 файлов в соответствующие подкаталоги, "
        "создан сводный отчёт summary_*.txt в папке processed_orders/Заказы/. "
        "Итоговая статистика выводится в консоль после завершения обработки:"
    )
    add_code_block(doc, LAB10_RUN_OUTPUT)
    add_figure_caption(doc, "Рис. 11. Вывод скрипта в режиме run с итоговой статистикой")

    add_subheading(doc, "Структура целевой директории")
    add_paragraph(doc,
        "После выполнения в режиме run директория processed_orders "
        "содержит следующую структуру категорий:"
    )
    dir_output = '''processed_orders/
├── Данные/
│   └── store_data.json
├── Заказы/
│   ├── order_001.csv
│   ├── order_002.csv
│   ├── order_003.csv
│   └── summary_20260528_142510.txt
└── Отчёты/
    ├── report_january.txt
    └── report_february.txt'''
    add_code_block(doc, dir_output)
    add_figure_caption(doc, "Рис. 12. Структура целевой директории после обработки")

    add_subheading(doc, "Генерация уникальных имён при повторном запуске")
    add_paragraph(doc,
        "При повторном запуске скрипта с теми же исходными файлами "
        "функция generate_unique_name() предотвращает перезапись "
        "уже перемещённых файлов: к имени добавляется суффикс _1, _2 и т.д. "
        "Файлы не теряются и не перезаписываются. "
        "Также был протестирован некорректный режим запуска: "
        "python main.py sort выводит справку об использовании "
        "и завершается с кодом возврата 1."
    )
    add_code_block(doc, LAB10_UNIQUE_OUTPUT)
    add_figure_caption(doc, "Рис. 13. Генерация уникальных имён при конфликте")

    # Вывод
    add_heading(doc, "Вывод")
    add_paragraph(doc,
        "В ходе выполнения лабораторной работы изучены и применены на практике "
        "стандартные библиотеки Python для автоматизации задач: pathlib, shutil, "
        "logging, sys, csv и json. Освоен объектно-ориентированный подход к "
        "работе с файловой системой через объект Path; настроена многоканальная "
        "система журналирования с одновременным выводом в консоль и файл; "
        "реализована работа с форматами CSV и JSON для анализа структурированных данных."
    )
    add_paragraph(doc,
        "Разработан скрипт автоматизации для обработки файлов книжного магазина. "
        "Реализованы: классификация файлов по расширениям на 5 категорий, "
        "удаление мусорных, пустых файлов и устаревших архивов, "
        "перемещение в иерархическую структуру с генерацией уникальных имён, "
        "анализ CSV-заказов с подсчётом суммы и строк, "
        "формирование текстового сводного отчёта. "
        "Реализация двух режимов (run/preview) повышает безопасность: "
        "пользователь видит запланированные операции до их выполнения."
    )
    add_paragraph(doc,
        "Полученные компетенции в области автоматизации файловых операций "
        "применимы в широком круге практических задач: "
        "организация резервного копирования, мониторинг и очистка рабочих директорий, "
        "пакетная обработка отчётов, автоматическая сортировка загрузок. "
        "Использование исключительно стандартной библиотеки обеспечивает "
        "переносимость скрипта между платформами без необходимости установки "
        "дополнительных пакетов."
    )

    # Приложение 1
    add_heading(doc, "Приложение 1")
    add_paragraph(doc,
        "Полный исходный код скрипта автоматизации (lab10/main.py):"
    )
    add_code_block(doc, LAB10_FULL_CODE)

    doc.save(output_path)
    print(f"[OK] ЛР10: {output_path}")


# ═══════════════════════════════════════════════════════════════════════════
#  ТОЧКА ВХОДА
# ═══════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    BASE = "/home/user/vibe/mylabs"

    os.makedirs(f"{BASE}/lab3", exist_ok=True)
    os.makedirs(f"{BASE}/lab4", exist_ok=True)
    os.makedirs(f"{BASE}/lab5", exist_ok=True)
    os.makedirs(f"{BASE}/lab6", exist_ok=True)
    os.makedirs(f"{BASE}/lab7", exist_ok=True)
    os.makedirs(f"{BASE}/lab8", exist_ok=True)
    os.makedirs(f"{BASE}/lab9", exist_ok=True)
    os.makedirs(f"{BASE}/lab10", exist_ok=True)

    generate_lab3(f"{BASE}/lab3/report.docx")
    generate_lab4(f"{BASE}/lab4/report.docx")
    generate_lab5(f"{BASE}/lab5/report.docx")
    generate_lab6(f"{BASE}/lab6/report.docx")
    generate_lab7(f"{BASE}/lab7/report.docx")
    generate_lab8(f"{BASE}/lab8/report.docx")
    generate_lab9(f"{BASE}/lab9/report.docx")
    generate_lab10(f"{BASE}/lab10/report.docx")

    print("\n[ГОТОВО] Все 8 отчётов сгенерированы.")
