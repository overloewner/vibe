"""
Генератор отчётов для ЛР9 и ЛР10
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ══════════════════════════════════════════════════════════
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# ══════════════════════════════════════════════════════════

def set_page_margins(doc, top=2, bottom=2, left=3, right=1.5):
    sec = doc.sections[0]
    sec.top_margin    = Cm(top)
    sec.bottom_margin = Cm(bottom)
    sec.left_margin   = Cm(left)
    sec.right_margin  = Cm(right)


def _fix_font(run, name="Times New Roman"):
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"),    name)
    rPr.insert(0, rFonts)


def add_paragraph(doc, text, bold=False, italic=False,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  font_size=14, first_indent=True,
                  space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_indent and alignment == WD_ALIGN_PARAGRAPH.LEFT:
        pf.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    _fix_font(run)
    return p


def add_heading(doc, text, font_size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after  = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)
    _fix_font(run)
    return p


def add_subheading(doc, text, font_size=14):
    """Подзаголовок с отступом первой строки."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after  = Pt(2)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)
    _fix_font(run)
    return p


def add_bullet(doc, text, font_size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.left_indent   = Cm(1.25)
    pf.first_line_indent = Cm(0)
    run = p.add_run(f"– {text}")
    run.font.size = Pt(font_size)
    _fix_font(run)
    return p


def add_code_block(doc, code_text, caption=None):
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(4)
        cp.paragraph_format.space_after  = Pt(2)
        r = cp.add_run(caption)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)

    for line in code_text.split("\n"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after  = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.first_line_indent = Cm(0)
        pf.left_indent = Cm(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(10)
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Courier New")
        rFonts.set(qn("w:hAnsi"), "Courier New")
        rFonts.set(qn("w:cs"),    "Courier New")
        rPr.insert(0, rFonts)
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        for side in ["top", "left", "bottom", "right"]:
            bdr = OxmlElement(f"w:{side}")
            bdr.set(qn("w:val"), "single")
            bdr.set(qn("w:sz"), "4")
            bdr.set(qn("w:space"), "1")
            bdr.set(qn("w:color"), "CCCCCC")
            pBdr.append(bdr)
        pPr.append(pBdr)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F5F5F5")
        pPr.append(shd)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Заголовок
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(12)
        _fix_font(run, "Times New Roman")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9E1F2")
        tcPr.append(shd)
    # Данные
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(cell_text))
            run.font.size = Pt(11)
            _fix_font(run, "Times New Roman")
    # Ширины столбцов
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def add_figure_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after  = Pt(6)
    pf.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.italic = True
    _fix_font(run)


def add_page_break(doc):
    doc.add_page_break()


def create_title_page(doc, lab_num, lab_title,
                      student="Аценкова М.В.", group="С22-СИБ"):
    C = WD_ALIGN_PARAGRAPH.CENTER

    def tp(text, bold=False, size=14, sb=0, sa=0):
        p = doc.add_paragraph()
        p.alignment = C
        pf = p.paragraph_format
        pf.space_before = Pt(sb)
        pf.space_after  = Pt(sa)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.first_line_indent = Cm(0)
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        _fix_font(run)
        return p

    tp("МИНОБРНАУКИ РОССИИ", bold=True, size=12)
    tp("Федеральное государственное бюджетное образовательное учреждение высшего образования", size=12)
    tp("НИЖЕГОРОДСКИЙ ГОСУДАРСТВЕННЫЙ ТЕХНИЧЕСКИЙ", bold=True, size=14)
    tp("УНИВЕРСИТЕТ им. Р.Е.АЛЕКСЕЕВА", bold=True, size=14)
    tp("Институт радиоэлектроники и информационных технологий", size=12, sb=4)
    tp("Кафедра «Информационная безопасность вычислительных систем и сетей»", size=12, sb=2)
    tp(f"«{lab_title}»", bold=True, size=14, sb=20)
    tp("ОТЧЕТ", bold=True, size=14, sb=6)
    tp(f"по лабораторной работе №{lab_num}", size=14)
    tp("по дисциплине", size=14)
    tp("Скриптовые языки программирования", bold=True, size=14, sa=20)

    for label, value in [("РУКОВОДИТЕЛЬ:", "Вайнбаум Д.А."),
                          ("СТУДЕНТ:", student)]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after  = Pt(2)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.first_line_indent = Cm(0)
        r1 = p.add_run(f"{label}  ")
        r1.font.size = Pt(14)
        _fix_font(r1)
        r2 = p.add_run(value)
        r2.font.size = Pt(14)
        _fix_font(r2)

    tp(f"Группа: {group}", size=14, sb=4)
    tp("Работа защищена «___» ____________", size=14, sb=10)
    tp("С оценкой ________________________", size=14)
    tp("Нижний Новгород 2026", size=14, sb=20)


# ══════════════════════════════════════════════════════════
# ЛР9 — СОЗДАНИЕ GUI-ПРИЛОЖЕНИЯ
# ══════════════════════════════════════════════════════════

def generate_lab9():
    doc = Document()
    set_page_margins(doc)
    create_title_page(doc, 9, "Создание GUI-приложения")
    add_page_break(doc)

    # ── Цель и задачи ──────────────────────────────────────
    add_heading(doc, "Цель работы")
    add_paragraph(doc,
        "Цель работы: изучение основ разработки приложений с графическим пользовательским "
        "интерфейсом (GUI) и приобретение практических навыков создания desktop-приложений "
        "с использованием библиотеки Tkinter языка программирования Python.")
    add_paragraph(doc,
        "В ходе выполнения работы необходимо решить следующие задачи:")
    for task in [
        "изучить принципы построения графических пользовательских интерфейсов и "
        "особенности событийно-ориентированного программирования;",
        "освоить работу с базовыми элементами интерфейса (виджетами: Label, Entry, Button, "
        "Listbox, Radiobutton, Text, Scrollbar);",
        "изучить способы размещения элементов с использованием менеджеров геометрии "
        "pack() и grid();",
        "разработать приложение с графическим интерфейсом для управления каталогом "
        "книжного магазина;",
        "реализовать обработку пользовательских событий с помощью функций-обработчиков;",
        "провести тестирование разработанного приложения.",
    ]:
        add_bullet(doc, task)

    # ── Теоретические сведения ─────────────────────────────
    add_heading(doc, "Теоретические сведения")
    add_paragraph(doc,
        "Графический пользовательский интерфейс (GUI) представляет собой способ "
        "взаимодействия пользователя с программой посредством визуальных элементов: "
        "окон, кнопок, полей ввода, списков и других компонентов. В отличие от консольных "
        "приложений, GUI делает работу с программой более наглядной и интуитивно понятной.")
    add_paragraph(doc,
        "Ключевой особенностью GUI-приложений является событийно-ориентированное "
        "программирование: вместо последовательного выполнения программа ожидает действий "
        "пользователя (нажатие кнопки, ввод текста, выбор элемента) и реагирует на них "
        "через специальные функции-обработчики.")

    add_subheading(doc, "Виджеты Tkinter")
    add_paragraph(doc,
        "Основу графического интерфейса составляют виджеты — отдельные элементы управления. "
        "Tkinter является стандартной библиотекой Python и не требует дополнительной установки. "
        "К основным виджетам относятся:")
    for item in [
        "Label — метка для отображения текста или изображения;",
        "Button — кнопка для запуска функции-обработчика;",
        "Entry — однострочное поле ввода текста;",
        "Text — многострочное текстовое поле;",
        "Listbox — список с возможностью выбора элемента;",
        "Radiobutton — переключатель для выбора одного из нескольких вариантов;",
        "Combobox (ttk) — выпадающий список;",
        "Scrollbar — полоса прокрутки.",
    ]:
        add_bullet(doc, item)

    add_subheading(doc, "Менеджеры геометрии")
    add_paragraph(doc,
        "Размещение элементов в окне осуществляется с помощью менеджеров геометрии:")
    for item in [
        "pack() — автоматическое последовательное размещение элементов;",
        "grid() — размещение по строкам и столбцам в виде таблицы;",
        "place() — точное позиционирование по пиксельным координатам.",
    ]:
        add_bullet(doc, item)

    add_subheading(doc, "Обработка событий")
    add_paragraph(doc,
        "Привязка обработчиков к действиям пользователя осуществляется двумя способами:")
    for item in [
        "через параметр command: tk.Button(root, command=func) — срабатывает при нажатии;",
        "через метод bind(): widget.bind('<Return>', func) — обрабатывает события клавиатуры "
        "и мыши (например, нажатие клавиши Enter).",
    ]:
        add_bullet(doc, item)

    add_paragraph(doc,
        "Главный цикл обработки событий запускается командой root.mainloop() — без него окно "
        "приложения не будет ожидать действий пользователя.")

    # ── Постановка задачи ──────────────────────────────────
    add_heading(doc, "Постановка задачи")
    add_paragraph(doc,
        "В качестве предметной области выбран книжный магазин, каталог которого разрабатывался "
        "в рамках лабораторных работ №6–8. На основе ранее выбранной предметной области "
        "разрабатывается десктопное GUI-приложение для управления каталогом книг.")
    add_paragraph(doc, "Приложение должно обеспечивать следующий функционал:")
    for item in [
        "отображение полного каталога книг в виде списка;",
        "поиск книг по названию, автору или жанру с переключением режима поиска;",
        "добавление новой книги в каталог с валидацией всех полей;",
        "отображение справочной информации о жанрах каталога;",
        "обработку ошибочного ввода с информированием пользователя через messagebox.",
    ]:
        add_bullet(doc, item)

    # ── Проектирование ─────────────────────────────────────
    add_heading(doc, "Проектирование приложения")
    add_paragraph(doc,
        "Основное окно приложения имеет размер 1000×580 пикселей и разделено на две "
        "функциональные области с помощью контейнеров Frame.")
    add_paragraph(doc,
        "Левая панель (left) — область поиска и отображения результатов — включает:")
    for item in [
        "строку поиска (Entry) с кнопками «Найти» и «✖ Очистить»;",
        "радиокнопки (Radiobutton) для выбора режима поиска: по названию, автору или жанру;",
        "кнопку «Показать все книги» для вывода полного каталога;",
        "Listbox со скроллбаром для отображения результатов поиска;",
        "строку статуса (Label) с количеством найденных/всего книг.",
    ]:
        add_bullet(doc, item)
    add_paragraph(doc,
        "Правая панель (right) — форма добавления книги — включает:")
    for item in [
        "поля ввода (Entry): Название, Автор, Год, Цена;",
        "выпадающий список (Combobox): Жанр;",
        "кнопку «Добавить в каталог»;",
        "текстовый блок (Text) со справочной информацией о жанрах.",
    ]:
        add_bullet(doc, item)

    # ── Реализация ─────────────────────────────────────────
    add_heading(doc, "Реализация")
    add_paragraph(doc,
        "Программная реализация выполнена на языке Python с использованием библиотеки Tkinter. "
        "Полный исходный код приведён в Приложении 1.")

    add_subheading(doc, "Инициализация главного окна")
    add_paragraph(doc,
        "Главное окно создаётся с помощью tk.Tk(). Задаются заголовок, размер и цветовая схема. "
        "Для организации двухколоночной компоновки используется метод columnconfigure() с "
        "параметром weight.")
    add_code_block(doc, """\
root = tk.Tk()
root.title("Книжный магазин — Каталог")
root.geometry("1000x580")
root.resizable(False, False)
root.configure(bg=BG)
root.columnconfigure(0, weight=3)
root.columnconfigure(1, weight=1)""")
    add_figure_caption(doc, "Рис. 1. Фрагмент кода инициализации главного окна")

    add_subheading(doc, "Разбивка на панели")
    add_paragraph(doc,
        "Интерфейс разделён на две независимые области с помощью контейнеров Frame, "
        "размещённых по сетке. Левая панель содержит элементы поиска, правая — форму "
        "добавления. Параметр sticky=\"nsew\" обеспечивает растяжение фреймов по всей "
        "доступной области.")
    add_code_block(doc, """\
left = tk.Frame(root, bg=BG)
left.grid(row=0, column=0, sticky="nsew", padx=15, pady=15)

right = tk.Frame(root, bg=BG, relief="groove", bd=1)
right.grid(row=0, column=1, sticky="nsew", padx=(0, 15), pady=15)""")
    add_figure_caption(doc, "Рис. 2. Фрагмент кода создания панелей")

    add_subheading(doc, "Список результатов со скроллом")
    add_paragraph(doc,
        "Для отображения книг используется виджет Listbox, связанный с вертикальным Scrollbar. "
        "Двусторонняя синхронизация обеспечивается параметром yscrollcommand и методом "
        "command=result_list.yview.")
    add_code_block(doc, """\
sb = tk.Scrollbar(lf)
sb.pack(side="right", fill="y")
result_list = tk.Listbox(lf, font=FONT_MAIN,
                          yscrollcommand=sb.set, height=16,
                          selectbackground=ACCENT)
result_list.pack(fill="both", expand=True)
sb.config(command=result_list.yview)""")
    add_figure_caption(doc, "Рис. 3. Фрагмент кода создания Listbox со Scrollbar")

    add_subheading(doc, "Функция поиска")
    add_paragraph(doc,
        "Функция search_books() считывает текст из поля поиска и режим из переменной "
        "search_mode. По выбранному критерию фильтруется список books, результаты "
        "добавляются в Listbox. Если ничего не найдено — выводится соответствующее "
        "сообщение.")
    add_code_block(doc, """\
def search_books():
    query = search_entry.get().strip().lower()
    mode = search_mode.get()
    result_list.delete(0, tk.END)
    if not query:
        messagebox.showwarning("Внимание", "Введите поисковый запрос.")
        return
    found = [b for b in books
             if query in b[mode].lower()]
    if not found:
        result_list.insert(tk.END, "— Книги не найдены —")
    else:
        for b in found:
            result_list.insert(tk.END,
                f"{b['title']}  |  {b['author']}  |  {b['genre']}  |  "
                f"{b['year']} г.  |  {b['price']:.0f} руб.")
    status_label.config(text=f"Найдено: {len(found)} книг(и)")""")
    add_figure_caption(doc, "Рис. 4. Фрагмент кода функции поиска книг")

    add_subheading(doc, "Функция добавления книги с валидацией")
    add_paragraph(doc,
        "Функция add_book() считывает значения из полей формы и выполняет проверку каждого "
        "поля. Строковые поля проверяются на пустоту. Год публикации преобразуется в int "
        "с проверкой диапазона 1000–2100. Цена преобразуется в float с проверкой знака. "
        "При ошибке валидации выводится messagebox.showerror() и функция прерывается. "
        "После успешной проверки книга добавляется в список books.")
    add_code_block(doc, """\
def add_book():
    title  = title_entry.get().strip()
    author = author_entry.get().strip()
    genre  = genre_var.get().strip()
    year_s  = year_entry.get().strip()
    price_s = price_entry.get().strip()

    if not title or not author or not genre:
        messagebox.showerror("Ошибка",
            "Заполните все обязательные поля.")
        return
    try:
        year = int(year_s)
        if year < 1000 or year > 2100:
            raise ValueError
    except ValueError:
        messagebox.showerror("Ошибка",
            "Год должен быть в диапазоне 1000–2100.")
        return
    try:
        price = float(price_s)
        if price <= 0:
            raise ValueError
    except ValueError:
        messagebox.showerror("Ошибка",
            "Цена — положительное число.")
        return
    books.append({"title": title, "author": author,
                  "genre": genre, "year": year, "price": price})
    messagebox.showinfo("Успешно", f"Книга «{title}» добавлена.")""")
    add_figure_caption(doc, "Рис. 5. Фрагмент кода функции добавления книги")

    # ── Тестирование ───────────────────────────────────────
    add_heading(doc, "Тестирование")
    add_paragraph(doc,
        "На этапе тестирования была выполнена проверка корректности работы разработанного "
        "приложения. Основной целью являлось подтверждение правильности функционирования "
        "интерфейса, обработки пользовательского ввода, а также логики поиска и добавления книг.")

    add_subheading(doc, "Запуск и отображение каталога")
    add_paragraph(doc,
        "После запуска программы отображается главное окно приложения с двумя панелями. "
        "В левой части автоматически выводится полный каталог из 10 книг. "
        "Строка статуса показывает общее количество книг в каталоге.")
    add_figure_caption(doc, "Рис. 6. Главное окно приложения при запуске (полный каталог)")

    add_subheading(doc, "Поиск по различным критериям")
    add_paragraph(doc,
        "Был протестирован поиск по каждому из трёх критериев — названию, автору и жанру. "
        "При поиске по автору «Булгаков» найдено 2 книги: «Мастер и Маргарита» и «Собачье сердце». "
        "При поиске по жанру «Роман-эпопея» найдены: «Война и мир» и «Тихий Дон». "
        "Строка статуса обновляется и отражает количество найденных книг.")
    add_figure_caption(doc, "Рис. 7. Результат поиска по автору «Булгаков»")
    add_figure_caption(doc, "Рис. 8. Результат поиска по жанру «Роман-эпопея»")

    add_subheading(doc, "Добавление книги")
    add_paragraph(doc,
        "Была протестирована форма добавления новой книги. При корректном вводе всех полей "
        "книга добавляется в каталог и появляется сообщение об успехе. "
        "Обновляется выпадающий список жанров и блок справки.")
    add_figure_caption(doc, "Рис. 9. Добавление новой книги в каталог")

    add_subheading(doc, "Обработка ошибочного ввода")
    add_paragraph(doc,
        "Была выполнена проверка обработки ошибочных ситуаций:")
    for case in [
        "пустое поле названия → сообщение «Заполните все обязательные поля»;",
        "год «19аб» (не число) → сообщение «Год должен быть в диапазоне 1000–2100»;",
        "отрицательная цена «-100» → сообщение «Цена — положительное число»;",
        "пустой поисковый запрос → предупреждение «Введите поисковый запрос».",
    ]:
        add_bullet(doc, case)
    add_figure_caption(doc, "Рис. 10. Диалоговое окно при ошибке валидации")

    add_paragraph(doc,
        "Таким образом, в ходе тестирования подтверждена работоспособность, "
        "функциональная и визуальная корректность разработанного приложения. "
        "Все элементы интерфейса функционируют стабильно, обработка ошибок работает корректно.")

    # ── Вывод ──────────────────────────────────────────────
    add_heading(doc, "Вывод")
    add_paragraph(doc,
        "В ходе выполнения лабораторной работы были изучены принципы разработки приложений "
        "с графическим интерфейсом и особенности событийно-ориентированного программирования. "
        "Получены практические навыки работы с базовыми виджетами Tkinter, освоены менеджеры "
        "геометрии pack() и grid().")
    add_paragraph(doc,
        "В рамках работы разработано приложение для управления каталогом книжного магазина. "
        "Реализованы поиск книг по трём критериям с использованием Radiobutton, добавление "
        "новых книг с многоуровневой валидацией входных данных, отображение результатов "
        "в Listbox со Scrollbar, а также справочная панель с информацией о жанрах.")
    add_paragraph(doc,
        "Полученные знания и навыки могут быть применены при разработке более сложных "
        "desktop-приложений с использованием других GUI-фреймворков Python (PyQt, wxPython), "
        "а также составляют основу для дальнейшего изучения разработки пользовательских "
        "интерфейсов.")

    # ── Приложение: исходный код ───────────────────────────
    add_page_break(doc)
    add_heading(doc, "Приложение 1")
    add_paragraph(doc, "Исходный код приложения с графическим интерфейсом (lab9/main.py)",
                  bold=True, first_indent=False)

    src = open(os.path.join(os.path.dirname(__file__), "lab9", "main.py"),
               encoding="utf-8").read()
    add_code_block(doc, src)

    os.makedirs(os.path.join(os.path.dirname(__file__), "lab9"), exist_ok=True)
    path = os.path.join(os.path.dirname(__file__), "lab9", "report.docx")
    doc.save(path)
    print(f"[OK] lab9/report.docx сохранён ({os.path.getsize(path)//1024} KB)")


# ══════════════════════════════════════════════════════════
# ЛР10 — АВТОМАТИЗАЦИЯ ЗАДАЧ
# ══════════════════════════════════════════════════════════

def generate_lab10():
    doc = Document()
    set_page_margins(doc)
    create_title_page(doc, 10, "Автоматизация задач")
    add_page_break(doc)

    # ── Цель и задачи ──────────────────────────────────────
    add_heading(doc, "Цель работы")
    add_paragraph(doc,
        "Цель работы: освоение возможностей языка Python для разработки программных решений, "
        "направленных на автоматизацию рутинных задач обработки файлов и повышение "
        "эффективности работы с файловой системой.")
    add_paragraph(doc, "В ходе выполнения работы необходимо решить следующие задачи:")
    for task in [
        "изучить инструменты работы с файловой системой (модули pathlib, shutil, sys, os);",
        "изучить модуль logging для ведения журнала операций;",
        "изучить работу с форматами CSV и JSON средствами стандартной библиотеки Python;",
        "разработать скрипт автоматизации обработки файлов заказов книжного магазина;",
        "реализовать два режима работы: run (реальное выполнение) и preview (просмотр);",
        "провести тестирование и сформулировать выводы.",
    ]:
        add_bullet(doc, task)

    # ── Теоретическая часть ────────────────────────────────
    add_heading(doc, "Теоретическая часть")
    add_subheading(doc, "Основы автоматизации с использованием Python")
    add_paragraph(doc,
        "Автоматизация — подход, при котором задачи, ранее выполняемые вручную, передаются "
        "программным системам. Основные цели: сокращение ручного труда, минимизация ошибок, "
        "ускорение обработки данных. Python является одним из наиболее популярных языков "
        "для автоматизации благодаря простому синтаксису, кроссплатформенности и богатой "
        "стандартной библиотеке.")
    add_paragraph(doc, "Основные направления автоматизации на Python:")
    for item in [
        "работа с файловой системой: переименование, перемещение, архивация, поиск файлов;",
        "обработка структурированных данных: CSV, JSON, XML;",
        "создание отчётов и обработка результатов;",
        "системное администрирование: резервное копирование, мониторинг, очистка.",
    ]:
        add_bullet(doc, item)

    add_subheading(doc, "Модуль pathlib")
    add_paragraph(doc,
        "Модуль pathlib представляет пути файловой системы в виде объектов класса Path "
        "с собственными методами. Это делает код более читаемым и кроссплатформенным "
        "по сравнению с использованием строк и модуля os.path.")
    add_code_block(doc, """\
from pathlib import Path

source = Path("orders")
for file in source.rglob("*.csv"):   # рекурсивный поиск
    print(file.name, file.suffix)    # имя и расширение

Path("sorted").mkdir(parents=True, exist_ok=True)  # создание
""")
    add_figure_caption(doc, "Рис. 1. Примеры использования pathlib")

    add_subheading(doc, "Модуль shutil")
    add_paragraph(doc,
        "Модуль shutil предназначен для высокоуровневых операций с файлами: копирование "
        "(shutil.copy), перемещение (shutil.move), удаление директорий (shutil.rmtree), "
        "архивация (shutil.make_archive).")

    add_subheading(doc, "Модуль logging")
    add_paragraph(doc,
        "Модуль logging обеспечивает гибкую систему журналирования событий. Поддерживает "
        "уровни: DEBUG, INFO, WARNING, ERROR, CRITICAL. Позволяет направлять вывод "
        "одновременно в консоль и файл через настройку обработчиков (handlers).")
    add_code_block(doc, """\
import logging

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)-8s | %(message)s",
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler("automation.log", encoding="utf-8"),
    ],
)
logging.info("Запуск автоматизации")
logging.warning("Файл не найден: data.csv")
""")
    add_figure_caption(doc, "Рис. 2. Настройка системы логирования")

    add_subheading(doc, "Работа с форматами CSV и JSON")
    add_paragraph(doc,
        "Модуль csv обеспечивает чтение и запись файлов формата CSV (Comma-Separated Values). "
        "Класс csv.DictReader позволяет читать строки как словари, используя заголовки столбцов "
        "в качестве ключей. Модуль json обеспечивает сериализацию/десериализацию объектов "
        "Python в формат JSON с поддержкой кодировки Unicode.")

    add_subheading(doc, "Модуль sys")
    add_paragraph(doc,
        "Модуль sys обеспечивает доступ к параметрам командной строки (sys.argv), управление "
        "завершением программы (sys.exit()), а также взаимодействие с потоками ввода-вывода. "
        "В задачах автоматизации sys часто используется для создания универсальных скриптов, "
        "принимающих параметры при запуске из командной строки.")

    # ── Постановка задачи ──────────────────────────────────
    add_heading(doc, "Постановка задачи")
    add_paragraph(doc,
        "В рамках предметной области книжного магазина, используемой в лабораторных "
        "работах №6–8, разрабатывается скрипт автоматизации обработки файлов заказов. "
        "Задача скрипта — анализировать директорию с входящими файлами, классифицировать "
        "их по типам, удалять мусорные и пустые файлы, а также формировать сводный "
        "отчёт по CSV-заказам.")
    add_paragraph(doc, "Скрипт должен реализовывать следующий функционал:")
    for item in [
        "анализ содержимого исходной директории;",
        "удаление пустых файлов и файлов с расширениями-мусором (.tmp, .bak, .old);",
        "удаление устаревших архивов (старше 30 дней);",
        "классификацию файлов по категориям: Заказы, Отчёты, Данные, Архивы, Другое;",
        "генерацию уникальных имён файлов при конфликте (имя_1, имя_2 и т.д.);",
        "перемещение файлов в целевую структуру каталогов;",
        "анализ содержимого CSV-файлов заказов (подсчёт строк и суммы заказа);",
        "формирование сводного текстового отчёта по всем обработанным заказам;",
        "два режима работы: run (реальные операции) и preview (просмотр без изменений);",
        "ведение журнала операций через модуль logging.",
    ]:
        add_bullet(doc, item)
    add_paragraph(doc,
        "Запуск с параметрами: python main.py [run|preview] [источник] [цель]. "
        "При отсутствии исходной директории автоматически создаются тестовые данные.")

    # ── Описание реализации ────────────────────────────────
    add_heading(doc, "Описание реализации программного решения")
    add_paragraph(doc,
        "Программная реализация выполнена на языке Python с использованием стандартных "
        "библиотек: pathlib (работа с путями), shutil (операции с файлами), sys (аргументы "
        "командной строки), logging (журналирование), csv, json (форматы данных), time "
        "(проверка возраста файлов). Полный исходный код приведён в Приложении 1.")

    add_subheading(doc, "Обработка параметров командной строки")
    add_paragraph(doc,
        "Входные параметры считываются из sys.argv. Первый аргумент — режим работы (run "
        "или preview). При некорректном значении программа выводит справку и завершается "
        "с кодом 1. Пути к директориям используются из аргументов или по умолчанию.")
    add_code_block(doc, """\
args = sys.argv
mode = args[1].lower() if len(args) > 1 else "run"

if mode not in {"run", "preview"}:
    print("Использование: python main.py [run|preview] [источник] [цель]")
    sys.exit(1)

source_dir = Path(args[2]) if len(args) > 2 else DEFAULT_SOURCE
target_dir = Path(args[3]) if len(args) > 3 else DEFAULT_TARGET""")
    add_figure_caption(doc, "Рис. 3. Обработка параметров командной строки")

    add_subheading(doc, "Функция определения категории файла")
    add_paragraph(doc,
        "Функция get_category(ext) принимает расширение файла и возвращает его категорию "
        "на основе словаря EXTENSIONS. Если расширение не найдено — возвращается «Другое».")
    add_code_block(doc, """\
EXTENSIONS = {
    "Заказы":  {".csv"},
    "Отчёты":  {".txt", ".log"},
    "Данные":  {".json"},
    "Архивы":  {".zip", ".tar", ".gz", ".rar"},
}

def get_category(ext):
    for category, extensions in EXTENSIONS.items():
        if ext in extensions:
            return category
    return "Другое"
""")
    add_figure_caption(doc, "Рис. 4. Классификация файлов по расширению")

    add_subheading(doc, "Функция генерации уникального имени")
    add_paragraph(doc,
        "Функция generate_unique_name() предотвращает перезапись существующих файлов. "
        "Имя файла приводится к нижнему регистру, пробелы заменяются подчёркиванием. "
        "При конфликте к имени добавляется числовой суффикс (файл_1, файл_2 и т.д.).")
    add_code_block(doc, """\
def generate_unique_name(directory, name, ext):
    base = Path(name).stem.lower().replace(" ", "_")
    counter = 1
    new_name = f"{base}{ext}"
    while (directory / new_name).exists():
        new_name = f"{base}_{counter}{ext}"
        counter += 1
    return new_name""")
    add_figure_caption(doc, "Рис. 5. Функция генерации уникального имени файла")

    add_subheading(doc, "Анализ CSV-файлов заказов")
    add_paragraph(doc,
        "Функция parse_csv_order() читает CSV-файл заказа с помощью csv.DictReader "
        "и подсчитывает общую сумму по формуле price × quantity для каждой строки. "
        "Возвращает количество строк и итоговую сумму.")
    add_code_block(doc, """\
def parse_csv_order(file_path):
    total = 0.0
    rows  = 0
    with open(file_path, encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            price = float(row.get("price", 0) or 0)
            qty   = int(row.get("quantity", 1) or 1)
            total += price * qty
            rows  += 1
    return rows, total""")
    add_figure_caption(doc, "Рис. 6. Анализ содержимого CSV-файла заказа")

    add_subheading(doc, "Основной цикл обработки файлов")
    add_paragraph(doc,
        "Основной цикл обходит все файлы исходной директории методом rglob(\"*\"). "
        "Для каждого файла последовательно выполняется: пропуск директорий и файлов "
        "целевой папки, удаление мусора и пустых файлов, удаление устаревших архивов, "
        "определение категории, перемещение в целевую папку.")
    add_code_block(doc, """\
for file_path in sorted(source_dir.rglob("*")):
    if not file_path.is_file():
        continue
    if target_dir in file_path.parents:
        continue

    ext  = file_path.suffix.lower()
    size = file_path.stat().st_size

    if ext in TRASH_EXTENSIONS:    # удаляем мусор
        file_path.unlink()
        stats["deleted"] += 1
        continue

    if size == 0:                  # удаляем пустые
        file_path.unlink()
        stats["deleted"] += 1
        continue

    category = get_category(ext)
    category_path = target_dir / category
    category_path.mkdir(parents=True, exist_ok=True)
    new_name = generate_unique_name(category_path, file_path.name, ext)
    shutil.move(str(file_path), str(category_path / new_name))
    stats["moved"] += 1""")
    add_figure_caption(doc, "Рис. 7. Основной цикл обработки файлов")

    add_subheading(doc, "Режим preview")
    add_paragraph(doc,
        "В режиме preview все операции логируются с префиксом [PREVIEW], но реальные "
        "изменения файловой системы не выполняются. Это позволяет пользователю заранее "
        "убедиться в корректности логики работы скрипта перед реальным запуском.")
    add_code_block(doc, """\
def log(message, level="info", preview=False):
    prefix = "[PREVIEW] " if preview else ""
    getattr(logging, level)(prefix + message)

# При перемещении:
if not preview:
    shutil.move(str(file_path), str(category_path / new_name))
    stats["moved"] += 1
else:
    log(f"Перемещение: {file_path.name} → {new_name}", preview=True)""")
    add_figure_caption(doc, "Рис. 8. Реализация режима preview")

    # ── Тестирование ───────────────────────────────────────
    add_heading(doc, "Тестирование")
    add_paragraph(doc,
        "Тестирование проводилось путём последовательного запуска скрипта в обоих режимах. "
        "Для обеспечения воспроизводимости тестов при отсутствии исходной директории "
        "скрипт автоматически создаёт тестовый набор файлов: 3 CSV-заказа, "
        "2 текстовых отчёта, 1 JSON-файл, 2 мусорных файла и 1 пустой файл.")

    add_subheading(doc, "Тестирование режима preview")
    add_paragraph(doc,
        "Первоначально скрипт запускался командой python main.py preview. "
        "В режиме preview все сообщения сопровождаются префиксом [PREVIEW]. "
        "Реальных изменений в файловой системе не происходит, что подтверждается "
        "ручной проверкой состояния директорий после выполнения.")
    add_figure_caption(doc, "Рис. 9. Лог выполнения скрипта в режиме preview")

    add_subheading(doc, "Тестирование режима run")
    add_paragraph(doc,
        "При запуске python main.py run скрипт выполняет реальные операции:")
    for item in [
        "удалены 2 мусорных файла (.tmp, .bak) и 1 пустой CSV-файл;",
        "перемещены 3 CSV-файла заказов в папку processed_orders/Заказы/;",
        "перемещены 2 текстовых файла в папку processed_orders/Отчёты/;",
        "перемещён JSON-файл в папку processed_orders/Данные/;",
        "сформирован сводный отчёт с итоговой суммой заказов 4 970,00 руб.",
    ]:
        add_bullet(doc, item)
    add_figure_caption(doc, "Рис. 10. Лог выполнения скрипта в режиме run")
    add_figure_caption(doc, "Рис. 11. Структура целевой директории после обработки")

    add_subheading(doc, "Тестирование генерации уникальных имён")
    add_paragraph(doc,
        "При повторном запуске скрипта на уже обработанных данных механизм уникальных "
        "имён корректно добавляет суффиксы order_001_1.csv, order_001_2.csv и т.д., "
        "предотвращая перезапись существующих файлов.")
    add_figure_caption(doc, "Рис. 12. Генерация уникальных имён при повторном запуске")

    add_subheading(doc, "Вывод итоговой статистики")
    add_paragraph(doc,
        "После завершения обработки программа выводит сводную статистику: режим работы, "
        "пути директорий, количество удалённых и перемещённых файлов, распределение "
        "по категориям. Одновременно ведётся файл журнала automation.log.")
    add_code_block(doc, """\
============================================================
ИТОГОВАЯ СТАТИСТИКА ОБРАБОТКИ
============================================================
Режим работы  : RUN
Источник      : orders
Цель          : processed_orders
Удалено файлов: 3
Перемещено    : 6
Ошибок        : 0

ПО КАТЕГОРИЯМ:
  Данные              : 1 файл(ов)
  Заказы              : 3 файл(ов)
  Отчёты              : 2 файл(ов)
============================================================""")
    add_figure_caption(doc, "Рис. 13. Итоговая статистика выполнения скрипта")

    add_paragraph(doc,
        "Тестирование также подтвердило корректность обработки некорректного режима запуска: "
        "при передаче неизвестного аргумента (например, «sort») программа выводит справку "
        "об использовании и завершается с кодом возврата 1.")

    # ── Вывод ──────────────────────────────────────────────
    add_heading(doc, "Вывод")
    add_paragraph(doc,
        "В ходе выполнения лабораторной работы были изучены основы автоматизации задач "
        "с использованием стандартных библиотек Python. Освоена работа с модулями pathlib, "
        "shutil, sys, logging, csv и json. Получены практические навыки разработки "
        "автономных скриптов обработки файлов.")
    add_paragraph(doc,
        "В рамках работы разработан скрипт автоматизации для обработки файлов книжного "
        "магазина. Реализованы: классификация файлов по расширениям, удаление мусорных и "
        "пустых файлов, удаление устаревших архивов, перемещение файлов в структуру "
        "категорий с генерацией уникальных имён, анализ CSV-заказов и формирование "
        "сводного отчёта, ведение журнала через logging.")
    add_paragraph(doc,
        "Реализация двух режимов работы (run и preview) повышает безопасность использования "
        "скрипта: пользователь может предварительно просмотреть запланированные действия "
        "без изменения файловой системы. Полученные навыки применимы для автоматизации "
        "широкого круга задач в области системного администрирования и обработки данных.")

    # ── Приложение ─────────────────────────────────────────
    add_page_break(doc)
    add_heading(doc, "Приложение 1")
    add_paragraph(doc, "Исходный код скрипта автоматизации (lab10/main.py)",
                  bold=True, first_indent=False)
    src = open(os.path.join(os.path.dirname(__file__), "lab10", "main.py"),
               encoding="utf-8").read()
    add_code_block(doc, src)

    os.makedirs(os.path.join(os.path.dirname(__file__), "lab10"), exist_ok=True)
    path = os.path.join(os.path.dirname(__file__), "lab10", "report.docx")
    doc.save(path)
    print(f"[OK] lab10/report.docx сохранён ({os.path.getsize(path)//1024} KB)")


# ══════════════════════════════════════════════════════════
if __name__ == "__main__":
    generate_lab9()
    generate_lab10()
    print("\nВсе отчёты сгенерированы.")
